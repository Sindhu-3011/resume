from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_from_directory, abort, send_file
import psycopg2
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass
from psycopg2.extras import RealDictCursor
from pathlib import Path
import re
import zipfile
import json
import logging
import os
import urllib.request
import urllib.error
from contextlib import contextmanager
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import easyocr
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import pymupdf as _pymupdf_check
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# ── Tesseract OCR binary path (new code — do not modify existing code) ────────
# winget installs to AppData\Local on per-user basis; set cmd path explicitly
# so pytesseract can find the binary regardless of PATH.
try:
    import pytesseract as _pytess_check
    _TESSERACT_PATHS = [
        r"C:\Users\sindhu.sundara\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _tpath in _TESSERACT_PATHS:
        if os.path.isfile(_tpath):
            _pytess_check.pytesseract.tesseract_cmd = _tpath
            break
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "resume-profile-secret-key")

@app.after_request
def set_cache_control(response):
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_FOLDER = BASE_DIR / "uploads"

# PostgreSQL connection — set DATABASE_URL in .env or environment.
# Default assumes a local "resume_profiles" database owned by the current OS user.
DATABASE_URL = os.environ.get(
    "DATABASE_URL",
    "postgresql://postgres:postgres@localhost:5432/resume_profiles",
)
UPLOAD_FOLDER.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx"}
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

SECTION_ALIASES = {
    # ── Identity / header ─────────────────────────────────────────────────────
    "summary": [
        "summary", "professional summary", "profile summary",
        "executive summary", "career summary", "career profile",
        "professional profile", "personal profile",
        "professional overview", "career overview", "overview",
        "objective", "career objective", "professional objective",
        "job objective", "employment objective", "work objective",
        "position objective", "career aim", "professional aim",
        "about", "about me", "about myself",
        "personal statement", "professional statement", "introduction",
        "professional synopsis", "career synopsis", "synopsis",
        "brief profile", "candidate profile", "career snapshot",
        "profile overview", "professional background summary",
        "qualifications summary", "summary of qualifications",
        "highlights of qualifications",
        "professional bio", "bio", "personal bio", "biography",
        "professional goals", "career goals",
        "professional mission", "mission statement",
        "at a glance", "professional snapshot",
        "career brief", "professional brief",
        "personal overview", "personal summary", "work summary",
        "candidate overview", "professional profile summary",
        "brief about me", "career statement", "profile",
        "value proposition", "professional pitch",
        "personal pitch", "candidate summary",
        "skills summary profile", "background",
        # abbreviated forms: "Prof. Summary" → "prof summary"
        "prof summary", "exec summary", "prof profile",
        "career obj", "prof objective",
    ],

    # ── Skills / competencies ─────────────────────────────────────────────────
    "skills": [
        "skills", "skill set", "skillset", "skills summary",
        "technical skills", "core skills", "key skills",
        "professional skills", "functional skills",
        "competencies", "core competencies", "key competencies",
        "professional competencies", "technical competencies",
        "skills and competencies", "competency profile",
        "technologies", "tools and technologies", "tools",
        "technical tools", "software tools", "tools and software",
        "tools and applications", "tools applications", "tools and skills",
        "tools used", "tools technologies",
        "areas of expertise", "area of expertise",
        "expertise", "domain expertise", "domain skills",
        "technical expertise", "professional expertise",
        "technical proficiencies", "proficiencies",
        "applications summary", "application summary", "applications",
        "it skills", "skills overview", "technical skill set",
        "skills and frameworks", "skills and technologies",
        "technical knowledge", "knowledge and skills",
        "skills and expertise", "key technical skills",
        "technology skills",
        "soft skills", "interpersonal skills", "personal skills",
        "transferable skills", "hard skills",
        "abilities", "capabilities", "technical abilities",
        "frameworks and tools", "programming skills",
        "languages and tools", "platforms and tools",
        "technologies and tools",
        "it proficiencies", "technical profile",
        "expertise areas", "areas of proficiency",
        "specializations", "specialties", "specialisation",
        "core technical skills", "relevant skills",
        # abbreviated / punctuated forms normalized by canonical_section_name regex
        # e.g. "I.T. Skills" → "i t skills", "Tech. Skills" → "tech skills"
        "i t skills", "tech skills", "i t proficiencies",
        "tech proficiencies", "tech expertise",
    ],

    # ── Work experience ───────────────────────────────────────────────────────
    "experience": [
        "EXPERIENCE",
        "Experience", "WORK EXPERIENCE", "Work Experience",
        "experience", "work experience", "professional experience",
        "employment history", "career history", "work history",
        "employment", "professional background",
        "job experience", "relevant experience", "industry experience",
        "work experience details", "experience summary",
        "employment details", "professional history",
        "career experience", "professional work experience",
        "work experience summary", "key experiences",
        "work details", "career details", "employment record",
        "professional record", "experience details",
        "professional engagements", "roles and responsibilities",
        "current experience", "past experience",
        "career overview details",
        "work profile", "working experience",
        "job history", "positions held",
        "appointments", "professional appointments",
        "corporate experience", "it experience",
        "domain experience",
        "relevant work experience", "key work experience",
        "career work history",
        # abbreviated forms: "Prof. Experience" → "prof experience"
        "prof experience", "prof work experience",
        "work exp", "professional exp",
    ],

    # ── Internships ───────────────────────────────────────────────────────────
    "internships": [
        "internship", "internships",
        "internship experience", "internship history",
        "industrial training", "industry training",
        "vocational training",
        "summer internship", "summer training",
        "academic internship", "intern experience",
        "trainee experience", "apprenticeship",
        "co op experience", "coop experience",
        "cooperative education", "practicum",
        "field experience", "work placement",
        "placement experience", "industrial placement",
        "sandwich year", "year in industry",
        "graduate training", "graduate placement",
        "part time experience", "student work experience",
    ],

    # ── Education ─────────────────────────────────────────────────────────────
    "education": [
        "EDUCATION",
        "education", "academic background", "qualifications",
        "academic qualification", "academics",
        "educational background", "academic credentials",
        "educational qualifications", "academic details",
        "educational details", "academic qualifications",
        "education details", "qualification details",
        "academic history", "academic information",
        "educational information", "scholastic details",
        "scholastic background", "educational profile",
        "education and training",
        "academic record", "educational record",
        "academic profile", "studies",
        "degrees", "degrees earned", "degrees obtained",
        "formal education", "academic education",
        "educational credentials",
        "examination details", "scholastic record",
        "qualifications and education", "degree information",
        "academic qualifications and education",
    ],

    # ── Certifications ────────────────────────────────────────────────────────
    "certifications": [
        "certifications", "certificates", "certification",
        "professional certifications", "it certifications",
        "technical certifications",
        "certifications and credentials", "credentials",
        "professional credentials",
        "industry certifications",
        "certified", "certified courses",
        "certification summary",
        "licenses and certifications",
        "certificates and licenses",
        "professional development",
        "courses and certifications",
        "professional courses",
        "online courses",
        "certification and training",
        "certifications and training",
        "professional development certifications",
        "microsoft certifications", "aws certifications",
        "google certifications", "oracle certifications",
        "cisco certifications",
    ],

    # ── Training / workshops ──────────────────────────────────────────────────
    "training": [
        "training", "trainings", "training and development",
        "professional training", "technical training",
        "corporate training", "training attended",
        "training programs", "training and courses",
        "courses", "workshops", "seminars",
        "workshops and seminars", "seminars and workshops",
        "training and workshops", "training and seminars",
        "continuing education", "continuing professional development",
        "cpd",
        "professional development courses",
        "training completed", "programmes attended",
        "programs attended",
    ],

    # ── Licenses ──────────────────────────────────────────────────────────────
    "licenses": [
        "licenses", "licence", "licences", "license",
        "professional licenses", "licenses and permits",
        "state licenses", "federal licenses",
        "regulatory licenses",
        "licenses and registrations",
        "registration", "registrations",
        "professional registrations",
        "license and certification",
        "professional license",
    ],

    # ── Projects ──────────────────────────────────────────────────────────────
    "projects": [
        "project", "projects", "key projects", "project experience", "assignments",
        "notable projects", "project highlights", "academic projects",
        "personal projects", "relevant project experience",
        "project details", "project organizational details",
        "organizational details", "relevant project organizational details",
        "project and organizational details",
        "project organizational details roles and responsibilities",
        "client details", "relevant projects",
        "project summary", "projects summary", "project overview",
        "major projects", "project work", "projects worked",
        "key projects worked", "significant projects", "project details summary",
        "live projects", "work projects", "project descriptions",
        "side projects", "open source projects",
        "portfolio", "portfolio projects",
        "technical projects", "professional projects",
        "freelance projects", "contract projects",
        "client projects", "project portfolio",
    ],

    # ── Achievements / accomplishments ────────────────────────────────────────
    "achievements": [
        "achievements", "achievement", "key achievements",
        "professional achievements", "career achievements",
        "awards and achievements", "achievements and awards",
        "accomplishments", "key accomplishments",
        "professional accomplishments", "notable accomplishments",
        "major accomplishments", "career accomplishments",
        "honors", "honours", "academic honors", "academic honours",
        "distinction", "distinctions",
        "recognition", "accolades",
        "performance highlights", "significant achievements",
        "academic achievements",
        "notable contributions",
    ],

    # ── Awards ────────────────────────────────────────────────────────────────
    "awards": [
        "awards", "award", "awards and recognition",
        "recognition and awards", "honors and recognition",
        "honours and recognition", "recognition and honors",
        "prizes", "prizes and awards",
        "scholarships", "fellowships",
        "scholarships and awards",
        "grants", "grants and awards",
        "commendations", "merits",
        "industry awards", "professional awards",
    ],

    # ── Publications ──────────────────────────────────────────────────────────
    "publications": [
        "publications", "publication", "published works",
        "published papers", "papers",
        "journal publications", "journal articles",
        "articles", "research publications",
        "peer reviewed publications", "peer reviewed articles",
        "technical publications",
        "books", "book chapters", "chapters",
        "conference papers", "conference publications",
        "technical papers",
        "white papers", "whitepapers",
        "case studies",
        "authored works", "writings", "written works",
        "papers and publications", "published research",
        "scholarly publications", "academic publications",
    ],

    # ── Research ──────────────────────────────────────────────────────────────
    "research": [
        "research", "research experience", "research work",
        "research background", "research and development",
        "r and d", "research projects",
        "research interests", "areas of research",
        "research contributions",
        "thesis", "dissertation", "research thesis",
        "academic research", "applied research",
        "research activities", "research assignments",
        "research summary", "research overview",
        "doctoral research", "postdoctoral research",
    ],

    # ── Patents ───────────────────────────────────────────────────────────────
    "patents": [
        "patents", "patent", "intellectual property",
        "patents filed", "patents granted",
        "patents and trademarks", "inventions",
        "innovations", "filed patents",
        "patent applications", "utility patents",
    ],

    # ── Conferences / presentations ───────────────────────────────────────────
    "conferences": [
        "conferences", "conference", "conference presentations",
        "presentations", "presentations and talks",
        "talks", "speaking engagements",
        "conference attendance", "conferences attended",
        "seminars attended",
        "workshops attended",
        "conferences and events", "events",
        "keynote presentations",
        "panel discussions",
        "posters", "poster presentations",
        "conference sessions", "invited talks",
        "technical talks",
    ],

    # ── Volunteer / community ─────────────────────────────────────────────────
    "volunteer": [
        "volunteer experience", "volunteering", "volunteer",
        "volunteer work", "community service",
        "community involvement", "community engagement",
        "social work", "charity work",
        "non profit experience", "nonprofit experience",
        "voluntary work", "voluntary experience",
        "pro bono work", "civic engagement",
        "community activities", "philanthropy",
        "service activities",
        "outreach", "community outreach",
        "social impact", "corporate social responsibility",
        "ngo experience",
    ],

    # ── Leadership ────────────────────────────────────────────────────────────
    "leadership": [
        "leadership", "leadership experience", "leadership roles",
        "leadership and management", "management experience",
        "team leadership", "leadership positions",
        "positions of responsibility", "leadership activities",
        "leadership and extracurricular",
        "leadership and involvement",
        "committee experience", "committee positions",
        "board experience", "board positions",
        "officer positions", "club officer",
        "leadership summary", "leadership and service",
        "management and leadership",
    ],

    # ── Extracurricular / activities / memberships ────────────────────────────
    "extracurricular": [
        "extracurricular activities", "extracurricular",
        "activities", "co curricular activities", "co curricular",
        "college activities", "campus activities",
        "student activities", "club activities",
        "student organizations", "organizations",
        "clubs and organizations", "clubs and activities",
        "university activities",
        "campus involvement", "student involvement",
        "affiliations", "professional affiliations",
        "memberships", "professional memberships",
        "association memberships", "associations",
        "clubs", "society memberships", "societies",
        "professional associations",
    ],

    # ── Languages ─────────────────────────────────────────────────────────────
    "languages": [
        "languages", "language", "language skills",
        "language proficiency", "linguistic skills",
        "languages known", "languages spoken",
        "known languages", "spoken languages",
        "communication languages",
        "foreign languages", "second languages",
        "bilingual", "multilingual",
        "natural languages", "language abilities",
        "language competencies",
    ],

    # ── Interests / hobbies ───────────────────────────────────────────────────
    "interests": [
        "interests", "areas of interest",
        "professional interests", "career interests",
        "personal interests",
        "interests and hobbies", "hobbies and interests",
        "personal interests and hobbies",
        "hobbies and activities",
        "hobbies", "hobby", "hobbies and pastimes",
        "pastimes", "activities and hobbies",
        "leisure activities", "leisure interests",
        "personal activities", "other interests",
        "recreational activities",
    ],

    # ── References ────────────────────────────────────────────────────────────
    "references": [
        "references", "reference", "professional references",
        "character references", "personal references",
        "referees", "referee", "references available",
        "references available upon request",
        "available upon request",
        "references upon request",
        "references furnished upon request",
        "reference details",
        "testimonials",
    ],

    # ── Contact information ───────────────────────────────────────────────────
    "contact": [
        "contact", "contact information", "contact details",
        "contact info", "personal details", "personal information",
        "personal data", "contact data",
        "personal contact", "address",
        "contact and personal details",
        "basic information", "basic details",
        "personal particulars", "particulars",
        "general information", "candidate information",
        "applicant information",
    ],

    # ── Social / online presence ──────────────────────────────────────────────
    "social_links": [
        "social links", "social media", "social profiles",
        "online profiles", "online presence",
        "digital profiles", "web presence",
        "linkedin", "github", "online portfolio",
        "digital portfolio", "portfolio links",
        "professional links", "professional profiles",
        "links", "websites", "website",
        "web profiles", "internet profiles",
        "social media profiles",
    ],

    # ── Career highlights ─────────────────────────────────────────────────────
    "career_highlights": [
        "career highlights", "highlights",
        "career accomplishments", "professional highlights",
        "core achievements", "top achievements",
        "key contributions", "key results",
        "significant contributions", "performance summary",
        "executive highlights",
        "key career highlights", "notable highlights",
        "career summary highlights",
    ],
}

# The 6 sections shown in the UI / stored in the DB.
_CORE_SECTIONS = ("summary", "skills", "experience", "education", "certifications", "projects")

# Maps every extended section to a core bucket (None = discard).
_SECTION_FOLD = {
    "internships":       "experience",
    "training":          "certifications",
    "licenses":          "certifications",
    "achievements":      "certifications",
    "awards":            "certifications",
    "publications":      "projects",
    "research":          "experience",
    "patents":           "certifications",
    "conferences":       "certifications",
    "volunteer":         "experience",
    "leadership":        "experience",
    "extracurricular":   None,
    "languages":         "skills",
    "interests":         None,
    "references":        None,
    "contact":           None,
    "social_links":      None,
    "career_highlights": "summary",
}

# ── DB context manager ────────────────────────────────────────────────────────

class _PgConn:
    """Thin wrapper that gives a psycopg2 connection a sqlite3-style conn.execute() API."""

    def __init__(self, raw):
        self._raw = raw

    def execute(self, sql, params=None):
        cur = self._raw.cursor(cursor_factory=RealDictCursor)
        cur.execute(sql, params if params is not None else ())
        return cur

    def commit(self):
        self._raw.commit()

    def rollback(self):
        self._raw.rollback()

    def close(self):
        self._raw.close()


@contextmanager
def db_conn():
    raw = psycopg2.connect(DATABASE_URL)
    conn = _PgConn(raw)
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


# ── Helpers ───────────────────────────────────────────────────────────────────

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def slugify(text):
    text = (text or "profile").lower().strip()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-") or "profile"


def unique_slug(conn, base_slug, resume_id=None):
    slug = slugify(base_slug)
    candidate = slug
    counter = 2
    while True:
        if resume_id:
            existing = conn.execute(
                "SELECT id FROM resume WHERE slug = %s AND id != %s", (candidate, resume_id)
            ).fetchone()
        else:
            existing = conn.execute("SELECT id FROM resume WHERE slug = %s", (candidate,)).fetchone()
        if not existing:
            return candidate
        candidate = f"{slug}-{counter}"
        counter += 1


def sync_skills(conn, resume_id, skills_text):
    """Populate the normalised resume_skill table from the raw skills text blob."""
    conn.execute("DELETE FROM resume_skill WHERE resume_id = %s", (resume_id,))
    if not skills_text:
        return
    seen = set()
    for line in skills_text.splitlines():
        for part in re.split(r"[,|•\n]", line):
            skill = part.strip().strip("•- ")
            if skill and len(skill) > 1 and "://" not in skill and skill.lower() not in seen:
                seen.add(skill.lower())
                conn.execute(
                    "INSERT INTO resume_skill (resume_id, skill) VALUES (%s, %s)", (resume_id, skill)
                )


# ── PDF / DOCX text extraction ────────────────────────────────────────────────

# Cached EasyOCR reader — initialised once, reused on every subsequent call.
_ocr_reader = None

def _get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        logger.info("Initialising EasyOCR reader (first call only)…")
        _ocr_reader = easyocr.Reader(["en"], gpu=False)
    return _ocr_reader


def _ocr_contact_strip(path):
    """OCR only the top 25 % of page 1 to recover email/phone from icon headers.
    Returns a short string with whatever text EasyOCR finds in that strip.
    """
    try:
        import pymupdf as fitz
        import numpy as np
        doc = fitz.open(str(path))
        pg  = doc[0]
        # Clip to top quarter of the page (where contact bars live)
        clip = fitz.Rect(0, 0, pg.rect.width, pg.rect.height * 0.25)
        mat  = fitz.Matrix(150 / 72, 150 / 72)   # 150 DPI — fast enough for OCR
        pix  = pg.get_pixmap(matrix=mat, clip=clip)
        img  = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
        doc.close()
        reader  = _get_ocr_reader()
        results = reader.readtext(img)
        return "\n".join(t for (_, t, c) in results if c > 0.2)
    except Exception as e:
        logger.debug(f"OCR contact strip failed: {e}")
        return ""


def _detect_two_col_split(page):
    """Return the x midpoint of the gap between two column clusters, or None.

    Uses block bounding boxes (coarser but reliable) to detect whether the page
    content clearly lives in two separate vertical strips.  Conditions:
      - ≥ 3 blocks whose right edge is left of the midpoint  (left column)
      - ≥ 3 blocks whose left  edge is right of the midpoint (right column)
      - ≤ 2 blocks that span across the midpoint (title / contact header)
      - A real gap exists: min x0 of right blocks > max x1 of left blocks
    """
    blocks = page.get_text("blocks")
    pw  = page.rect.width
    mid = pw / 2
    margin = pw * 0.08   # 8 % — tolerates slight layout asymmetry

    left     = [b for b in blocks if b[2] < mid - margin]
    right    = [b for b in blocks if b[0] > mid + margin]
    spanning = [b for b in blocks if b[0] < mid - margin and b[2] > mid + margin]

    if len(left) < 3 or len(right) < 3 or len(spanning) > 2:
        return None

    max_left_x1  = max(b[2] for b in left)
    min_right_x0 = min(b[0] for b in right)
    if min_right_x0 <= max_left_x1:
        return None   # columns overlap — not a clean 2-col layout

    return (max_left_x1 + min_right_x0) / 2


def _detect_col_gutter_words(page):
    """Find a two-column split via a vertical word gutter, or None.

    Block-based detection (`_detect_two_col_split`) misses layouts where wide
    bullet lines in the left column extend past the page midpoint — the blocks
    "span" the centre and the heuristic bails.  This word-level pass instead looks
    for a vertical strip in the middle of the page that almost no word box crosses:
    a true column gutter.  Robust because a single-column page has lines spanning
    the centre, so every candidate split is crossed by many words.
    """
    words = page.get_text("words")
    if len(words) < 40:
        return None

    pw = page.rect.width
    lo, hi = pw * 0.33, pw * 0.67          # only hunt for a gutter in the middle third
    n = len(words)

    best_x, best_cross = None, None
    x = lo
    while x <= hi:
        cross = sum(1 for w in words if w[0] < x < w[2])
        if best_cross is None or cross < best_cross:
            best_cross, best_x = cross, x
        x += 2.0

    if best_x is None:
        return None
    # The gutter must be almost empty (≤1 % of words straddle it).
    if best_cross > max(2, n * 0.01):
        return None
    # Both columns must hold a real share of the content.
    left  = sum(1 for w in words if w[2] <= best_x)
    right = sum(1 for w in words if w[0] >= best_x)
    if left < n * 0.15 or right < n * 0.15:
        return None

    # Final guard: in a true 2-column layout each line lives mostly in one
    # column.  On a single-column page the "gutter" is just a ragged
    # right-margin gap — most lines still have words on BOTH sides.
    # If >25 % of y-level buckets have words on both sides, reject the split.
    _y_tol = 6.0
    _y_sides = {}
    for w in words:
        b = round(w[1] / _y_tol)
        prev = _y_sides.get(b, [False, False])
        _y_sides[b] = [prev[0] or w[2] <= best_x, prev[1] or w[0] >= best_x]
    _dual = sum(1 for v in _y_sides.values() if v[0] and v[1])
    if _dual > 0.25 * len(_y_sides):
        return None

    return best_x


def _words_to_text(raw_words):
    """Convert a flat list of PyMuPDF word tuples to a text string.

    Words on the same y-level (±5 pt) form a row; wide horizontal gaps (> 30 pt)
    within a row indicate sub-columns (e.g. company | role | dates).
    Section headings are always emitted on their own line so the section parser
    can pick them up reliably.
    """
    if not raw_words:
        return ""

    raw_words = sorted(raw_words, key=lambda w: w[1])   # top-to-bottom

    Y_TOL   = 5
    COL_GAP = 30

    def emit_row(row):
        if not row:
            return []
        row.sort()
        groups = [[row[0]]]
        for i in range(1, len(row)):
            if row[i][0] - row[i - 1][1] > COL_GAP:
                groups.append([])
            groups[-1].append(row[i])

        if len(groups) == 1:
            return [" ".join(wd for _, _, wd in groups[0])]

        gtexts = [" ".join(wd for _, _, wd in g) for g in groups]

        if canonical_section_name(gtexts[0]):
            return [gtexts[0]] + gtexts[1:]

        sec_idxs = [i for i in range(1, len(groups)) if canonical_section_name(gtexts[i])]
        if sec_idxs:
            sec_set = set(sec_idxs)
            non_sec = []
            for i, g in enumerate(groups):
                if i not in sec_set:
                    non_sec.extend(g)
            out = [gtexts[i] for i in sec_idxs]
            if non_sec:
                non_sec.sort()
                out.insert(0, " ".join(wd for _, _, wd in non_sec))
            return out

        all_words = [w for g in groups for w in g]
        all_words.sort()
        return [" ".join(wd for _, _, wd in all_words)]

    lines = []
    cur_row, cur_y = [], None
    for w in raw_words:
        x0, y0, x1, y1, word = w[0], w[1], w[2], w[3], w[4]
        if not word.strip():
            continue
        if cur_y is None or abs(y0 - cur_y) <= Y_TOL:
            cur_row.append((x0, x1, word))
            cur_y = cur_y if cur_y is not None else y0
        else:
            lines.extend(emit_row(cur_row))
            cur_row, cur_y = [(x0, x1, word)], y0
    lines.extend(emit_row(cur_row))
    return "\n".join(lines)


def _pdfplumber_two_col_text(path):
    """Extract text from a sidebar-style 2-column PDF using pdfplumber crops.

    For each page:
      1. Collect all word bounding boxes via extract_words().
      2. Find the X position in the 20-65% range where the fewest word boxes
         cross — this is the column gutter.
      3. Crop the page at that X, extract each side independently.
      4. Concatenate left-column text then right-column text so section
         headings (KEY SKILLS on left, WORK EXPERIENCE on right) appear in the
         correct order for find_sections().

    Returns the combined multi-page string, or None if pdfplumber is
    unavailable, the file is not a PDF, or no 2-column layout is detected.
    """
    if not HAS_PDFPLUMBER:
        return None
    try:
        import pdfplumber
        pages_text = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                # extract_words() with no extra kwargs — compatible with all versions
                try:
                    words = page.extract_words()
                except Exception:
                    pages_text.append(page.extract_text() or "")
                    continue

                if not words or len(words) < 15:
                    pages_text.append(page.extract_text() or "")
                    continue

                pw = float(page.width)
                ph = float(page.height)
                n  = len(words)

                # Scan x from 20 % to 65 % of page width for the column gutter
                lo, hi = pw * 0.20, pw * 0.65
                best_x, best_cross = None, n + 1
                x = lo
                while x <= hi:
                    cross = sum(
                        1 for w in words
                        if float(w.get('x0', 0)) < x < float(w.get('x1', 0))
                    )
                    if cross < best_cross:
                        best_cross, best_x = cross, x
                    x += 2.0

                # Allow up to 5% of words to straddle the gutter.
                # Raise tolerance to 10% when this looks like a sidebar layout
                # (right column has 2x+ more words than left): centered headings
                # in the right column can straddle a nearby gutter.
                _left_est  = sum(1 for w in words if float(w.get('x1', 0)) <= best_x)
                _right_est = sum(1 for w in words if float(w.get('x0', 0)) >= best_x)
                _cross_limit = (max(8, n * 0.10)
                                if _right_est > 2 * max(_left_est, 1)
                                else max(4, n * 0.05))
                logger.debug(
                    f"pdfplumber gutter scan: best_x={best_x:.1f} best_cross={best_cross} "
                    f"limit={_cross_limit:.1f} left_est={_left_est} right_est={_right_est}"
                )
                if best_x is None or best_cross > _cross_limit:
                    logger.debug("pdfplumber gutter rejected by cross-count — using extract_text()")
                    pages_text.append(page.extract_text() or "")
                    continue

                # Right column must be substantial (>=8%). Left can be as thin as 3%
                # because on page 2+ of sidebar PDFs the sidebar shrinks to just a few
                # personal-detail lines while the main content fills the right column.
                left_n  = sum(1 for w in words if float(w.get('x1', 0)) <= best_x)
                right_n = sum(1 for w in words if float(w.get('x0', 0)) >= best_x)
                if left_n < n * 0.03 or right_n < n * 0.08:
                    pages_text.append(page.extract_text() or "")
                    continue

                # Final guard: in a true 2-column layout each y-level has words on
                # one side only. On a single-column page most lines span both sides
                # (ragged right margin). If >40% of y-buckets are dual-side → single column.
                # Threshold is 40% (not 25%) because sidebar+main layouts with a photo
                # at the top-left have the top half populated only on the right, while
                # the bottom half has both columns — this legitimately raises dual-Y to ~30%.
                _y_tol = 6.0
                _y_sides = {}
                for w in words:
                    b = round(float(w.get('top', 0)) / _y_tol)
                    prev = _y_sides.get(b, [False, False])
                    _y_sides[b] = [prev[0] or float(w.get('x1', 0)) <= best_x,
                                   prev[1] or float(w.get('x0', 0)) >= best_x]
                _dual = sum(1 for v in _y_sides.values() if v[0] and v[1])
                if _y_sides and _dual > 0.40 * len(_y_sides):
                    pages_text.append(page.extract_text() or "")
                    continue

                # Crop each column and extract text in reading order
                try:
                    left_text  = (page.crop((0,       0, best_x, ph)).extract_text() or "").strip()
                    right_text = (page.crop((best_x,  0, pw,     ph)).extract_text() or "").strip()
                except Exception:
                    pages_text.append(page.extract_text() or "")
                    continue

                # Post-crop sanity: character-level word splits (caused by a false gutter
                # cutting through the middle of full-width lines) produce fragments like
                # "agement", "ifecycle", "anagement" that start with VOWELS (because the
                # split falls at a consonant-vowel boundary within polysyllabic words).
                # Legitimate right-column continuation lines start with full English words
                # which predominantly start with CONSONANTS ("medical", "marking", etc.).
                # If ≥2 lowercase-starting right-column lines exist and ≥60% of them start
                # with a vowel, the gutter is a false split — fall back to extract_text().
                _right_lines = [l for l in right_text.splitlines() if l.strip()]
                _lower_first_chars = [
                    l.strip()[0] for l in _right_lines[:12]
                    if l.strip() and l.strip()[0].islower()
                ]
                _vowel_frag_starts = sum(1 for c in _lower_first_chars if c in 'aeiou')
                if (len(_lower_first_chars) >= 2
                        and _vowel_frag_starts >= len(_lower_first_chars) * 0.6):
                    logger.debug(
                        "pdfplumber: right column has %d/%d vowel-starting lowercase lines "
                        "— word fragments from false 2-col split, using extract_text()",
                        _vowel_frag_starts, len(_lower_first_chars)
                    )
                    pages_text.append(page.extract_text() or "")
                    continue

                logger.info(f"pdfplumber 2-col split at x={best_x:.1f} "
                            f"(left={left_n} words, right={right_n} words)")

                # If the left column contains only sidebar-metadata headings (personal
                # details, hobbies, languages, extra-curricular) and no valuable resume
                # sections, emit ONLY the right column. This prevents continuation content
                # on the right (WE bullets, PROJECTS) from being broken by left-column
                # skip-headings that reset find_sections() state to None.
                _SIDEBAR_ONLY_RE = re.compile(
                    r'^(?:other personal details|personal information|personal details|'
                    r'hobbies?|interests?|extra\s*curricular|declaration|references|languages?)',
                    re.I
                )
                _VALUABLE_SECTION_RE = re.compile(
                    r'\b(?:skills|experience|education|projects|certifications|summary|profile)\b',
                    re.I
                )
                left_first_line = left_text.split('\n')[0].strip() if left_text else ""
                _left_is_sidebar_only = (
                    bool(_SIDEBAR_ONLY_RE.match(left_first_line))
                    and not _VALUABLE_SECTION_RE.search(left_text)
                )

                if _left_is_sidebar_only:
                    # Drop the left-column sidebar: its section headings (HOBBIES, LANGUAGES,
                    # etc.) would corrupt section state and cause right-column continuation
                    # content to be discarded by find_sections().
                    logger.info(f"Page left column is sidebar-only — emitting right column only")
                    pages_text.append(right_text)
                else:
                    # Detect label+content layout: narrow label column (SUMMARY / EXPERIENCE /
                    # SKILLS headings only) + wide content column.  If we concatenate
                    # left+right all headings appear before all content, breaking find_sections.
                    # Instead, reconstruct in true reading order (row-interleaved) so each
                    # heading immediately precedes its paragraph.
                    _left_lines  = [l.strip() for l in left_text.split('\n')  if l.strip()]
                    _right_lines = [l.strip() for l in right_text.split('\n') if l.strip()]
                    _avg_left_len  = (sum(len(l) for l in _left_lines)  / len(_left_lines))  if _left_lines  else 999
                    _short_left    = (sum(1 for l in _left_lines if len(l.split()) <= 3) / len(_left_lines)) if _left_lines else 0
                    _line_ratio    = len(_right_lines) / max(len(_left_lines), 1)
                    # word-count ratio (right_n/left_n) is more reliable than line ratio
                    # when pdfplumber compresses right-column text into few lines
                    _word_ratio    = right_n / max(left_n, 1)
                    # heading fraction: label+content layouts have ONLY section headings on the
                    # left (SKILLS / EDUCATION / etc.); content-sidebar layouts have actual data
                    # (email, phone, skill items) mixed in.  Require ≥30% of left lines to be
                    # recognised headings so we don't y-pair content-sidebar pages.
                    _left_hdr_cnt = sum(1 for l in _left_lines if canonical_section_name(l) is not None)
                    _left_heading_frac = _left_hdr_cnt / max(len(_left_lines), 1)
                    _is_label_layout = (
                        _avg_left_len < 25 and _short_left > 0.70
                        and (_line_ratio > 3 or _word_ratio > 5)
                        and _left_heading_frac >= 0.30
                    )
                    logger.debug(
                        f"label+content check: avg_left={_avg_left_len:.1f} "
                        f"short_left={_short_left:.2f} line_ratio={_line_ratio:.1f} "
                        f"word_ratio={_word_ratio:.1f} hdr_frac={_left_heading_frac:.2f} "
                        f"→ {_is_label_layout}"
                    )

                    if _is_label_layout:
                        # Label+content layout: pair each right-column content line with
                        # the heading in the left column that sits closest above it by y.
                        # Uses per-crop word coordinates (more reliable than full-page
                        # words which can have PDF-stream ordering artefacts).
                        logger.info("Label+content layout — crop-based y-pairing")
                        try:
                            _lw = page.crop((0,      0, best_x, ph)).extract_words() or []
                            _rw = page.crop((best_x, 0, pw,     ph)).extract_words() or []
                            if not _lw or not _rw:
                                raise ValueError("empty crop words")

                            # Build left-column lines: wy → text
                            _LYT = 4.0
                            _lrows = {}
                            for w in _lw:
                                wy = round(float(w.get('top', 0)) / _LYT)
                                _lrows.setdefault(wy, []).append(
                                    (float(w.get('x0', 0)), w.get('text', '')))
                            _l_lines = sorted(
                                [(wy * _LYT,
                                  ' '.join(t[1] for t in sorted(v, key=lambda t: t[0])))
                                 for wy, v in _lrows.items()],
                                key=lambda x: x[0])
                            _l_lines = [(y, t) for y, t in _l_lines if t.strip()]

                            # Build right-column lines: wy → text
                            _RYT = 3.0
                            _rrows = {}
                            for w in _rw:
                                wy = round(float(w.get('top', 0)) / _RYT)
                                _rrows.setdefault(wy, []).append(
                                    (float(w.get('x0', 0)), w.get('text', '')))
                            _r_lines = sorted(
                                [(wy * _RYT,
                                  ' '.join(t[1] for t in sorted(v, key=lambda t: t[0])))
                                 for wy, v in _rrows.items()],
                                key=lambda x: x[0])
                            _r_lines = [(y, t) for y, t in _r_lines if t.strip()]

                            # Pair each right line to the nearest left heading above it.
                            # Tolerance: heading may start up to 15pt above first content line.
                            _PTOL   = 15.0
                            _l_ys   = [y for y, _ in _l_lines]
                            _segs   = {y: [] for y, _ in _l_lines}
                            _pre    = []  # content before first heading (e.g. experience continuation)

                            for r_y, r_txt in _r_lines:
                                _bhy = None
                                for l_y in reversed(_l_ys):
                                    if l_y <= r_y + _PTOL:
                                        _bhy = l_y
                                        break
                                (_segs[_bhy] if _bhy is not None else _pre).append(r_txt)

                            _out = _pre[:]
                            for l_y, l_txt in _l_lines:
                                _out.append(l_txt)
                                _out.extend(_segs[l_y])

                            pages_text.append(
                                '\n'.join(l for l in _out if l.strip()))

                        except Exception as _exc:
                            logger.debug(f"label+content crop-pair failed: {_exc}")
                            pages_text.append(left_text + "\n" + right_text)
                    else:
                        # Content-sidebar layout: left column has sidebar data (personal info +
                        # skills), right column has the main content (name, title, summary…).
                        # Prepend the right column's header lines (name, title — everything
                        # before the first section heading) so that parse_resume_text sees
                        # the candidate's name and title at the very top of the document,
                        # even though the sidebar content is much longer.
                        # Only prepend the right-column header when the right column
                        # genuinely starts with a person's name (i.e. this is a true
                        # name-on-right sidebar layout, not a table split).
                        _rh_lines = []
                        _first_right = next(
                            (_rl.strip() for _rl in right_text.splitlines() if _rl.strip()), ""
                        )
                        if _looks_like_name(_first_right):
                            for _rl in right_text.splitlines():
                                _rl = _rl.strip()
                                if not _rl:
                                    continue
                                if canonical_section_name(_rl):  # first section heading → stop
                                    break
                                _rh_lines.append(_rl)
                                if len(_rh_lines) >= 3:  # name, title, optional tagline
                                    break
                        if _rh_lines:
                            pages_text.append(
                                "\n".join(_rh_lines) + "\n" + left_text + "\n" + right_text
                            )
                        else:
                            pages_text.append(left_text + "\n" + right_text)

        combined = "\n".join(pages_text).strip()
        return combined if combined else None
    except Exception as e:
        logger.debug(f"_pdfplumber_two_col_text failed: {e}")
        return None


def _pymupdf_page_text(page):
    """Extract readable text from a page, handling 2-column layouts correctly.

    If the page has a clear 2-column structure (detected via block positions),
    the left and right columns are processed independently and concatenated —
    preventing content from the two columns from being mixed on the same line.
    Falls back to single-stream extraction for 1-column pages.
    """
    raw_words = page.get_text("words")
    if not raw_words:
        return ""

    # Primary: conservative block-based detection. Fallback: word-gutter detection
    # for interleaved layouts where wide left-column lines span the page centre.
    split_x = _detect_two_col_split(page) or _detect_col_gutter_words(page)
    if split_x:
        left_words  = [w for w in raw_words if w[0] <  split_x]
        right_words = [w for w in raw_words if w[0] >= split_x]
        return (_words_to_text(left_words) + "\n" + _words_to_text(right_words)).strip()

    return _words_to_text(raw_words)


def _fix_wrapped_email(text):
    """Join emails that the PDF wrapped mid-TLD or mid-local-part across two lines.

    Case 1 (post-@ split):  "sindhusundaramoorthy30@gmail.c"  +  "om"
    Case 2 (pre-@ split):   "manibharathi601rav"  +  "i@gmail.com"
    """
    lines = text.split("\n")
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if i + 1 < len(lines) and "@" in line:
            next_line = lines[i + 1].strip()
            # Case 1: current line ends with partial TLD (≤2 chars), next line is TLD suffix
            if re.match(r"^[a-zA-Z]{1,4}$", next_line):
                if re.search(r"@[A-Za-z0-9.\-]+\.[A-Za-z]{1,2}$", line):
                    line = line.rstrip() + next_line
                    i += 1
        elif i + 1 < len(lines) and "@" not in line:
            next_line = lines[i + 1].strip()
            # Case 2: next line starts with a short prefix (≤5 chars) then @domain.tld
            # and current line ends with 5+ alphanum chars (looks like partial local part)
            m = re.match(r"^([A-Za-z0-9]{1,5})(@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})", next_line)
            if m and re.search(r"[A-Za-z0-9]{5,}$", line):
                line = line.rstrip() + next_line
                i += 1
        result.append(line)
        i += 1
    return "\n".join(result)


def extract_text_from_pdf(path):
    """Extract text from PDF.

    Order: pymupdf (fastest) → pdfplumber → PyPDF2.
    If email or phone are still missing after text extraction, a targeted OCR
    pass runs on the top 25 % of page 1 only (icon-header contact bars).
    """
    try:
        logger.info(f"Extracting text from PDF: {path}")
        text = ""

        # ── Primary: pymupdf ──────────────────────────────────────────────────
        if HAS_PYMUPDF:
            try:
                import pymupdf as fitz
                doc = fitz.open(str(path))
                parts = [_pymupdf_page_text(page) for page in doc]
                doc.close()
                text = "\n".join(parts)
                logger.info(f"pymupdf extracted {len(parts)} pages")
            except Exception as e:
                logger.debug(f"pymupdf text extract failed: {e}")

        # ── Fallback 1: pdfplumber ────────────────────────────────────────────
        if not text.strip() and HAS_PDFPLUMBER:
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    parts = [p for page in pdf.pages if (p := page.extract_text())]
                    if parts:
                        text = "\n".join(parts)
                        logger.info(f"pdfplumber extracted {len(pdf.pages)} pages")
            except Exception as e:
                logger.debug(f"pdfplumber failed: {e}")

        # ── Fallback 2: PyPDF2 ────────────────────────────────────────────────
        if not text.strip():
            try:
                reader = PdfReader(str(path))
                if not reader.pages:
                    raise ValueError("PDF has no pages")
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                logger.info(f"PyPDF2 extracted {len(reader.pages)} pages")
            except Exception as e:
                logger.debug(f"PyPDF2 failed: {e}")

        if not text.strip():
            raise ValueError("Could not extract text from PDF using any method")

        # ── Targeted OCR for icon-header contact info ─────────────────────────
        # Only runs when email or phone are absent AND EasyOCR is installed.
        has_email = bool(re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text))
        has_phone_m = re.search(r"[\+]?\(?\d[\d\s\-\.\(\)]{7,20}\d", text)
        has_phone = bool(has_phone_m and len(re.sub(r"\D", "", has_phone_m.group())) >= 10)
        logger.info(f"Text layer — has_email={has_email}, has_phone={has_phone}")
        if HAS_PYMUPDF and (not has_email or not has_phone):
            if not HAS_OCR:
                logger.warning(
                    "EasyOCR is NOT installed — email/phone in icon-font headers cannot be read. "
                    "Fix: stop Flask, run  pip install easyocr  then restart."
                )
            else:
                ocr_text = _ocr_contact_strip(path)
                if ocr_text:
                    # Extract only email/phone from OCR and inject as labelled lines.
                    # Never prepend raw OCR — it duplicates content that is already in
                    # the text layer, which corrupts summary/section parsing.
                    contact_lines = []
                    if not has_email:
                        em = _extract_email(ocr_text)
                        if em:
                            contact_lines.append(f"Email: {em}")
                    if not has_phone:
                        ph = _extract_phone(ocr_text)
                        if ph:
                            contact_lines.append(f"Phone: {ph}")
                    if contact_lines:
                        text = "\n".join(contact_lines) + "\n" + text
                        logger.info(f"OCR contact injected: {contact_lines}")

        text = _fix_wrapped_email(text)
        return text

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}", exc_info=True)
        raise


def _extract_canvas_docx(body):
    """Extract and reconstruct text from DOCX files that use floating text boxes
    (Word Processing Canvas / Group layout).  Standard paragraph iteration yields
    nothing for these files because all content lives in anchored drawing shapes.
    """
    W_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    MC_NS  = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    A_NS   = "http://schemas.openxmlformats.org/drawingml/2006/main"
    FALLBACK_TAG = f"{{{MC_NS}}}Fallback"

    def _in_fallback(el):
        for anc in el.iterancestors():
            if anc.tag == FALLBACK_TAG:
                return True
        return False

    def _box_lines(txbx):
        lines = []
        for p in txbx.findall(f"{{{W_NS}}}p"):
            t = "".join(r.text or "" for r in p.findall(f".//{{{W_NS}}}t"))
            if t.strip():
                lines.append(t.strip())
        return lines

    def _box_pos(txbx):
        """(y, x) position of the wps:wsp shape that contains this text box."""
        parent = txbx.getparent()          # wps:txbx
        if parent is None:
            return (0, 0)
        wsp = parent.getparent()           # wps:wsp
        if wsp is None or not wsp.tag.endswith("}wsp"):
            return (0, 0)
        spPr = wsp.find(f"{{{WPS_NS}}}spPr")
        if spPr is None:
            return (0, 0)
        xfrm = spPr.find(f"{{{A_NS}}}xfrm")
        if xfrm is None:
            return (0, 0)
        off = xfrm.find(f"{{{A_NS}}}off")
        if off is None:
            return (0, 0)
        return (int(off.get("y", 0)), int(off.get("x", 0)))

    # Collect non-VML-fallback text boxes
    boxes = []
    for txbx in body.findall(f".//{{{W_NS}}}txbxContent"):
        if _in_fallback(txbx):
            continue
        lines = _box_lines(txbx)
        text  = "\n".join(lines)
        if not text.strip():
            continue
        boxes.append({"pos": _box_pos(txbx), "text": text, "lines": lines})

    _JOB_WORDS = {
        "engineer", "manager", "lead", "analyst", "developer", "designer",
        "specialist", "consultant", "director", "officer", "executive",
        "programmer", "architect", "scientist", "researcher", "coordinator",
        "technician", "associate", "assistant", "head", "intern", "trainee",
    }

    # ── Classify each box ────────────────────────────────────────────────────
    for box in boxes:
        txt   = box["text"]
        first = box["lines"][0]
        pos   = box["pos"]

        # Icon / decorative single character
        if len(txt.strip()) <= 2:
            box["kind"] = "icon"
            continue

        # Shapes with explicit non-zero (y, x) positions are anchored inside the
        # experience section of the canvas — treat all such content as experience.
        if pos != (0, 0):
            box["kind"] = "experience_content"
            continue

        # Section label (matches SECTION_ALIASES, short text)
        sec = canonical_section_name(first)
        if sec and len(first.split()) <= 4:
            box["kind"]    = "label"
            box["section"] = sec
            continue

        # Contact info (phone / email present)
        if re.search(
            r"\+?\d[\d\s\-]{7,}\d|[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
            txt,
        ):
            box["kind"] = "contact"
            continue

        # Education content (degree / university keywords)
        if re.search(
            r"\b(?:university|college|school|b\.?com|b\.?sc|b\.?tech|m\.?sc|mba"
            r"|bachelor|master|degree)\b",
            txt, re.I,
        ):
            box["kind"] = "education"
            continue

        # Long single paragraph → professional summary
        if len(txt) > 200 and len(box["lines"]) <= 4:
            box["kind"] = "summary"
            continue

        # Multi-item list → skills
        if len(box["lines"]) >= 3:
            box["kind"] = "skills"
            continue

        # Short proper-case block: distinguish name vs. professional title
        words = first.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words[:3] if w.isalpha()):
            if any(w.lower() in _JOB_WORDS for w in words):
                box["kind"] = "title"   # e.g. "Validation Engineer"
            else:
                box["kind"] = "name"    # e.g. "Sindhu Sundaramoorthy"
        else:
            box["kind"] = "misc"

    # ── Reconstruct text in semantic reading order ────────────────────────────
    parts = []

    def _add(iterable):
        for b in iterable:
            parts.append(b["text"])

    # 1. Name first, then title
    _add(b for b in boxes if b.get("kind") == "name")
    _add(b for b in boxes if b.get("kind") == "title")

    # 2. Contact info — expand bare LinkedIn handle to full URL if present
    for box in boxes:
        if box.get("kind") != "contact":
            continue
        contact_lines = []
        for line in box["lines"]:
            line = line.strip()
            if re.match(r"^[a-zA-Z][a-zA-Z0-9\-]{4,}$", line):
                # Bare handle (no @, no digits only, no spaces) → LinkedIn slug
                line = f"https://www.linkedin.com/in/{line}"
            contact_lines.append(line)
        parts.append("\n".join(contact_lines))

    # 3. Professional summary
    lbl = next((b for b in boxes if b.get("kind") == "label" and b.get("section") == "summary"), None)
    if lbl:
        parts.append(lbl["lines"][0])
    _add(b for b in boxes if b.get("kind") == "summary")

    # 4. Experience — all explicitly-positioned shapes sorted by y (page order)
    exp_content = sorted(
        [b for b in boxes if b.get("kind") == "experience_content"],
        key=lambda b: b["pos"],
    )
    if exp_content:
        parts.append("EXPERIENCE")
        for b in exp_content:
            parts.append(b["text"])

    # 5. Education
    edu_lbl = next((b for b in boxes if b.get("kind") == "label" and b.get("section") == "education"), None)
    if edu_lbl:
        parts.append(edu_lbl["lines"][0])
    _add(b for b in boxes if b.get("kind") == "education")

    # 6. Skills
    sk_lbl = next((b for b in boxes if b.get("kind") == "label" and b.get("section") == "skills"), None)
    if sk_lbl:
        parts.append(sk_lbl["lines"][0])
    _add(b for b in boxes if b.get("kind") == "skills")

    return "\n".join(parts)


def extract_text_from_docx(path):
    try:
        logger.info(f"Extracting text from DOCX: {path}")
        doc = Document(str(path))
        parts = []
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph

        # Read document section headers first — some templates put name/contact info there
        try:
            for section in doc.sections:
                for hdr in (section.header, section.first_page_header, section.even_page_header):
                    if hdr is not None:
                        for para in hdr.paragraphs:
                            t = para.text.strip()
                            if t:
                                parts.append(t)
                        for tbl in hdr.tables:
                            for row in tbl.rows:
                                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                                if row_text:
                                    parts.append(row_text)
        except Exception:
            pass  # Headers are optional — skip silently if unavailable

        for child in doc.element.body:
            if child.tag == qn("w:p"):
                para = DocxParagraph(child, doc)
                if para.text.strip():
                    parts.append(para.text.strip())
            elif child.tag == qn("w:tbl"):
                tbl = DocxTable(child, doc)
                for row in tbl.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)

        body_text = "\n".join(parts).strip()

        # If the standard body extraction yielded nothing, the file likely uses
        # a floating text-box canvas layout — try the canvas extractor.
        if not body_text:
            logger.info("No body text found; trying canvas/text-box extraction")
            canvas_text = _extract_canvas_docx(doc.element.body)
            if canvas_text.strip():
                logger.info(f"Canvas extraction yielded {len(canvas_text)} chars")
                return canvas_text
        else:
            # Even when body text exists, supplement with canvas text boxes — some
            # templates have a sidebar (skills, contact) in floating boxes not in body.
            try:
                canvas_text = _extract_canvas_docx(doc.element.body)
                if canvas_text.strip():
                    # Add only lines that are not already present in the body
                    existing = set(l.strip().lower() for l in parts if l.strip())
                    for canvas_line in canvas_text.split('\n'):
                        cl = canvas_line.strip()
                        if cl and cl.lower() not in existing:
                            parts.append(cl)
                            existing.add(cl.lower())
            except Exception:
                pass

        logger.info(f"Successfully extracted {len(parts)} parts from DOCX")
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}", exc_info=True)
        raise


def extract_resume_text(path, ext):
    if ext == "pdf":
        return extract_text_from_pdf(path)
    if ext == "docx":
        return extract_text_from_docx(path)
    raise ValueError("Unsupported file type")


# ── Resume parsing ─────────────────────────────────────────────────────────────

def normalize_lines(text):
    cleaned = text.replace("•", "•")
    return [re.sub(r"\s+", " ", line).strip(" :-\t") for line in cleaned.splitlines() if line.strip()]


def canonical_section_name(line):
    cleaned = re.sub(r"[^a-zA-Z ]", " ", line).lower()
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return None
    for field, aliases in SECTION_ALIASES.items():
        if cleaned in aliases:
            return field
    return None


def _looks_like_company_name(line):
    """
    Detect if a line looks like a company name.
    E.g., "HCL TECH", "TATA Elxsi", "Lyrus Life Sciences Pvt. Ltd.", "Google India"
    """
    line = line.strip()

    # Remove trailing punctuation
    line = line.rstrip('.,;:!?')
    words = line.split()

    # Must be 2-6 words
    if not (2 <= len(words) <= 6):
        return False

    # Check if all words start with uppercase
    if not all(w and w[0].isupper() for w in words):
        return False

    # Exclude common non-company words that appear in job titles or descriptions
    exclude_words = {
        "the", "and", "or", "for", "with", "in", "at", "by", "as",
        "to", "from", "of", "on", "is", "are", "be", "been",
        "experience", "duration", "designation", "project",
        "highly", "skilled", "professional", "years", "month"
    }
    lower_words = {w.lower().rstrip('.,;:') for w in words}
    if lower_words & exclude_words:
        return False

    # Company keywords (including variations)
    company_keywords = {
        "tech", "technologies", "systems", "solutions", "services",
        "group", "corporation", "corp", "inc", "ltd", "llc", "llp",
        "company", "consulting", "consultants", "labs", "designs",
        "pvt", "private", "limited", "healthcare", "pharma", "bio",
        "sciences", "elxsi", "tcs", "infosys", "wipro", "accenture"
    }
    has_company_keyword = any(w.lower().rstrip('.,;:') in company_keywords for w in words)

    # Pattern 1: "XYZ Corp/Ltd/Inc" or "XYZ Group" etc
    if has_company_keyword:
        return True

    # Pattern 2: All words are uppercase (e.g., "HCL TECH", "TATA ELXSI")
    if all(w.isupper() for w in words):
        return True

    # Pattern 3: 2-3 capitalized words, likely a company (e.g., "Google India", "Amazon AWS")
    # But exclude if it looks like a title (contains job keywords) or a soft-skill descriptor
    # (e.g. "Good Team Player", "Optimistic Thinker" are personal-quality phrases, not companies).
    job_keywords = {
        "engineer", "manager", "lead", "analyst", "developer", "designer",
        "specialist", "consultant", "director", "officer", "executive"
    }
    soft_skill_words = {
        "team", "player", "thinker", "learner", "communicator", "listener",
        "optimistic", "motivated", "adaptable", "flexible", "reliable",
        "proactive", "dedicated", "driven", "focused", "creative",
        "innovative", "passionate", "responsible", "organized", "punctual",
        "enthusiastic", "diligent", "hardworking", "versatile", "empathetic",
        "confident", "assertive", "collaborative", "analytical", "strategic",
    }
    has_job_keyword = any(w.lower() in job_keywords for w in words)
    has_soft_skill = any(w.lower() in soft_skill_words for w in words)
    if not has_job_keyword and not has_soft_skill and len(words) <= 4:
        return True

    return False


def find_sections(lines):
    sections = {key: "" for key in SECTION_ALIASES}
    _SKIP_HEADINGS = {
        "personal", "personal information", "personal details", "other personal details",
        "references", "declaration", "interests", "languages",
        "websites portfolios and profiles", "websites and profiles",
        "websites portfolios", "and profiles",
        "additional information", "other information", "extra curricular",
        "soft skills",
        # "achievements" is NOT here — it has its own SECTION_ALIASES key and folds into certifications via _SECTION_FOLD
    }
    _SKIP_STARTSWITH = {
        "websites portfolios and profiles", "websites and profiles",
        "websites portfolios", "and profiles",
    }
    current = None
    bucket = []
    lines = list(lines)  # materialise so we can do lookahead
    _work_meta_re = re.compile(r'^\s*(duration|designation|project)\s*:', re.I)
    for _line_idx, line in enumerate(lines):
        # A single lowercase word ending with a period is a sentence continuation
        # (e.g. "projects." wrapped from "...strategies for GxP projects."), not a heading.
        words_raw = line.split()
        if len(words_raw) == 1 and words_raw[0][-1] == "." and words_raw[0][0].islower():
            if current:
                bucket.append(line)
            continue

        cleaned_line = re.sub(r"[^a-zA-Z ]", " ", line).lower()
        cleaned_line = re.sub(r"\s+", " ", cleaned_line).strip()

        _is_skip = cleaned_line in _SKIP_HEADINGS or any(
            cleaned_line.startswith(sh) for sh in _SKIP_STARTSWITH
        )
        if _is_skip:
            if current and bucket:
                sections[current] = "\n".join(bucket).strip()
            current = None
            bucket = []
            continue

        section = canonical_section_name(line)
        inline_remainder = None

        if not section:
            # An unlabeled company-name line starts the experience section only when the
            # next 1-2 non-empty lines look like work-experience metadata
            # ("Duration :", "Designation :", "Project :"). This prevents the person's
            # own name (e.g. "SIVARANJANI D" in a page-2 repeated header) from being
            # misidentified as a company name when current is None.
            _next_lines = lines[_line_idx + 1 : _line_idx + 3]
            _followed_by_work_meta = any(
                _work_meta_re.match(nl) for nl in _next_lines if nl.strip()
            )
            if _looks_like_company_name(line) and _followed_by_work_meta:
                section = "experience"
                inline_remainder = line  # include the company name as the first experience line
            else:
                words = line.split()
                # Lines starting with a non-letter character (bullet markers ✓, •, –,
                # numbered-list "1.", etc.) are list items, never section headings.
                # Skip the prefix-heading scan to avoid e.g. "✓ Expertise in Defect Tracking"
                # being parsed as a 2-word "✓ Expertise" → skills heading.
                if words and not words[0][0].isalpha():
                    if current:
                        bucket.append(line)
                    continue
                for n in range(min(3, len(words)), 0, -1):
                    prefix_section = canonical_section_name(" ".join(words[:n]))
                    if prefix_section:
                        remainder = " ".join(words[n:]).strip()
                        # Single-word prefix only counts as a heading when:
                        #   • it IS the entire line (no remainder), OR
                        #   • it ends with a separator like "Skills:", OR
                        #   • it is ALL-CAPS (e.g. "EXPERIENCE VALIDATION ENGINEER..." from
                        #     row-interleaved PDF extraction — heading + inline content on same line).
                        # Title-Case words like "Experience on all ALM modules" are content.
                        _word0_bare = words[0].rstrip('.,;:!?')
                        if n == 1 and remainder and not words[0].endswith(':') and not _word0_bare.isupper():
                            break  # "Experience on all ALM modules" → content, not heading
                        section = prefix_section
                        inline_remainder = remainder if remainder else None
                        break

        if section:
            # Guard: a single Title-case word (e.g. "Qualifications") that immediately
            # follows an incomplete line (no terminal punctuation) is a wrapped continuation
            # of that line — NOT a new section heading. E.g. the bullet "…and Software\n
            # Qualifications." wraps across lines but both chunks belong to the same section.
            # ALL-CAPS words (e.g. "EDUCATION") are always genuine headings; skip the guard.
            if (len(words_raw) == 1 and current and bucket
                    and not words_raw[0].rstrip('.,;:!?').isupper()):
                # Look past any trailing empty lines to the last substantive bucket line.
                # An empty line between wrapped content (e.g. PDF page rendering gaps)
                # must not prevent the guard from seeing the prior non-empty line.
                _prev = next(
                    (bl.rstrip() for bl in reversed(bucket) if bl.strip()),
                    ""
                )
                if _prev and _prev[-1] not in '.!?;:':
                    bucket.append(line)
                    continue
            if section == current:
                # "Summary" appearing mid-summary (e.g. wrapped "Validation Summary Report.")
                # — ignore it so the current section doesn't restart and lose its content.
                continue
            if current and bucket:
                sections[current] = "\n".join(bucket).strip()
            current = section
            bucket = []
            if inline_remainder:
                bucket.append(inline_remainder)
            continue

        if current:
            # Drop bare page-number lines (1–3 digit standalone numbers from PDF
            # page breaks, e.g. "1" between two halves of a skills list, "3" at the
            # end of the last page). Legitimate content never appears as a bare digit.
            if not re.match(r'^\s*\d{1,3}\s*$', line):
                bucket.append(line)

    if current and bucket:
        sections[current] = "\n".join(bucket).strip()

    # Fold extended sections into the 6 core display sections, then return only those.
    for src, dst in _SECTION_FOLD.items():
        content = sections.get(src, "")
        if content and dst:
            sep = "\n" if sections.get(dst) else ""
            sections[dst] = (sections.get(dst) or "") + sep + content

    return {k: sections.get(k, "") for k in _CORE_SECTIONS}


def parse_label_value(lines, labels):
    label_pattern = "|".join(re.escape(label) for label in labels)
    pattern = re.compile(rf"(?:^|\b)({label_pattern})\s*[:\-]\s*(.+)$", re.I)
    for line in lines[:30]:
        match = pattern.search(line)
        if match:
            return match.group(2).strip()
    return ""


def _extract_email(text):
    """Return the first valid email found in text, handling common OCR artifacts."""
    # Pre-pass: join lines where an email wraps mid-TLD (e.g. "foo@gmail.c\nom").
    # Require the domain to end with exactly 1 alpha char (incomplete TLD stub) so
    # we don't accidentally join a COMPLETE email (e.g. "foo@gmail.com") with the
    # unrelated word on the next line (e.g. "Mobile").
    text = re.sub(
        r"([A-Za-z0-9._%+\-]+@(?:[A-Za-z0-9\-]+\.)+[A-Za-z]{1})\n([A-Za-z]{2,5})\b",
        r"\1\2", text,
    )

    # Pass 1: standard clean email
    m = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    if m:
        return m.group().strip().rstrip(".,;)")

    # Pass 2: scan lines — rebuild from the @ position handling OCR noise
    for line in text.splitlines():
        if "@" not in line:
            continue
        at_idx = line.find("@")

        # Local part: grab word(s) immediately before @
        before = line[:at_idx]
        local_match = re.search(r"([A-Za-z0-9][A-Za-z0-9.\s_%+\-]{0,50})$", before)
        if not local_match:
            continue
        local = local_match.group(1).strip()
        local = re.sub(r"\s+", ".", local)          # OCR spaces → dots
        local = re.sub(r"\.{2,}", ".", local)        # collapse duplicate dots

        # Domain part: first non-space token after @
        after_raw = line[at_idx + 1:].strip()
        domain = after_raw.split()[0] if after_raw else ""
        domain = domain.strip(".,;:()[]")
        domain = re.sub(r"\s+", "", domain)
        # OCR often drops the dot before TLD: "outlookcom" → "outlook.com"
        domain = re.sub(
            r"(?<=[a-zA-Z])(com|net|org|in|io|co|uk|edu|gov|info|biz)$",
            r".\1", domain, flags=re.I,
        )

        if local and domain and "." in domain and len(domain.split(".")[-1]) >= 2:
            return f"{local}@{domain}"
    return ""


def _extract_phone(text):
    """Return the first phone number with 10–15 digits, fixing common OCR artifacts."""
    for line in text.splitlines():
        for m in re.finditer(r"\(?\+?\(?\d[\d\s\-\.\(\)]{7,20}\d", line):
            candidate = m.group().strip()
            digits = re.sub(r"\D", "", candidate)
            if 10 <= len(digits) <= 15:
                # OCR often reads '+' as '4' or '1' for country-code prefix
                # e.g. "491 8220 133 233" → "+91 8220 133 233"
                if re.match(r"^[41](\d{2}[\s\-])", candidate) and not candidate.startswith("+"):
                    candidate = "+" + candidate[1:]
                return candidate
    return ""


# ── Name / Title candidate helpers ────────────────────────────────────────────

_NAME_STOPWORDS = frozenset({
    "resume", "curriculum", "vitae", "cv", "profile", "candidate",
    "the", "a", "an", "and", "or", "but", "for", "of", "in", "on",
    "at", "to", "from", "with", "by", "as", "who", "that", "which",
    "highly", "skilled", "experienced", "results", "driven", "oriented",
    "seeking", "passionate", "motivated", "proactive", "dedicated",
    "dynamic", "detail", "focused", "professional", "over", "years", "year",
    "strong", "excellent", "extensive", "hands", "proven", "seasoned",
    "total", "experience", "current", "employer", "designation", "qualification",
    "key", "skills", "competencies", "summary", "objectives", "highlights",
    "personal", "information", "education", "project", "work", "certification",
    "achievement", "language", "hobby", "extracurricular", "tools", "applications",
    "core", "section", "header", "details", "accomplishment", "award",
})


def _looks_like_name(line):
    """Return True only if *line* could plausibly be a person's full name."""
    line = line.strip()
    if not line or len(line) > 55:
        return False
    if re.search(r'[,;:!?]|\.{2,}|—|–|\(|\)', line):
        return False
    if line.endswith('.') and ' ' in line:
        return False
    words = line.split()
    if not (1 <= len(words) <= 5):
        return False
    for w in words:
        clean = re.sub(r"[-'.]", "", w)
        if not clean.isalpha():
            return False
        if len(clean) > 3 and not w[0].isupper():
            return False
    lower_words = {w.lower().strip(".,;:-") for w in words}
    if lower_words & _NAME_STOPWORDS:
        return False
    return bool(words) and words[0][0].isupper()


_TITLE_SECTION_WORDS = frozenset({
    "summary", "profile", "objective", "overview", "experience", "education",
    "skills", "certifications", "projects", "employment", "history", "background",
    "qualifications", "competencies", "expertise", "achievements", "declaration",
    "references", "hobbies", "interests", "languages",
})


def _looks_like_title(line):
    """Return True only if *line* could plausibly be a professional job title."""
    line = line.strip()
    if not line or len(line) > 80 or len(line) < 2:
        return False
    words = line.split()
    if len(words) > 8:
        return False
    if re.search(
        r'\bhighly\b|\bseeking\b|\blooking for\b|with \d+\s*years?|over \d+\s*years?|'
        r'\byears? of\b|\bexperienced in\b|\bdetail.oriented\b|\bresults.driven\b|'
        r'\bdedicated to\b|\bresponsible for\b|\bwho has\b',
        line, re.I,
    ):
        return False
    # Reject common resume section headings (e.g. "PROFILE SUMMARY", "WORK EXPERIENCE")
    lower_words = {w.lower().rstrip('.:') for w in words}
    if len(lower_words & _TITLE_SECTION_WORDS) >= 1 and len(words) <= 3:
        return False
    return any(w[0].isupper() for w in words if w and w[0].isalpha())


# ── Contact-header pattern: "Name    Mobile: +91...   E-Mail: ..." ──────────

_CONTACT_LABEL_PAT = re.compile(
    r'\b(?:mobile|phone|tel|email|e[\-\.\ ]?mail|contact)\s*[:\-]',
    re.I,
)


def _extract_name_from_header_line(lines):
    """
    Pull the candidate name from a line that bundles name + contact details, e.g.
    'Santhoshkumar K    Mobile: +91 8608109310    E-Mail: xxx@yyy.com'
    The name precedes the first contact label keyword.
    """
    for line in lines[:8]:
        m = _CONTACT_LABEL_PAT.search(line)
        if m:
            before = line[:m.start()].strip()
            if before and _looks_like_name(before):
                return before
    return ""


def _extract_title_from_para(text):
    """
    Extract a job title embedded in a summary sentence, e.g.:
    'Highly skilled ... Validation and Compliance Lead with 7 years of experience...'
    → 'Validation and Compliance Lead'
    """
    # Pattern 1: <Title Phrase> with N year(s)
    m = re.search(r'\s+with\s+(?:over\s+)?\d+\+?\s*years?\b', text[:700], re.I)
    if m:
        before = text[:m.start()].strip()
        # Grab the last run of capitalized words (with optional "and/&" connectors)
        title_m = re.search(
            r'([A-Z][a-zA-Z]+(?:\s+(?:and|&|[A-Z][a-zA-Z]+)){1,5})$',
            before,
        )
        if title_m:
            candidate = title_m.group(1).strip()
            words = candidate.split()
            if 2 <= len(words) <= 7:
                return candidate
    # Pattern 2: as a/an <Title>
    m2 = re.search(
        r'\bas\s+an?\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,5})',
        text[:700],
    )
    if m2:
        candidate = m2.group(1).strip()
        if 1 <= len(candidate.split()) <= 7:
            return candidate
    return ""


_INTRO_SKIP_RE = re.compile(
    r'@|www\.|linkedin|github|\d{6,}|\+\d{1,3}[\s\-]?\d{3,}|(?<!\w)\d{10,}|'
    r'\b(?:mobile|phone|tel|email|e[\-\.]mail)\s*[:\-]',
    re.I,
)
_LOCATION_LINE_RE = re.compile(
    r'\b(?:india|karnataka|bangalore|bengaluru|chennai|hyderabad|pune|mumbai|delhi|'
    r'noida|gurgaon|gurugram|coimbatore|trivandrum|kochi|jaipur|ahmedabad|kolkata|'
    r'ludhiana|chandigarh|tamil\s*nadu|maharashtra|andhra|telangana|kerala|'
    r'uttar\s*pradesh|rajasthan|gujarat|west\s*bengal)\b',
    re.I,
)


def _extract_intro_paragraph(lines):
    """
    Collect the unheaded introductory block that precedes the first section heading.
    Many resumes start with a professional overview that has no label — this function
    captures it so it can be used as the Professional Summary field.
    Contact-info lines and short standalone location lines are excluded.
    """
    result = []
    for line in lines:
        if canonical_section_name(line):
            break
        stripped = line.strip()
        if not stripped or len(stripped) < 20:
            continue
        if _INTRO_SKIP_RE.search(stripped):
            continue
        # Reject short lines that are purely location data (e.g. "Coimbatore, India")
        if len(stripped) < 45 and _LOCATION_LINE_RE.search(stripped):
            continue
        result.append(stripped)
    return "\n".join(result).strip()


# Role keyword → known group role (checked in order; first match wins) ────────

_ROLE_MAPPING = [
    # CSV (checked before plain "Validation" — more specific)
    ({"csv"},                              {"lead", "senior", "head", "manager", "principal"}, "CSV Lead"),
    ({"csv"},                              set(),                                               "CSV Analyst"),
    ({"computer", "system", "validation"}, {"lead", "senior", "head", "manager"},              "CSV Lead"),
    ({"computer", "system", "validation"}, set(),                                              "CSV Analyst"),
    # CQV
    ({"cqv"},                              {"lead", "senior", "head"},                         "CQV Lead"),
    ({"cqv"},                              set(),                                               "CQV Engineer"),
    ({"commissioning", "qualification"},   {"lead", "senior"},                                 "CQV Lead"),
    ({"commissioning", "qualification"},   set(),                                              "CQV Engineer"),
    # Tosca
    ({"tosca"},                            {"lead", "senior", "head"},                         "Tosca Lead"),
    ({"tosca"},                            set(),                                               "Tosca Engineer"),
    # Automation (after Tosca so "Tosca Automation Lead" → Tosca Lead, not Automation Lead)
    ({"automation"},                       {"lead", "senior", "head", "manager"},               "Automation Lead"),
    ({"automation"},                       set(),                                               "Automation Engineer"),
    # Validation (after CSV/CQV)
    ({"validation"},                       {"lead", "compliance", "senior", "head", "manager"}, "Validation Lead"),
    ({"validation"},                       set(),                                               "Validation Engineer"),
    # Test / QA
    ({"test"},                             {"lead", "senior", "manager", "head"},               "Test Lead"),
    ({"qa"},                               {"lead", "senior", "manager", "head"},               "Test Lead"),
    ({"quality", "assurance"},             {"lead", "senior", "manager"},                       "Test Lead"),
    ({"test"},                             set(),                                               "Test Engineer"),
    ({"qa"},                               set(),                                               "QA Engineer"),
    # IT roles
    ({"full", "stack"},                    set(),                                               "Full Stack Developer"),
    ({"frontend"},                         set(),                                               "Frontend Developer"),
    ({"front", "end"},                     set(),                                               "Frontend Developer"),
    ({"backend"},                          set(),                                               "Backend Developer"),
    ({"back", "end"},                      set(),                                               "Backend Developer"),
    ({"devops"},                           set(),                                               "DevOps Engineer"),
    ({"cloud"},                            {"engineer", "architect", "developer"},              "Cloud Engineer"),
    ({"data"},                             {"engineer"},                                        "Data Engineer"),
    ({"data"},                             {"analyst"},                                         "Data Analyst"),
    ({"business"},                         {"analyst"},                                         "Business Analyst"),
    ({"scrum"},                            set(),                                               "Scrum Master"),
    ({"ui"},                               {"ux"},                                              "UI/UX Designer"),
    ({"ux"},                               set(),                                               "UI/UX Designer"),
    ({"solution"},                         {"architect"},                                       "Solution Architect"),
    ({"cybersecurity"},                    set(),                                               "Cybersecurity Engineer"),
    ({"security"},                         {"engineer", "analyst", "specialist"},              "Cybersecurity Engineer"),
    ({"database"},                         {"administrator", "admin"},                          "Database Administrator"),
    ({"dba"},                              set(),                                               "Database Administrator"),
    ({"artificial", "intelligence"},       set(),                                               "AI/ML Engineer"),
    ({"machine", "learning"},              set(),                                               "AI/ML Engineer"),
    ({"sap"},                              set(),                                               "SAP Consultant"),
    ({"salesforce"},                       set(),                                               "Salesforce Developer"),
    ({"software"},                         set(),                                               "Software Developer"),
    ({"developer"},                        set(),                                               "Software Developer"),
    ({"project"},                          {"manager"},                                         "Project Manager"),
    ({"design", "control"},               set(),                                               "Design Control Consultant"),
    ({"document", "control"},             set(),                                               "Design Control Consultant"),
]


def _map_to_group_role(raw_title):
    """Map a free-form job title to the closest known group role, or return unchanged."""
    if not raw_title:
        return raw_title
    lower_words = set(re.findall(r'\b\w+\b', raw_title.lower()))
    for required, also_has, role in _ROLE_MAPPING:
        if not (required <= lower_words):
            continue
        if also_has and not (also_has & lower_words):
            continue
        return role
    return raw_title


_PDF_CONTACT_RE = re.compile(
    r'@|\d{6,}|\+\d|\b(?:mobile|phone|tel|email|e[\-\. ]?mail|website|www\.|http)\b',
    re.I,
)


def _extract_name_from_pdf_fonts(path):
    """
    Use PyMuPDF font metadata to find the candidate name at the top of page 1.
    Looks for bold+underlined, then bold-only, then largest-font text that passes
    _looks_like_name(). Returns "" when PyMuPDF is unavailable or nothing found.
    """
    if not HAS_PYMUPDF:
        return ""
    try:
        import pymupdf as fitz
        doc = fitz.open(str(path))
        if not doc.page_count:
            doc.close()
            return ""
        page = doc[0]
        top_zone = page.rect.height * 0.30
        dict_data = page.get_text("dict")
        doc.close()
    except Exception:
        return ""

    candidates = []  # (score, font_size, text)
    max_size = 0.0

    for block in dict_data.get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            y0 = line.get("bbox", [0, 0, 0, 9999])[1]
            if y0 > top_zone:
                continue
            name_parts = []
            line_max_size = 0.0
            for span in line.get("spans", []):
                txt = span.get("text", "").strip()
                if not txt or _PDF_CONTACT_RE.search(txt):
                    continue
                flags = span.get("flags", 0)
                size = span.get("size", 0.0)
                is_bold = bool(flags & 16)
                is_underline = bool(flags & 4)
                line_max_size = max(line_max_size, size)
                name_parts.append((txt, size, is_bold, is_underline))
            if not name_parts:
                continue
            max_size = max(max_size, line_max_size)
            line_text = " ".join(p[0] for p in name_parts)
            bold = any(p[2] for p in name_parts)
            underline = any(p[3] for p in name_parts)
            # Try from longest prefix down; take the first that looks like a name
            words = line_text.split()
            for n in range(min(5, len(words)), 0, -1):
                candidate = " ".join(words[:n])
                if _looks_like_name(candidate):
                    score = (bold and underline) * 4 + bold * 2
                    candidates.append((score, line_max_size, candidate))
                    break

    if not candidates:
        return ""

    # Add large-font bonus now that global max_size is known
    scored = [
        (score + (1 if size >= max_size * 0.85 else 0), size, text)
        for score, size, text in candidates
    ]
    scored.sort(key=lambda x: (-x[0], -x[1]))
    return scored[0][2][:80]


def parse_resume_text(text, name_hint=None):
    lines = normalize_lines(text)
    parsed = {
        "full_name": "", "title": "", "email": "", "phone": "", "linkedin": "", "location": "",
        "summary": "", "skills": "", "experience": "", "education": "",
        "certifications": "", "projects": "",
    }
    if not lines:
        return parsed

    # Font-extracted name from PDF metadata takes top priority — skip all other name logic.
    if name_hint and _looks_like_name(name_hint):
        parsed["full_name"] = name_hint[:80]

    # ── Email extraction ──────────────────────────────────────────────────────
    # Strategy: scan every line for a token containing @; require a dot after @.
    parsed["email"] = _extract_email(text)

    # ── Phone extraction ──────────────────────────────────────────────────────
    # Strategy: scan every line; take first candidate whose digit-only count >= 10.
    parsed["phone"] = _extract_phone(text)

    # Join LinkedIn URLs that wrap across lines (e.g. "malthesh-\nkarnam-29a6b416a")
    _lnk_text = re.sub(
        r'(linkedin\.com/in/[a-zA-Z0-9\-_%]*)-[ \t]*\n[ \t]*([a-zA-Z0-9])',
        lambda m: m.group(1) + '-' + m.group(2),
        text, flags=re.I,
    )
    linkedin_match = re.search(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9\-_%]+)", _lnk_text, re.I
    )
    if linkedin_match:
        parsed["linkedin"] = "https://www.linkedin.com/in/" + linkedin_match.group(1)
    else:
        # Bare handle format: "LinkedIn: john-doe" or "LinkedIn ID: john-doe-123abc"
        _bare_match = re.search(
            r'(?:linkedin|linked\s*in)(?:\s*id|\s*profile|\s*handle)?[\s:]+([a-zA-Z0-9][a-zA-Z0-9\-_%]{2,})',
            text, re.I,
        )
        if _bare_match:
            _handle = _bare_match.group(1).strip('-').strip()
            # Must look like a handle (not a common word)
            if len(_handle) >= 3 and not re.match(r'^(profile|handle|id|url|link|page|account)$', _handle, re.I):
                parsed["linkedin"] = "https://www.linkedin.com/in/" + _handle

    # Remove all LinkedIn URLs and fragments from text
    text = re.sub(r'https?://[^\s]*linkedin\.com[^\s]*', '', text, flags=re.I)
    text = re.sub(r'www\.linkedin\.com[^\s]*', '', text, flags=re.I)
    text = re.sub(r'linkedin\.com/in/[a-zA-Z0-9\-_%/@?=&#]*', '', text, flags=re.I)
    text = re.sub(r'\?skipRedirect[^\s]*|\&skipRedirect[^\s]*|#skipRedirect[^\s]*', '', text, flags=re.I)
    # Remove standalone LinkedIn profile ID fragments (e.g., "karnam-29a6b416a/")
    text = re.sub(r'^\s*[a-z]+\-[a-z0-9]+/?$', '', text, flags=re.MULTILINE | re.I)

    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]

    if not parsed["full_name"]:
        parsed["full_name"] = parse_label_value(lines, ["name", "full name", "candidate name"])
    parsed["title"] = parse_label_value(lines, ["title", "designation", "role", "current role"])
    parsed["location"] = parse_label_value(lines, ["location", "address", "city"])

    # Pull name from a header line like "John Smith   Mobile: +91 xxx   E-Mail: yyy"
    # Must run before useful_top_lines filter, which would discard such lines.
    if not parsed["full_name"]:
        parsed["full_name"] = _extract_name_from_header_line(lines)

    _email_fragments = set()
    if parsed.get("email"):
        _em = parsed["email"]
        for _tld_len in range(1, 5):
            _frag = _em[-_tld_len:]
            if _frag.isalpha():
                _email_fragments.add(_frag.lower())

    ignored = ["resume", "curriculum vitae", "cv"]
    useful_top_lines = [
        line for line in lines[:20]
        if line.lower() not in ignored
        and line.lower().strip() not in _email_fragments
        and not (len(line.strip()) <= 4 and line.strip().isalpha() and line.strip().islower())
        and not re.search(
            r"@|www\.|linkedin|github|\d{6,}|\+\d{1,3}[\s\-]?\d{3,}|(?<!\w)\d{10,}"
            r"|date.{0,5}birth|nationality|dob\b|gender|marital|\d{1,2}[./]\d{1,2}[./]\d{2,4}"
            r"|total.{0,10}experience|current employer|current designation|qualification|designation"
            # standalone contact-sidebar labels that must never be mistaken for job titles
            r"|^email$|^mobile$|^phone$|^telephone$|^tel$|^fax$|^social\s+links?$"
            # experience-duration metadata lines like "7 Years 0 Month"
            r"|\d+\s+years?\s+\d+\s+months?",
            line, re.I,
        )
    ]
    # Name fallback: look for a line that actually looks like a person's name.
    # Prevents professional summary sentences from being assigned as the name.
    # Prioritize all-caps names (likely to be actual names in resume headers)
    if not parsed["full_name"]:
        # First pass: look for all-caps names in FIRST 10 lines (e.g., "SIVARANJANI D")
        for _nl in useful_top_lines[:10]:
            if _nl.isupper() and _looks_like_name(_nl):
                parsed["full_name"] = _nl[:80]
                break
        # Second pass: look for TWO-WORD names (common name pattern) in first 15 lines
        if not parsed["full_name"]:
            for _nl in useful_top_lines[:15]:
                _words = _nl.split()
                if len(_words) == 2 and _looks_like_name(_nl):
                    parsed["full_name"] = _nl[:80]
                    break
        # Third pass: look for any proper-cased name in all useful lines
        if not parsed["full_name"]:
            for _nl in useful_top_lines:
                if _looks_like_name(_nl):
                    parsed["full_name"] = _nl[:80]
                    break
        if not parsed["full_name"]:
            # Looser pass: short proper-cased line with 2-5 words and no stopwords
            for _nl in useful_top_lines:
                _nwords = _nl.split()
                if (2 <= len(_nwords) <= 5 and _nl[:1].isupper() and len(_nl) <= 50
                        and not {w.lower().strip(".,;:-") for w in _nwords} & _NAME_STOPWORDS):
                    parsed["full_name"] = _nl[:80]
                    break
    # Title priority 1: extract from first paragraph sentence pattern
    # ("Validation and Compliance Lead with 7 years...") — done FIRST so sentence-embedded
    # titles beat ambiguous short lines from the top-lines scan.
    if not parsed["title"]:
        parsed["title"] = _extract_title_from_para(text[:700])
    # Title priority 2: first short proper-cased line after the name that reads
    # like a job title (not a summary sentence, not a section heading).
    if not parsed["title"]:
        _name_line = parsed.get("full_name", "")
        for _tl in useful_top_lines:
            if _tl == _name_line:
                continue
            if not canonical_section_name(_tl) and _looks_like_title(_tl):
                parsed["title"] = _tl[:80]
                break

    if not parsed["location"]:
        _loc_city_re = re.compile(
            r"\b(india|karnataka|bangalore|bengaluru|chennai|hyderabad|pune|mumbai|delhi|"
            r"noida|gurgaon|gurugram|coimbatore|trivandrum|kochi|jaipur|ahmedabad|kolkata|"
            r"ludhiana|chandigarh)\b", re.I
        )
        for line in lines[:20]:
            if _loc_city_re.search(line):
                # Split on pipe/bullet separators first; take only the segment
                # that contains the city keyword — avoids grabbing "Open to
                # Relocation", "LinkedIn:", years-of-experience text, etc.
                _segs = re.split(r'\s*[|•·]\s*', line)
                _city_seg = next((s for s in _segs if _loc_city_re.search(s)), line)
                # Strip email, phone, and icon characters from the city segment
                loc_line = re.sub(r'[\w.+-]+@[\w.-]+\.[a-z]{2,}', '', _city_seg, flags=re.I)
                loc_line = re.sub(r'[\+\(]?[\d\s\-\(\)]{7,}', '', loc_line)
                loc_line = re.sub(r'[^\x20-\x7E]', ' ', loc_line)
                loc_line = re.sub(r'[,;]+', ',', loc_line)
                loc_line = re.sub(r',\s*,', ',', loc_line)
                loc_line = ', '.join(p.strip() for p in loc_line.split(',') if p.strip())
                if loc_line and len(loc_line) > 2:
                    parsed["location"] = loc_line[:180]
                break

    parsed.update(find_sections(lines))

    # ── Post-process: Extract skills from 2-column layouts ──────────────────────
    # In some PDFs (e.g., Sivaranjani's), skills are interleaved with summary due to
    # column extraction. If skills section is missing, extract from summary AND certifications.
    if not parsed.get("skills"):
        skills_items = []

        # Collect lines from multiple sections that might contain skills
        sections_to_scan = []
        for _sec_key in ("summary", "certifications", "training", "achievements", "career_highlights"):
            if parsed.get(_sec_key):
                sections_to_scan.extend(parsed[_sec_key].split('\n'))

        for line in sections_to_scan:
            line_stripped = line.strip()

            if not line_stripped or len(line_stripped) < 5:
                continue

            # Must start with capital letter
            if not line_stripped[0].isupper():
                continue

            # Skip full sentences that end with periods
            if line_stripped.endswith('.'):
                continue

            # Extract skill, handling cases where skill + summary text appear on same line
            # Pattern: skill text ends where summary paragraph begins (indicated by "Experienced", "professional", etc.)
            skill_part = line_stripped
            if any(marker in line_stripped for marker in [' Experienced ', ' professional ', ' compliance ', ' and compliance']):
                # Find the skill part before summary text
                for marker in [' Experienced ', ' professional ']:
                    if marker in line_stripped:
                        skill_part = line_stripped.split(marker)[0].strip()
                        break

            # Pattern 1: Ends with & or ( (incomplete lines are skills continuing to next line)
            if skill_part.endswith('&') or skill_part.endswith('('):
                skills_items.append(skill_part)
                continue

            # Pattern 2: Contains skill indicators (&, /, parens, colons) and is reasonably short
            has_skill_indicator = any(c in skill_part for c in ['&', '/', '(', ')', ':'])
            if has_skill_indicator and len(skill_part) <= 90:
                # Accept if it has the skill pattern
                word_count = len(skill_part.split())
                if word_count <= 8:  # Allow slightly longer for multi-part skills
                    skills_items.append(skill_part)
                    continue

            # Pattern 3: Short items (1-2 words, 5-40 chars) - likely single skill names
            word_count = len(skill_part.split())
            if 1 <= word_count <= 2 and 5 <= len(skill_part) <= 40:
                skills_items.append(skill_part)
                continue

            # Pattern 4: Multi-word items with acronyms/standards (3-6 words with & or /)
            # Examples: "Medical Device Lifecycle Management & UDI", "ASTM F1980 / ISO 20417"
            if 3 <= word_count <= 6 and (has_skill_indicator or any(c.isdigit() for c in skill_part)):
                # Likely a skill with special formatting
                if len(skill_part) <= 90:
                    skills_items.append(skill_part)
                    continue

            # Pattern 5: Lines with acronyms/numbers that look like skills even without & or /
            # Examples: "MDD & EU MDR 745/2017", "ASTM F1980 / ISO 20417"
            # Accept short-medium lines (up to 5 words) that have numbers and caps
            if 2 <= word_count <= 5 and len(skill_part) <= 90:
                has_number = any(c.isdigit() for c in skill_part)
                has_caps = any(w.isupper() for w in skill_part.split())
                if has_number and has_caps and not skill_part.endswith(' that') and not skill_part.endswith(' and'):
                    skills_items.append(skill_part)

        if skills_items:
            parsed["skills"] = "\n".join(skills_items)
            # Remove extracted skills from summary to clean it up
            if parsed.get("summary"):
                summary_lines = parsed["summary"].split('\n')
                summary_cleaned = '\n'.join(l for l in summary_lines if l.strip() not in skills_items)
                parsed["summary"] = summary_cleaned.strip()

    # Capture the unheaded intro paragraph that precedes the first section heading.
    # Use it only when no explicitly labelled summary section was found; explicit
    # section headings like "PROFILE SUMMARY" take priority.
    _intro = _extract_intro_paragraph(lines)
    if _intro and not parsed.get("summary"):
        parsed["summary"] = _intro

    # Fallback: if summary is still empty, do a direct line-by-line scan for a
    # "Profile Summary" or "Professional Summary" heading and collect what follows.
    # This covers PDFs where the heading's special formatting (decorative chars,
    # 2-column extraction order, etc.) prevented the section parser from capturing it.
    if not parsed.get("summary"):
        _in_ps, _ps_lines = False, []
        for _ln in lines:
            _key = re.sub(r"[^a-zA-Z ]", " ", _ln).lower()
            _key = re.sub(r"\s+", " ", _key).strip()
            if _key in ("profile summary", "professional summary"):
                _in_ps = True
                continue
            if _in_ps:
                _sec = canonical_section_name(_ln)
                if _sec and _sec != "summary":
                    break
                if _ln.strip():
                    _ps_lines.append(_ln.strip())
        if _ps_lines:
            parsed["summary"] = "\n".join(_ps_lines)

    if parsed.get("skills"):
        _url_line_pat = re.compile(
            r"^(https?://|www\.|linkedin\.com|github\.com|gitlab\.com|portfolio|website|blog|"
            r"twitter|facebook|instagram|behance|dribbble|stackoverflow|medium\.com|kaggle)\S*$",
            re.I,
        )
        _portfolio_label_pat = re.compile(
            r"^(portfolio|website|blog|github|gitlab|twitter|facebook|instagram|behance|"
            r"dribbble|stackoverflow|medium|kaggle|profile|link|url)\s*[:\-]",
            re.I,
        )
        _non_skill_heading_pat = re.compile(
            r"^(websites?|portfolios?|websites?[,\s]+portfolios?|websites?.*profiles?|"
            r"and\s+profiles?|profiles?\s+and|"
            r"personal|references?|declaration|additional information|"
            r"other information|education|experience|certifications?|projects?)\s*$",
            re.I,
        )
        _inline_url_pat = re.compile(r"https?://\S+|www\.\S+", re.I)
        clean_skill_lines = []
        for sl in parsed["skills"].splitlines():
            stripped_sl = sl.strip()
            if _url_line_pat.match(stripped_sl):
                continue
            if _portfolio_label_pat.match(stripped_sl):
                continue
            bare = stripped_sl.strip("•-| ")
            if _url_line_pat.match(bare) or _inline_url_pat.fullmatch(bare):
                continue
            cleaned_sl = _inline_url_pat.sub("", sl).strip(" •-,|")
            if not cleaned_sl:
                continue
            if _non_skill_heading_pat.match(re.sub(r"[^a-zA-Z\s]", " ", cleaned_sl).strip()):
                continue
            clean_skill_lines.append(cleaned_sl)
        parsed["skills"] = "\n".join(clean_skill_lines).strip()

    if parsed.get("education"):
        _personal_pat = re.compile(
            r"\b(father|mother|parent|guardian|date.{0,5}birth|born|dob\b|gender|sex\b|"
            r"marital|nationality|citizen|passport|religion|caste|languages known|"
            r"permanent address|current address|pin\s*code|aadhar|pan\b|"
            r"place of birth|age\b|blood group|reference)\b",
            re.I,
        )
        edu_lines = [l for l in parsed["education"].splitlines() if not _personal_pat.search(l)]
        parsed["education"] = "\n".join(edu_lines).strip()

    _TITLE_REJECTS = {
        "summary", "professional summary", "profile", "career profile",
        "professional profile", "objective", "career objective", "overview",
        "about", "about me", "executive summary",
    }
    if not parsed["title"] and parsed["summary"]:
        first_line = parsed["summary"].strip().split("\n")[0].strip()
        title_match = re.match(r"^((?:[A-Z][a-zA-Z]*(?:\s+|$)){1,5})", first_line)
        if title_match:
            candidate = title_match.group(1).strip()
            if 1 <= len(candidate.split()) <= 5 and candidate.lower() not in _TITLE_REJECTS:
                parsed["title"] = candidate

    if not parsed["title"] and parsed["experience"]:
        for exp_line in parsed["experience"].splitlines()[:5]:
            exp_line = exp_line.strip()
            if exp_line and "|" not in exp_line and not re.search(r"\d", exp_line):
                title_match = re.match(r"^((?:[A-Z][a-zA-Z]*(?:\s+|$)){1,6})$", exp_line)
                if title_match:
                    candidate = title_match.group(1).strip()
                    if 1 <= len(candidate.split()) <= 6:
                        parsed["title"] = candidate
                        break

    _TITLE_REJECTS_SET = {
        "summary", "professional summary", "profile", "career profile",
        "professional profile", "objective", "career objective", "overview",
        "about", "about me", "executive summary", "skills", "experience",
        "education", "certifications", "projects",
        # Single descriptor words that are NOT job titles
        "highly", "skilled", "experienced", "dedicated", "motivated",
        "dynamic", "proactive", "seasoned", "accomplished", "passionate",
    }
    if parsed["title"]:
        title_clean = parsed["title"].strip()
        title_clean = re.split(r"[,|/\\–—]", title_clean)[0].strip()
        title_clean = re.sub(r"[\(\[].*$", "", title_clean).strip()
        title_clean = re.sub(
            r"^(summary|professional\s+summary|profile|career\s+profile|"
            r"professional\s+profile|objective|career\s+objective|overview|"
            r"about\s+me?|executive\s+summary)\s+",
            "", title_clean, flags=re.I,
        ).strip()
        title_clean = re.sub(
            r"^(experienced|skilled|dedicated|results.driven|dynamic|seasoned|"
            r"highly experienced|passionate|motivated|proactive|hands.on)\s+",
            "", title_clean, flags=re.I,
        ).strip()
        if title_clean.lower() in _TITLE_REJECTS_SET:
            title_clean = ""
        if title_clean:
            words = title_clean.split()
            if len(words) > 5:
                m = re.match(r"^((?:[A-Z][a-zA-Z]*(?:\s+|$)){1,5})", title_clean)
                title_clean = m.group(1).strip() if m else " ".join(words[:5])
        parsed["title"] = title_clean[:80]

    # Map the extracted/cleaned title to the closest known group role
    if parsed["title"]:
        parsed["title"] = _map_to_group_role(parsed["title"])

    return parsed


def merge_resume_data(form_data, parsed_data, overwrite=False):
    merged = dict(form_data)
    for key, value in parsed_data.items():
        if value and (overwrite or not merged.get(key)):
            merged[key] = value.strip()
    return merged


# ── Ollama "Resume Intelligence" parser (text LLM) ─────────────────────────────

_OLLAMA_BASE         = os.environ.get("OLLAMA_BASE_URL", "http://127.0.0.1:11434")
_TEXT_MODEL          = os.environ.get("OLLAMA_TEXT_MODEL", "resume-parser")
_OLLAMA_TEXT_TIMEOUT = int(os.environ.get("OLLAMA_TEXT_TIMEOUT", "10"))   # per-call cap
_OLLAMA_BUDGET_SECS  = int(os.environ.get("OLLAMA_BUDGET_SECS", "10"))    # total AI budget


# ── Pure-Ollama extraction — NO regex field engine in this path ────────────────
#
# A 1B model cannot one-shot a whole resume (it drifts, mislabels, or returns empty
# above ~1 000 chars of input).  The reliable pattern, confirmed by testing, is:
#   • feed it small, homogeneous pieces, and
#   • use flat JSON for short identity fields, plain text for long sections.
# So we slice the (column-split-cleaned) text into section chunks by heading, then
# make one small Ollama call per piece.  Ollama does ALL the field extraction; the
# only non-LLM step is cutting the text at heading lines.

_BLANK_FIELDS = (
    "full_name", "title", "email", "phone", "linkedin", "location",
    "summary", "skills", "experience", "education", "certifications", "projects",
)

# Heading text → which field its content belongs to. Extends the canonical section
# map with a few all-caps headings small models' resumes use.
_HEADING_FIELD = {
    "technology": "skills", "technologies": "skills", "technical skills": "skills",
    "publications": "certifications", "awards": "certifications",
    "achievements": "certifications", "certificates": "certifications",
}


def _ollama_chat(prompt, *, as_json, num_predict, num_ctx=8192):
    """Single Ollama call. Returns the raw assistant string (or {} dict if as_json)."""
    body = {
        "model": _TEXT_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "stream": False,
        "options": {"temperature": 0, "num_predict": num_predict, "num_ctx": num_ctx},
    }
    if as_json:
        body["format"] = "json"
    url = f"{_OLLAMA_BASE}/api/chat"
    logger.info(f"Ollama call → {url} model={_TEXT_MODEL}")
    req = urllib.request.Request(
        url,
        data=json.dumps(body).encode(),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=_OLLAMA_TEXT_TIMEOUT) as r:
        content = json.loads(r.read()).get("message", {}).get("content", "").strip()
    if not as_json:
        return content
    try:
        return json.loads(content)
    except (ValueError, TypeError):
        return {}


def _split_resume_chunks(text):
    """Cut clean resume text into {field: section_text} pieces at heading lines.

    Handles 2-column PDFs by scanning for all section headings first, then extracting
    content between them by position rather than linear order. This prevents columns
    from getting scrambled (e.g. "KEY SKILLS" appearing after "WORK EXPERIENCE").
    """
    lines = text.split("\n")

    # First pass: identify all section headings and their line numbers
    headings = []  # [(line_num, field_name), ...]
    for i, line in enumerate(lines):
        s = line.strip()
        if not s:
            continue
        key = s.lower().strip(":·•- ")
        field = canonical_section_name(s) or _HEADING_FIELD.get(key)
        # Treat as a heading only if it's short (real headings are 1-3 words).
        # Count only words that contain at least one letter so decorative tokens
        # like "────", "===", "•" etc. do not inflate the word count and cause
        # headings such as "── PROFILE SUMMARY ──" to be silently skipped.
        _alpha_word_count = sum(1 for w in s.split() if any(c.isalpha() for c in w))
        if field and _alpha_word_count <= 3:
            headings.append((i, field))

    # Second pass: extract content between heading positions
    chunks = {}
    for idx, (line_num, field) in enumerate(headings):
        # Start from the line after the heading
        start = line_num + 1
        # End at the next heading (or end of text)
        end = headings[idx + 1][0] if idx + 1 < len(headings) else len(lines)

        # Collect non-empty lines in this range
        content_lines = []
        for line in lines[start:end]:
            s = line.strip()
            if s:
                content_lines.append(s)

        if content_lines:
            content = "\n".join(content_lines).strip()
            # Handle duplicate section headings by appending content
            if field in chunks:
                chunks[field] = chunks[field] + "\n" + content
            else:
                chunks[field] = content

    return chunks


def _pymupdf_rawdict_scan(path):
    """Extract email/phone from page 1 using pymupdf's character-level rawdict mode.

    `get_text("rawdict")` traverses the PDF content stream character by character,
    including text in coloured/styled areas that `get_text("words")` sometimes misses
    when glyph-to-Unicode mappings are non-standard (e.g. icon fonts in contact bars).

    Returns a dict with 'email' and/or 'phone' keys — only the found ones.
    """
    if not HAS_PYMUPDF:
        return {}
    try:
        import pymupdf as fitz
        doc = fitz.open(str(path))
        page = doc[0]
        raw = page.get_text("rawdict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        doc.close()

        line_texts = []
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_str = ""
                for span in line.get("spans", []):
                    for char_obj in span.get("chars", []):
                        ch = char_obj.get("c", "")
                        if ch:
                            line_str += ch
                if line_str.strip():
                    line_texts.append(line_str)

        full_text = "\n".join(line_texts)
        logger.debug(f"rawdict scan page-1 first 400 chars: {full_text[:400]!r}")

        found = {}
        em = _extract_email(full_text)
        if em:
            found["email"] = em
        ph = _extract_phone(full_text)
        if ph:
            found["phone"] = ph
        if found:
            logger.info(f"rawdict scan found: {found}")
        return found
    except Exception as exc:
        logger.debug(f"_pymupdf_rawdict_scan: {exc}")
        return {}


def _tesseract_contact_scan(path):
    """OCR the top portion of page 1 with pytesseract to recover email/phone.

    Designed for styled PDFs where the contact bar (dark background, white text,
    icon fonts) is NOT accessible through any PDF text-layer extraction method.
    Renders page 1 at 200 DPI, crops to the top 35 %, and runs tesseract OCR.

    Returns a dict with 'email' and/or 'phone' keys — only the found ones.
    """
    try:
        import pytesseract
        from PIL import Image
        import pymupdf as fitz
        import numpy as np

        for _tp in [
            r"C:\Users\sindhu.sundara\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.isfile(_tp):
                pytesseract.pytesseract.tesseract_cmd = _tp
                break

        doc = fitz.open(str(path))
        pg = doc[0]
        clip = fitz.Rect(0, 0, pg.rect.width, pg.rect.height * 0.35)
        mat  = fitz.Matrix(200 / 72, 200 / 72)   # 200 DPI — good quality for OCR
        pix  = pg.get_pixmap(matrix=mat, clip=clip, colorspace=fitz.csRGB)
        doc.close()

        img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, 3
        )
        img = Image.fromarray(img_array)

        # --psm 6: assume a uniform block of text (contact bar)
        # --oem 3: LSTM engine
        ocr_text = pytesseract.image_to_string(
            img, config="--psm 6 --oem 3 -c tessedit_char_whitelist="
            "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789@.+\\-_ "
        )
        logger.info(f"Tesseract OCR top-35%: {ocr_text[:300]!r}")

        found = {}
        em = _extract_email(ocr_text)
        if em:
            found["email"] = em
        ph = _extract_phone(ocr_text)
        if ph:
            found["phone"] = ph
        if found:
            logger.info(f"Tesseract contact scan found: {found}")
        return found
    except Exception as exc:
        logger.debug(f"_tesseract_contact_scan: {exc}")
        return {}


def _tesseract_inverted_ocr_text(path):
    """Render page 1, invert the image, and return the raw Tesseract OCR string.

    WHY INVERSION:
    Dark-background contact bars (navy/black with white text) give Tesseract very
    poor results because it expects dark ink on a light page.  Inverting the image
    first (dark background → light, white text → dark) gives full-contrast black
    text on white, which Tesseract reads with high accuracy.

    The RED channel is used for the inversion pass: a navy bar has near-zero red
    values, so inverting the red channel produces a near-white background with
    dark text — maximum contrast for the OCR engine.

    Returns the raw OCR string (may be empty on error).
    """
    try:
        import pytesseract
        from PIL import Image, ImageOps
        import pymupdf as fitz
        import numpy as np

        for _tp in [
            r"C:\Users\sindhu.sundara\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        ]:
            if os.path.isfile(_tp):
                pytesseract.pytesseract.tesseract_cmd = _tp
                break

        doc = fitz.open(str(path))
        pg  = doc[0]
        # Render top 40 % at 300 DPI — higher res improves small-font accuracy
        clip = fitz.Rect(0, 0, pg.rect.width, pg.rect.height * 0.40)
        mat  = fitz.Matrix(300 / 72, 300 / 72)
        pix  = pg.get_pixmap(matrix=mat, clip=clip, colorspace=fitz.csRGB)
        doc.close()

        img_arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, 3
        )
        img = Image.fromarray(img_arr)

        # Pass 1 — invert full colour image → dark-on-light
        img_inv = ImageOps.invert(img)

        # Pass 2 — use red channel only: navy bar → near-black → inverted = near-white
        r_ch = img.split()[0]               # red channel
        r_inv = ImageOps.invert(r_ch)       # invert: 0→255 (white bg), 255→0 (dark text)

        ocr1 = pytesseract.image_to_string(img_inv.convert("L"), config="--psm 6 --oem 3")
        ocr2 = pytesseract.image_to_string(r_inv,                 config="--psm 6 --oem 3")
        combined = ocr1 + "\n" + ocr2
        logger.info(f"Inverted OCR (pass1+pass2 first 400): {combined[:400]!r}")
        return combined
    except Exception as exc:
        logger.debug(f"_tesseract_inverted_ocr_text: {exc}")
        return ""


def _ollama_vision_contact_scan(path):
    """Send the rendered contact bar image to an Ollama vision model to read email/phone.

    The dark-background contact bars in styled resumes are NOT in the PDF text layer
    and cannot be read by any text-extraction method.  This function renders the top
    30 % of page 1 as a PNG, encodes it as base64, and sends it to the first vision-
    capable model found in Ollama (moondream, llava, bakllava, minicpm-v, etc.).

    To enable: run  `ollama pull moondream`  in a terminal (≈1.7 GB one-time download).

    Returns a dict with 'email' and/or 'phone' keys — only the validated ones.
    """
    try:
        import base64
        import io
        import json as _json
        import requests as _req
        import pymupdf as fitz
        import numpy as np
        from PIL import Image

        # Find a vision-capable model in Ollama
        _tags = _req.get(f"{_OLLAMA_BASE}/api/tags", timeout=5).json()
        _vision_kw = ("moondream", "llava", "bakllava", "minicpm-v", "vision",
                      "qwen2-vl", "gemma3", "pixtral")
        _vision_models = [
            m["name"] for m in _tags.get("models", [])
            if any(kw in m["name"].lower() for kw in _vision_kw)
        ]
        if not _vision_models:
            logger.info(
                "No vision model found in Ollama — skipping vision contact scan. "
                "Run: ollama pull moondream"
            )
            return {}

        vision_model = _vision_models[0]
        logger.info(f"Vision model selected: {vision_model}")

        # Render top 30 % of page 1 at 200 DPI
        doc  = fitz.open(str(path))
        pg   = doc[0]
        clip = fitz.Rect(0, 0, pg.rect.width, pg.rect.height * 0.30)
        mat  = fitz.Matrix(200 / 72, 200 / 72)
        pix  = pg.get_pixmap(matrix=mat, clip=clip, colorspace=fitz.csRGB)
        doc.close()

        img_arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, 3
        )
        img = Image.fromarray(img_arr)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        prompt = (
            "This is the top of a resume. There is a dark-colored contact bar with "
            "an email address and a phone number written in white text.\n"
            "Extract ONLY the email address and phone number.\n"
            "Return ONLY valid JSON with no extra text:\n"
            '{"email": "someone@domain.com", "phone": "+91 XXXX XXX XXX"}'
        )

        payload = {
            "model": vision_model,
            "messages": [{"role": "user", "content": prompt, "images": [img_b64]}],
            "stream": False,
            "options": {"temperature": 0, "num_predict": 80},
        }
        resp = _req.post(
            f"{_OLLAMA_BASE}/api/chat", json=payload,
            timeout=_OLLAMA_TEXT_TIMEOUT
        )
        resp.raise_for_status()
        response_text = resp.json().get("message", {}).get("content", "")
        logger.info(f"Vision model raw response: {response_text[:300]!r}")

        # Extract JSON from the response
        json_m = re.search(r'\{[^}]+\}', response_text, re.DOTALL)
        if not json_m:
            return {}

        data = _json.loads(json_m.group())
        found = {}
        _em = str(data.get("email", "")).strip().rstrip(".,;)")
        _ph = str(data.get("phone", "")).strip()

        _em_re = re.compile(r'^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$')
        if _em and _em_re.match(_em):
            found["email"] = _em

        _ph_digits = re.sub(r"\D", "", _ph)
        if 10 <= len(_ph_digits) <= 15:
            found["phone"] = _ph

        if found:
            logger.info(f"Vision model contact found: {found}")
        return found

    except Exception as exc:
        logger.debug(f"_ollama_vision_contact_scan: {exc}")
        return {}


def _pdfplumber_whitetext_scan(path):
    """Extract email/phone by filtering pdfplumber chars to white/light-colored text.

    Dark-background contact bars (navy, black) render text in white.  pdfplumber's
    `page.chars` exposes the `non_stroking_color` (fill color) for every character
    in the PDF content stream.  Filtering to light colors isolates the contact bar
    text from the rest of the page content, then email/phone patterns are applied.

    Returns a dict with 'email' and/or 'phone' keys — only the found ones.
    """
    if not HAS_PDFPLUMBER:
        return {}
    try:
        import pdfplumber
        from collections import defaultdict
        with pdfplumber.open(str(path)) as pdf:
            if not pdf.pages:
                return {}
            chars = pdf.pages[0].chars

        if not chars:
            return {}

        def _is_light_color(color):
            """Return True for white / near-white fill colors."""
            if color is None:
                return False
            if isinstance(color, (int, float)):
                return float(color) >= 0.70
            if isinstance(color, (list, tuple)):
                vals = [float(v) for v in color]
                # RGB or CMYK (inverted for K channel)
                if len(vals) == 3:
                    return all(v >= 0.65 for v in vals)
                if len(vals) == 1:
                    return vals[0] >= 0.70
            return False

        white_chars = [c for c in chars if _is_light_color(c.get("non_stroking_color"))]
        logger.debug(f"whitetext scan: {len(white_chars)} light-colored chars out of {len(chars)}")
        if not white_chars:
            return {}

        # Group characters into lines by y-position (tolerance 4 pts)
        lines_dict = defaultdict(list)
        for c in white_chars:
            y_key = round(float(c.get("top", 0)) / 4) * 4
            lines_dict[y_key].append(c)

        line_texts = []
        for y in sorted(lines_dict.keys()):
            line_chars = sorted(lines_dict[y], key=lambda c: float(c.get("x0", 0)))
            text = "".join(c.get("text", "") for c in line_chars)
            if text.strip():
                line_texts.append(text)

        full_text = "\n".join(line_texts)
        logger.debug(f"whitetext scan text: {full_text[:300]!r}")

        found = {}
        em = _extract_email(full_text)
        if em:
            found["email"] = em
        ph = _extract_phone(full_text)
        if ph:
            found["phone"] = ph
        if found:
            logger.info(f"whitetext scan found: {found}")
        return found
    except Exception as exc:
        logger.debug(f"_pdfplumber_whitetext_scan: {exc}")
        return {}


def _pdfplumber_contact_scan(path):
    """Extract email/phone from page 1 of a styled PDF (icon-font contact bars).

    pdfplumber's extract_words() retrieves all glyph boxes including those on
    coloured backgrounds that pymupdf sometimes misses when fonts are embedded
    with a non-standard encoding.  Used as a second-chance pass before the AI
    fallback fires for email/phone.

    Returns a dict with 'email' and/or 'phone' keys — only the found ones.
    """
    if not HAS_PDFPLUMBER:
        return {}
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            if not pdf.pages:
                return {}
            page = pdf.pages[0]
            words = page.extract_words(x_tolerance=5, y_tolerance=5)
            plain = page.extract_text() or ""
            word_line = " ".join(w.get("text", "") for w in words if w.get("text", "").strip())
            combined = word_line + "\n" + plain
        found = {}
        em = _extract_email(combined)
        if em:
            found["email"] = em
        ph = _extract_phone(combined)
        if ph:
            found["phone"] = ph
        if found:
            logger.info(f"pdfplumber contact scan found: {list(found.keys())}")
        return found
    except Exception as exc:
        logger.debug(f"_pdfplumber_contact_scan: {exc}")
        return {}


def parse_resume_with_llm_text(path):
    """Pure-Ollama resume parsing — no regex field engine.

    1. PDF → clean text (column-split aware).
    2. One Ollama call extracts the identity header into JSON (name/title/contact) —
       this is where a model genuinely beats rules, and the short header keeps it fast
       and safe from hallucination.
    3. Content sections are cut from the text at their heading lines and used verbatim
       — faithful to the resume (no invented content) and instant.
    Returns (fields_dict, "llm_text").
    """
    raw_text = extract_resume_text(path, "pdf")

    # For sidebar-style 2-column PDFs the pymupdf word-order extractor interleaves
    # left-column content (KEY SKILLS) with right-column content (WORK EXPERIENCE)
    # when the gutter is too close to the page edge for the block-based detector.
    # Try the pdfplumber crop-based extractor which reads each column independently.
    if str(path).lower().endswith(".pdf"):
        # Preserve any OCR-injected contact info BEFORE possibly replacing raw_text.
        # Icon-based contact bars (image text) are only visible to OCR; pdfplumber
        # reads the PDF vector layer and would miss them.
        _ocr_email = _extract_email(raw_text)
        _ocr_phone = _extract_phone(raw_text)
        two_col = _pdfplumber_two_col_text(path)
        if two_col and len(two_col.strip()) >= len(raw_text.strip()) * 0.5:
            raw_text = two_col
            # Re-inject OCR contact info if not already present in the new text
            _post = []
            if _ocr_email and _ocr_email not in raw_text:
                _post.append(f"Email: {_ocr_email}")
            if _ocr_phone and _ocr_phone not in raw_text:
                _post.append(f"Phone: {_ocr_phone}")
            if _post:
                raw_text = "\n".join(_post) + "\n" + raw_text

    # Strip standalone page numbers (e.g. "\n1\n", "\n 2 \n", "\n3" at end)
    raw_text = re.sub(r'\n[ \t]*\d{1,3}[ \t]*(?=\n|$)', '', raw_text)

    # ── 1. Base parse via fast regex (instant) ───────────────────────────────────
    font_name = _extract_name_from_pdf_fonts(path)
    result = parse_resume_text(raw_text, name_hint=font_name)

    # Clear name/title if quick_parse grabbed a degree string instead of real values
    _degree_abbr_re = re.compile(
        r'^(M\.?\s?Sc\.?|B\.?\s?Sc\.?|M\.?\s?B\.?\s?A\.?|Ph\.?\s?D\.?'
        r'|M\.?\s?Tech\.?|B\.?\s?Tech\.?|B\.?\s?E\.?|D\.?\s?Pharm\.?|B\.?\s?Pharm\.?)$',
        re.IGNORECASE
    )
    _degree_phrase_re = re.compile(
        r'\b(master|bachelor|m\.?sc|b\.?sc|m\.?tech|b\.?tech|diploma|'
        r'certificate|analytical chemistry|pharmacy|life science)\b',
        re.IGNORECASE
    )
    _name = result.get("full_name", "").strip()
    if not _name or _degree_abbr_re.match(_name):
        logger.info(f"Name '{_name}' looks like a degree — clearing for AI fallback")
        result["full_name"] = ""
    _title = result.get("title", "").strip()
    if _degree_phrase_re.search(_title):
        logger.info(f"Title '{_title}' looks like a degree — clearing for AI fallback")
        result["title"] = ""

    # Detect name/title swap: if the "name" field contains a role keyword and the
    # "title" field looks like a person's name (2+ words, no role keyword), swap them.
    # Happens when font-size extractor picks the job-title line as the largest text.
    _role_kw_re = re.compile(
        r'\b(analyst|engineer|lead|manager|consultant|developer|specialist|'
        r'executive|officer|director|architect|scientist|associate|coordinator|'
        r'validation|quality|compliance|pharmacist|chemist|technician|programmer)\b',
        re.IGNORECASE
    )
    _cur_name  = result.get("full_name", "").strip()
    _cur_title = result.get("title", "").strip()
    if (_cur_name and _cur_title
            and _role_kw_re.search(_cur_name)
            and not _role_kw_re.search(_cur_title)
            and len(_cur_title.split()) >= 2):
        logger.info(f"Name/title swap detected — swapping '{_cur_name}' ↔ '{_cur_title}'")
        result["full_name"], result["title"] = _cur_title, _cur_name

    # ── 2. Verbatim section extraction ─────────────────────────────────────────────
    # Exact keywords only. CRITICAL: filter out word-wrap false matches by requiring
    # a blank line (double \n) before any heading — "Software \nQualifications" is a
    # wrapped sentence, not a section heading; real headings always follow a blank line.
    _EXACT_HEADINGS = [
        "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY",
        "PROFESSIONAL SUMMARY", "PROFILE SUMMARY", "APPLICATIONS SUMMARY",
        "AREAS OF EXPERTISE", "TECHNICAL SKILLS", "CORE COMPETENCIES", "KEY SKILLS",
        "ACADEMIC BACKGROUND", "ACADEMIC QUALIFICATIONS",
        "SUMMARY", "EXPERIENCE", "EMPLOYMENT", "SKILLS", "EDUCATION",
        "QUALIFICATIONS", "CERTIFICATIONS", "PROJECTS", "ACHIEVEMENTS",
        "AWARDS", "REFERENCES",
        "TOOLS & APPLICATIONS", "TOOLS AND APPLICATIONS", "TOOLS",
        # Title-case headings (Arun-style resumes)
        "Professional Summary", "Work Experience", "Relevant Project Experience",
        "Key Skills", "Technical Skills",
        "Tools & Applications", "Tools and Applications",
    ]
    _section_heading_re = re.compile(
        r'(?:^|(?<=\n))\s*('
        + '|'.join(re.escape(k) for k in _EXACT_HEADINGS)
        # Allow any trailing non-word, non-newline chars (e.g. " :-", " —", " :") after
        # the keyword so "Education :-" and "Certifications :-" are also matched.
        + r')[^\w\n]*(?=\n|$)',
        re.IGNORECASE
    )

    def _heading_field(h):
        u = h.strip().upper()
        if u in ("SUMMARY", "PROFILE SUMMARY", "PROFESSIONAL SUMMARY",
                 "APPLICATIONS SUMMARY"):
            return "summary"
        if u in ("EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
                 "EMPLOYMENT HISTORY", "EMPLOYMENT"):
            return "experience"
        if u in ("SKILLS", "TECHNICAL SKILLS", "KEY SKILLS", "CORE COMPETENCIES",
                 "AREAS OF EXPERTISE"):
            return "skills"
        if u in ("EDUCATION", "QUALIFICATIONS", "ACADEMIC BACKGROUND",
                 "ACADEMIC QUALIFICATIONS", "EDUCATIONAL QUALIFICATIONS",
                 "EDUCATIONAL BACKGROUND"):
            return "education"
        if u in ("CERTIFICATIONS", "CERTIFICATION"):
            return "certifications"
        if u in ("ACHIEVEMENTS", "ACHIEVEMENT", "AWARDS", "AWARD"):
            return "achievements"
        if u in ("PROJECTS", "PROJECT", "RELEVANT PROJECT EXPERIENCE"):
            return "projects"
        if u in ("REFERENCES",):
            return "references"
        return None

    # Validate headings: must be ALL CAPS (letter characters only).
    # This filters word-wrapped sentence continuations like "...Software\nQualifications"
    # (mixed case) while accepting real headings like "PROFILE SUMMARY", "EXPERIENCE".
    # Title-case headings are accepted as fallback if preceded by a blank line OR if the
    # preceding line ends with a sentence terminator (the regex already ensures the keyword
    # appears alone on its line, so a prior sentence-end is a reliable heading signal).
    def _preceded_by_blank_or_sentence_end(pos):
        if pos <= 1:
            return True
        line_end = pos - 1
        line_start = raw_text.rfind('\n', 0, line_end)
        prev_line = raw_text[line_start + 1:line_end] if line_start != -1 else raw_text[:line_end]
        stripped = prev_line.strip()
        if not stripped:
            return True  # blank line before heading
        # A sentence/list-item ending immediately before the heading
        return stripped[-1] in '.):–—'

    # Single-word headings that are so section-specific they're safe without
    # requiring blank-line / sentence-end context (they almost never appear
    # mid-sentence in resume text).
    _STRONG_HEADINGS = frozenset([
        'CERTIFICATIONS', 'CERTIFICATION', 'ACHIEVEMENTS', 'ACHIEVEMENT',
        'AWARDS', 'AWARD', 'REFERENCES', 'PROJECTS', 'QUALIFICATIONS',
    ])

    candidates = [(m.start(), m.group(1).strip()) for m in _section_heading_re.finditer(raw_text)]
    headings = []
    for pos, hd in candidates:
        letters = re.sub(r'[^a-zA-Z]', '', hd)
        is_all_caps = bool(letters) and letters == letters.upper()
        # Multi-word known headings (e.g. "Key Skills", "Profile Summary") are accepted
        # without requiring all-caps or a preceding blank line — they are unambiguous
        # as section headings because the regex already anchors them to their own line.
        is_multiword_known = len(hd.split()) >= 2
        # Unambiguously section-specific single words are accepted without context
        is_strong = hd.upper() in _STRONG_HEADINGS
        if is_all_caps or is_multiword_known or is_strong or _preceded_by_blank_or_sentence_end(pos):
            headings.append((pos, hd))
        else:
            logger.debug(f"Skipped '{hd}' at {pos}: single-word mixed-case with no blank line before")
    logger.info(f"Validated headings: {[h for _, h in headings[:12]]}")

    _qual_exp_date_re = re.compile(
        r'(?:JAN(?:UARY)?|FEB(?:RUARY)?|MAR(?:CH)?|APR(?:IL)?|MAY|JUN(?:E)?|'
        r'JUL(?:Y)?|AUG(?:UST)?|SEP(?:TEMBER)?|OCT(?:OBER)?|NOV(?:EMBER)?|DEC(?:EMBER)?)'
        r'\s+\d{4}\s*[-–]',
        re.I,
    )

    verbatim = {}
    for i, (pos, heading) in enumerate(headings):
        field = _heading_field(heading)
        nl_pos = raw_text.find("\n", pos)
        content_start = nl_pos + 1 if nl_pos != -1 else pos + len(heading)
        content_end = headings[i + 1][0] if i + 1 < len(headings) else len(raw_text)
        text = raw_text[content_start:content_end].strip()

        # "Qualifications" headings are ambiguous — check content first.
        # If the content has experience date ranges (MONTH YEAR – ...) it's a
        # job-entry block mis-routed by a label-sidebar page-2 "Qualifications" label.
        # Append it to experience instead of education.
        if heading.upper() == "QUALIFICATIONS" and text and _qual_exp_date_re.search(text):
            if "experience" in verbatim:
                verbatim["experience"] += "\n" + text
            else:
                verbatim["experience"] = text
            logger.info(
                "Verbatim 'experience' +%d chars re-routed from 'Qualifications' (date ranges found)",
                len(text),
            )
            continue

        if not field:
            continue
        # Explicit heading labels (e.g. "EDUCATION") always override an earlier
        # ambiguous heading (e.g. "Qualifications") for the same field.
        _is_explicit = heading.upper() in (
            "EDUCATION", "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
            "SUMMARY", "PROFESSIONAL SUMMARY", "SKILLS", "TECHNICAL SKILLS",
            "CERTIFICATIONS", "PROJECTS",
        )
        if field in verbatim and not _is_explicit:
            continue
        if text:
            verbatim[field] = text
            logger.info(f"Verbatim '{field}' ({len(text)} chars) from '{heading}'")

    for field, text in verbatim.items():
        result[field] = text

    # ── 2a-post. Clean verbatim skills: remove website headings, URLs, split merged bullets ──
    if result.get("skills"):
        _vskill_url = re.compile(
            r"^(https?://|www\.|linkedin|github|gitlab|portfolio|website|blog)\S*$",
            re.I,
        )
        _vskill_hdr = re.compile(
            r"^(websites?(\s*,\s*|\s+)portfolios?(\s+and\s+profiles?)?|websites?|"
            r"portfolios?\s+and\s+profiles?|and\s+profiles?|profiles?\s+and|"
            r"personal\s+information|personal\s+details?|declaration|"
            r"additional\s+information|other\s+information)\s*$",
            re.I,
        )
        vsk_lines = []
        for sl in result["skills"].splitlines():
            s = sl.strip()
            if not s:
                continue
            if _vskill_url.match(s):
                continue
            bare = re.sub(r"[^a-zA-Z\s]", " ", s).strip()
            if _vskill_hdr.match(bare):
                continue
            # Split merged bullet items: "• Skill A • Skill B" → two separate items
            if s.count("•") > 1 or s.count("●") > 1:
                for p in re.split(r"\s*[•●▪◦]\s*", s):
                    p = p.strip()
                    if p:
                        vsk_lines.append(p)
            else:
                vsk_lines.append(s.lstrip("•● -").strip())
        # Remove the person's own name/title if they leaked into skills from the
        # column boundary (left sidebar ends → right column starts with name+title).
        _sk_name = result.get("full_name", "").strip().upper()
        _sk_title = result.get("title", "").strip().lower()
        vsk_lines = [
            l for l in vsk_lines
            if not (l.strip().upper() == _sk_name or
                    (_sk_title and l.strip().lower() == _sk_title))
        ]
        result["skills"] = "\n".join(l for l in vsk_lines if l.strip())

    # ── 2a-post. Clean verbatim education: stop at PERSONAL INFORMATION, remove personal lines ──
    if result.get("education"):
        # "PERSONAL INFORMATION" can be on one line OR split across two lines
        _edu_stop_pi = re.compile(
            r"\n[ \t]*(?:PERSONAL(?:[ \t]+(?:INFORMATION|DETAILS?))?|"
            r"OTHER[ \t]+PERSONAL|DECLARATION)\s*(?:\n|$)",
            re.I,
        )
        _em_pi = _edu_stop_pi.search(result["education"])
        if _em_pi:
            trimmed_pi = result["education"][:_em_pi.start()].strip()
            if trimmed_pi:
                logger.info(f"Education trimmed at PERSONAL INFORMATION boundary")
                result["education"] = trimmed_pi
        _edu_pi_pat = re.compile(
            r"\b(father|mother|parent|guardian|date.{0,5}birth|born|dob\b|gender|sex\b|"
            r"marital|nationality|citizen|passport|religion|caste|languages\s+known|"
            r"permanent\s+address|current\s+address|pin\s*code|aadhar|pan\b|"
            r"place\s+of\s+birth|age\b|blood\s+group|declaration|personal\s+information|"
            r"personal\s+details?)\b",
            re.I,
        )
        # Also drop standalone noise words like bare "PERSONAL" or "INFORMATION"
        _edu_noise = {"PERSONAL", "INFORMATION", "DETAILS", "DECLARATION", "OTHER"}
        edu_clean = [
            l for l in result["education"].splitlines()
            if not _edu_pi_pat.search(l) and l.strip().upper() not in _edu_noise
        ]
        result["education"] = "\n".join(edu_clean).strip()

    # ── 2a-edu-garbage: If education content looks like experience, clear or rescue ──
    # Symptom: label-sidebar CVs with multiple page labels (e.g. "Qualifications" on
    # page 2 sidebar) cause experience entries from later pages to land in the education
    # field.  Two sub-cases:
    #   A) Content has experience date ranges (MONTH YEAR – MONTH YEAR/PRESENT) →
    #      it's real experience that got mis-routed → merge back into experience, clear edu.
    #   B) Content has work-verb bullets but NO date ranges and NO edu keywords →
    #      it's a professional-skills/qualifications blurb → clear, let AI find real education.
    if result.get("education"):
        _edu_text = result["education"]
        _has_edu_kw = re.search(
            r'\b(university|college|institute|school|bachelor|master|b\.e\b|b\.tech\b|'
            r'b\.com\b|m\.tech\b|m\.e\b|mba\b|phd\b|diploma|degree|cgpa|gpa|'
            r'10\+2|hsc\b|sslc\b|12th|10th|graduation|affiliated)\b',
            _edu_text, re.I,
        )
        # Count experience-style lines — bullet char is optional (some PDFs use spaces only).
        _exp_verbs = re.findall(
            r'(?:^|\n)[ \t]*[•\-]?[ \t]*(?:Expertise|Managed|Performed|Prepared|Preparati|'
            r'Worked|Working|Ensure|Contribut|Execut|Conduct|Complet|Monitor|Analys|'
            r'Maintained|Develop|Review|Coordinat|Having|Exposure|Responsible|'
            r'Handling|Support|Assist|Troubleshoot)',
            _edu_text,
        )
        if not _has_edu_kw and len(_exp_verbs) >= 2:
            # Check whether the content contains actual experience date ranges
            # (month name + year → dashes → month/year/present).
            _exp_date_re = re.compile(
                r'(?:JAN(?:UARY)?|FEB(?:RUARY)?|MAR(?:CH)?|APR(?:IL)?|MAY|JUN(?:E)?|'
                r'JUL(?:Y)?|AUG(?:UST)?|SEP(?:TEMBER)?|OCT(?:OBER)?|NOV(?:EMBER)?|DEC(?:EMBER)?)'
                r'\s+\d{4}\s*[-–]',
                re.I,
            )
            if _exp_date_re.search(_edu_text):
                # Real experience entries mis-routed to education — rescue them.
                logger.info(
                    "Education field contains experience entries (date ranges found) "
                    "— merging into experience and clearing education"
                )
                if result.get("experience"):
                    result["experience"] = result["experience"] + "\n" + _edu_text
                else:
                    result["experience"] = _edu_text
            else:
                logger.info(
                    "Education field contains %d experience-style bullets (no date ranges) "
                    "— clearing for AI fallback", len(_exp_verbs)
                )
            result["education"] = ""

    # ── 2a-post. Reject title if it looks like a sentence fragment ────────────────
    _title_check = result.get("title", "").strip()
    if _title_check:
        _tw0 = _title_check.split()[0]
        _is_frag = (
            _tw0[:1].islower()            # starts lowercase → mid-sentence fragment
            or _tw0.lower().endswith("ing")  # gerund (e.g. "enhancing", "managing")
            or ("." in _title_check[:-1])    # period in the middle → sentence boundary
            or len(_title_check.split()) > 7  # too long for a title
        )
        if _is_frag:
            logger.info(f"Title '{_title_check[:60]}' looks like sentence fragment — clearing")
            result["title"] = ""

    # ── 2b. Post-process education: stop at project/experience boundary ───────
    # Handles resumes (e.g. Arun Sivakumar) where "Relevant Project Experience"
    # is directly below EDUCATION with no blank line — so heading detection missed
    # it as a boundary and the full project table was pulled into education.
    if result.get("education"):
        _edu_trim_re = re.compile(
            r'(?:\n|^)[ \t]*(?:Relevant\s+)?Project\s+Experience[:\s]*(?=\n|$)'
            r'|(?:\n|^)[ \t]*Work\s+Experience[:\s]*(?=\n|$)'
            r'|(?:\n|^)[ \t]*Employment\s+History[:\s]*(?=\n|$)',
            re.IGNORECASE,
        )
        _em = _edu_trim_re.search(result["education"])
        if _em:
            trimmed = result["education"][:_em.start()].strip()
            if trimmed:
                logger.info(
                    f"Education trimmed at project/experience boundary "
                    f"({len(result['education'])} → {len(trimmed)} chars)"
                )
                result["education"] = trimmed

    # ── 2b-fix. Restore truncated last lines (education + experience) ───────────
    # The two-column text splitter can clip the tail of a line that spans the
    # full page width, e.g. "Maharaja Engineering Colle" instead of
    # "Maharaja Engineering College, Coimbatore - 2016 to 2010".
    # Heuristic: if the last line ends with a plain letter (no punctuation,
    # digit, or closing bracket) the line was cut mid-word.
    # Fix: re-read each page with pdfplumber extract_text() (no column split)
    # and replace the truncated line with the full version found there.
    def _restore_truncated_last_line(field_text, full_page_text):
        last_line = field_text.rsplit("\n", 1)[-1].strip()
        if not last_line or not last_line[-1].isalpha():
            return field_text
        prefix = last_line[:20]
        idx = full_page_text.find(prefix)
        if idx == -1:
            return field_text
        line_end = full_page_text.find("\n", idx)
        full_line = full_page_text[
            idx : line_end if line_end != -1 else idx + 300
        ].strip()
        if len(full_line) > len(last_line):
            return field_text[: -len(last_line)] + full_line
        return field_text

    if str(path).lower().endswith(".pdf") and any(
        result.get(f) for f in ("education", "experience")
    ):
        try:
            import pdfplumber as _fix_plumber
            _fix_full_text = ""
            with _fix_plumber.open(str(path)) as _fix_pdf:
                for _fix_pg in _fix_pdf.pages:
                    _fix_full_text += (_fix_pg.extract_text() or "") + "\n"
            for _fix_field in ("education", "experience"):
                if result.get(_fix_field):
                    _before = result[_fix_field]
                    _after = _restore_truncated_last_line(_before, _fix_full_text)
                    if _after != _before:
                        result[_fix_field] = _after
                        logger.info(
                            "%s last line restored: %r → %r",
                            _fix_field,
                            _before.rsplit("\n", 1)[-1].strip(),
                            _after.rsplit("\n", 1)[-1].strip(),
                        )
        except Exception as _fix_exc:
            logger.debug("Truncated last line restore failed: %s", _fix_exc)

    # ── 2c. Styled-bar contact scan ────────────────────────────────────────────
    # For PDFs with icon-font contact bars (dark background, envelope/phone icons)
    # the main pymupdf pass sometimes fails to decode the email/phone characters.
    # Run a dedicated pdfplumber word-box scan on page 1 before the AI fallback.
    if str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        _extra_contact = _pdfplumber_contact_scan(path)
        for _cf, _cv in _extra_contact.items():
            if not result.get(_cf) and _cv:
                result[_cf] = _cv
                logger.info(f"Styled-bar contact scan filled '{_cf}': {_cv}")

    # ── 2d. Garbage detection — clear fields that contain header metadata so the
    # AI fallback can re-extract them correctly.  This handles layouts where the
    # person's all-caps name (e.g. "SIVARANJANI D") trips the company-name heuristic
    # in find_sections and causes header lines to fill the experience field.
    _header_meta_re = re.compile(
        r'total\s+experience|current\s+employer|current\s+designation'
        r'|current\s+ctc|expected\s+ctc|notice\s+period',
        re.IGNORECASE,
    )
    _exp_val = result.get("experience", "")
    if _exp_val and _header_meta_re.search(_exp_val):
        logger.info(
            f"Clearing garbage experience (header metadata detected) — AI fallback will fill it"
        )
        result["experience"] = ""

    # A single-word title that is not a recognised role descriptor is likely a
    # fragment from a split header box (e.g. "Product" from "Product Regulatory
    # Analyst" truncated at the column gutter).  Clear it so the next step can fill it.
    _title_val = result.get("title", "").strip()
    _role_fragment_ok = re.compile(
        r'\b(analyst|engineer|lead|manager|consultant|developer|specialist|'
        r'designer|director|officer|architect|executive|scientist|associate|'
        r'coordinator|technician|programmer|advisor|writer|researcher)\b',
        re.IGNORECASE,
    )
    if _title_val and len(_title_val.split()) == 1 and not _role_fragment_ok.search(_title_val):
        logger.info(
            f"Clearing single-word title fragment '{_title_val}' — regex/AI fallback will fill it"
        )
        result["title"] = ""

    # Try to recover title directly from "Current Designation : <value>" or
    # "Designation : <value>" pattern in the raw text — reliable for header-box resumes.
    if not result.get("title"):
        # Iterate all Designation matches: the header may have a truncated value
        # (e.g. "Current Designation : Product" split across columns) while the
        # work-experience section further down has the full title.
        # Use the first match that is at least 2 words or contains a role keyword.
        for _desig_m in re.finditer(
            r'(?:current\s+)?designation\s*[:\-]\s*(.+)',
            raw_text[:5000], re.IGNORECASE
        ):
            _desig_val = _desig_m.group(1).strip().split('\n')[0].strip()
            if _desig_val and (len(_desig_val.split()) >= 2 or _role_fragment_ok.search(_desig_val)):
                result["title"] = _desig_val
                logger.info(f"Title from Designation regex: '{_desig_val}'")
                break

    # Title last-resort before AI: ALL-CAPS role title at the top of experience
    # (e.g. "VALIDATION ENGINEER, 11/2021-Current") — avoids an AI call entirely.
    if not result.get("title") and result.get("experience"):
        _exp_fl = result["experience"].strip().split('\n')[0].strip()
        _exp_title_m = re.match(
            r'^([A-Z][A-Z0-9\s/&\-]{2,45}?)(?:\s*[,\-–]\s*\d|\s*$)', _exp_fl
        )
        if _exp_title_m:
            _cand = _exp_title_m.group(1).strip()
            if _role_fragment_ok.search(_cand):
                result["title"] = _cand.title()
                logger.info(f"Title from experience first line: '{result['title']}'")

    # ── 3. AI fallback — fires only when a field is still empty ─────────────────
    # Email/phone also get AI fallback for icon-based contact bars where regex fails.
    _AI_FALLBACK = {
        "full_name": (
            "What is the full name of the person in this resume? "
            "Return ONLY the name, nothing else.\n\nRESUME HEADER:\n",
            raw_text[:600], False, 20, 512,
        ),
        "title": (
            "What is the current professional job title of the person in this resume? "
            "Look for lines labelled 'Current Designation:', 'Designation:', 'Title:', "
            "or a role title just below the person's name. "
            "Return ONLY the job title (e.g. 'Product Regulatory Analyst'). "
            "Do NOT return a sentence, a degree, or a summary.\n\nRESUME HEADER:\n",
            raw_text[:1200], False, 25, 512,
        ),
        "summary": (
            "Extract the professional summary section verbatim from this resume. "
            "Do not paraphrase.\n\nRESUME:\n",
            raw_text[:3000], False, 600, 2048,
        ),
        "experience": (
            "Extract the work experience from this resume. "
            "The experience section may NOT have a heading like 'Work Experience' — "
            "look for company/employer names (often in ALL CAPS or Title Case) followed by "
            "Duration, Designation/Title, and bullet-point responsibilities. "
            "Return ONLY the verbatim work experience text — company, dates, role, and bullets. "
            "Do NOT return the profile summary or certifications.\n\nRESUME:\n",
            raw_text, False, 1600, 4096,
        ),
        "skills": (
            "Extract all skills and tools from this resume. "
            'Return ONLY JSON: {"skills":["s1","s2",...]}.\n\nRESUME:\n',
            raw_text[:3000], True, 400, 1024,
        ),
        "education": (
            "Extract the education section verbatim from this resume.\n\nRESUME:\n",
            raw_text, False, 300, 1024,
        ),
        "email": (
            "Find the email address in this resume. Return ONLY the email address, nothing else.\n\nRESUME:\n",
            raw_text[:800], False, 15, 512,
        ),
        "phone": (
            "Find the phone number in this resume. Return ONLY the phone number "
            "(include country code if present), nothing else.\n\nRESUME:\n",
            raw_text[:800], False, 15, 512,
        ),
    }

    import time as _time
    _ai_budget_start = _time.monotonic()

    for field, (prompt, ctx_text, as_json, num_predict, num_ctx) in _AI_FALLBACK.items():
        if result.get(field):
            continue  # already filled — skip AI
        if _time.monotonic() - _ai_budget_start > _OLLAMA_BUDGET_SECS:
            logger.warning(f"AI budget exhausted — skipping remaining fallback fields")
            break
        logger.info(f"AI fallback for '{field}'")
        try:
            raw = _ollama_chat(prompt + ctx_text, as_json=as_json,
                               num_predict=num_predict, num_ctx=num_ctx)
            if as_json:
                items = raw.get("skills", []) if isinstance(raw, dict) else []
                if items:
                    result[field] = "\n".join(str(s) for s in items if str(s).strip())
            else:
                if raw and raw.strip():
                    result[field] = raw.strip()
            if result.get(field):
                logger.info(f"AI filled '{field}' ({len(result[field])} chars)")
        except Exception as e:
            logger.warning(f"AI fallback '{field}' failed: {e}", exc_info=True)

    # ── 3b. Post-AI title validation ───────────────────────────────────────────
    # The AI sometimes returns a summary paragraph or "SUMMARY: ..." as the title.
    # Apply the same fragment check used at step 2a-post; if invalid, try to
    # recover the title from the first line of the experience section.
    _ai_title = result.get("title", "").strip()
    if _ai_title:
        _tw0 = _ai_title.split()[0] if _ai_title.split() else ""
        _ai_title_bad = (
            _tw0[:1].islower()
            or _tw0.lower().endswith("ing")
            or "." in _ai_title[:-1]
            or len(_ai_title.split()) > 7
            or ":" in _tw0           # "SUMMARY:" / "OBJECTIVE:" type prefix
            or _ai_title.startswith("```")  # markdown block
        )
        if _ai_title_bad:
            logger.info(f"Post-AI: title '{_ai_title[:60]}' looks like paragraph — clearing")
            result["title"] = ""

    # ── 4. Post-AI contact validation ──────────────────────────────────────────
    # The AI model sometimes returns text fragments instead of real email/phone
    # (e.g. "sion Report." for email, or a hallucinated phone number).
    # Validate and clear any field that does not match the expected format.
    _valid_email_re = re.compile(
        r'^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$'
    )
    _valid_phone_re = re.compile(
        r'^\+?[\d][\d\s\-\.\(\)]{6,24}$'
    )
    for _vf, _vr in [("email", _valid_email_re), ("phone", _valid_phone_re)]:
        _val = result.get(_vf, "").strip()
        if _val and not _vr.match(_val):
            logger.warning(f"Post-AI validation: clearing invalid {_vf}='{_val}'")
            result[_vf] = ""

    # ── 4b. Cross-validate AI phone: clear if its digits aren't in the raw text ─
    # The AI sometimes hallucinate a plausible-looking phone number that passes the
    # format check above but was never in the document.  Compare digit sequences:
    # if the AI phone's digits don't appear anywhere in the raw text, discard it.
    if result.get("phone"):
        _ai_ph_digits = re.sub(r'\D', '', result["phone"])
        _raw_digits   = re.sub(r'\D', '', raw_text)
        if len(_ai_ph_digits) >= 10 and _ai_ph_digits not in _raw_digits:
            logger.warning(
                f"Phone '{result['phone']}' digits not in raw text — AI hallucination, clearing"
            )
            result["phone"] = ""

    # ── 4c. Aggressive +91 / 10-digit regex on full raw text ──────────────────
    # Runs after validation + cross-check have cleared garbage values.
    # Targets Indian numbers explicitly (common in uploaded resumes).
    if not result.get("email") or not result.get("phone"):
        _email_agr_re = re.compile(
            r'\b[A-Za-z0-9][A-Za-z0-9._%+\-]*@[A-Za-z0-9.\-]+\.[A-Za-z]{2,6}\b'
        )
        _ph_91_re = re.compile(r'\+91[\s\-\.]?\d[\d\s\-\.]{7,14}\d')
        _ph_10_re = re.compile(r'(?<!\d)[6-9]\d{9}(?!\d)')

        if not result.get("email"):
            _agr_em = _email_agr_re.search(raw_text)
            if _agr_em:
                _em_val = _agr_em.group().strip().rstrip(".,;)")
                result["email"] = _em_val
                logger.info(f"Aggressive regex email: {_em_val}")

        if not result.get("phone"):
            _agr_ph = _ph_91_re.search(raw_text) or _ph_10_re.search(raw_text)
            if _agr_ph:
                result["phone"] = _agr_ph.group().strip()
                logger.info(f"Aggressive regex phone: {result['phone']}")

    # ── 4d. White-text pdfplumber scan (contact bar on dark background) ────────
    # pdfplumber chars expose non_stroking_color (fill color).  White text on a
    # dark background (navy contact bar) has color ≥ 0.7 in all channels.
    # This specifically targets the contact bar without touching the body text.
    if str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        _wt_contact = _pdfplumber_whitetext_scan(path)
        for _wtf, _wtv in _wt_contact.items():
            if not result.get(_wtf) and _wtv:
                result[_wtf] = _wtv
                logger.info(f"White-text scan filled '{_wtf}': {_wtv}")

    # ── 5. Final rawdict contact scan ──────────────────────────────────────────
    # Last resort: use pymupdf rawdict (char-level) to find email/phone that may
    # live in styled header bars with non-standard font encoding.
    # Runs only after validation has cleared garbage values.
    if str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        _rd_contact = _pymupdf_rawdict_scan(path)
        for _rdf, _rdv in _rd_contact.items():
            if not result.get(_rdf) and _rdv:
                result[_rdf] = _rdv
                logger.info(f"rawdict scan filled '{_rdf}': {_rdv}")

    # ── 5d. Inverted Tesseract OCR — highest accuracy for dark contact bars ──────
    # MUST run before step 6 (non-inverted OCR).  Step 6 struggles with white text
    # on a dark/navy background and produces OCR artefacts (e.g. "Barun.sivakumar"
    # instead of "arun.sivakumar").  Running the inverted pass first sets correct
    # values so step 6 and 6b are skipped entirely.
    if str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        _inv_ocr_early = _tesseract_inverted_ocr_text(path)
        if _inv_ocr_early.strip():
            if not result.get("email"):
                _em_early = _extract_email(_inv_ocr_early)
                if _em_early:
                    result["email"] = _em_early
                    logger.info(f"Step 5d inverted OCR email: {_em_early}")
            if not result.get("phone"):
                _ph_early = _extract_phone(_inv_ocr_early)
                if _ph_early:
                    _ph_early_digits = re.sub(r"\D", "", _ph_early)
                    if 10 <= len(_ph_early_digits) <= 15:
                        result["phone"] = _ph_early
                        logger.info(f"Step 5d inverted OCR phone: {_ph_early}")

    # ── 6. Tesseract OCR contact scan ──────────────────────────────────────────
    # pytesseract renders page 1 as an image and OCRs the top 35% — the only
    # reliable method for dark-background contact bars (white text + icon fonts)
    # that are invisible to all PDF text-layer extraction methods.
    # Runs as a final pass; fills email/phone directly from OCR output.
    if str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        _tess_contact = _tesseract_contact_scan(path)
        for _tf, _tv in _tess_contact.items():
            if not result.get(_tf) and _tv:
                result[_tf] = _tv
                logger.info(f"Tesseract OCR filled '{_tf}': {_tv}")

    # ── 6b. Inverted-image OCR + Ollama extraction ────────────────────────────
    # The standard tesseract pass (step 6) renders white-on-dark text which gives
    # Tesseract poor contrast.  This step:
    #   1. Re-renders at 300 DPI and INVERTS the image (dark→light, white→dark).
    #   2. Runs Tesseract on both the inverted RGB image and the inverted red channel
    #      (maximum contrast for navy backgrounds).
    #   3. Tries direct regex on the combined OCR output.
    #   4. If email/phone are still missing, passes the OCR text to Ollama to
    #      extract and return them as JSON — Ollama acts as a smart OCR post-processor.
    if str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        _inv_ocr = _tesseract_inverted_ocr_text(path)
        if _inv_ocr.strip():
            # Step A: direct regex on inverted OCR text
            if not result.get("email"):
                _em_inv = _extract_email(_inv_ocr)
                if _em_inv:
                    result["email"] = _em_inv
                    logger.info(f"Inverted OCR direct email: {_em_inv}")
            if not result.get("phone"):
                _ph_inv = _extract_phone(_inv_ocr)
                if _ph_inv:
                    _ph_digits = re.sub(r"\D", "", _ph_inv)
                    if 10 <= len(_ph_digits) <= 15:
                        result["phone"] = _ph_inv
                        logger.info(f"Inverted OCR direct phone: {_ph_inv}")

            # Step B: Ollama extracts from OCR text when regex still failed
            if not result.get("email") or not result.get("phone"):
                _ollama_ocr_prompt = (
                    "The text below was OCR'd from the contact bar of a resume. "
                    "Find the email address and phone number and return ONLY this JSON:\n"
                    "{\"email\": \"the_email@domain.com\", \"phone\": \"+91 XXXX XXX XXX\"}\n\n"
                    "Rules:\n"
                    "- email must contain @ and end with .com / .in / .net etc.\n"
                    "- phone must be a number starting with +91 or a 10-digit mobile\n"
                    "- if you cannot find a field, use empty string \"\"\n\n"
                    "OCR TEXT:\n" + _inv_ocr[:600]
                )
                try:
                    _ocr_json = _ollama_chat(
                        _ollama_ocr_prompt, as_json=True, num_predict=80, num_ctx=1024
                    )
                    if isinstance(_ocr_json, dict):
                        _oj_email = str(_ocr_json.get("email", "")).strip().rstrip(".,;)")
                        _oj_phone = str(_ocr_json.get("phone", "")).strip()
                        _em_re = re.compile(
                            r'^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$'
                        )
                        if not result.get("email") and _oj_email and _em_re.match(_oj_email):
                            result["email"] = _oj_email
                            logger.info(f"Ollama-OCR email: {_oj_email}")
                        if not result.get("phone") and _oj_phone:
                            _ph_d = re.sub(r"\D", "", _oj_phone)
                            if 10 <= len(_ph_d) <= 15:
                                result["phone"] = _oj_phone
                                logger.info(f"Ollama-OCR phone: {_oj_phone}")
                except Exception as exc:
                    logger.debug(f"Ollama-OCR extraction: {exc}")

    # ── 7. Ollama with OCR context ─────────────────────────────────────────────
    # If Tesseract found partial results or nothing, build a combined context from
    # Tesseract OCR + raw_text and ask Ollama specifically for the missing fields.
    # Response is validated strictly — must match email/phone format exactly.
    _step7_ok = _time.monotonic() - _ai_budget_start <= _OLLAMA_BUDGET_SECS
    if _step7_ok and str(path).lower().endswith(".pdf") and (not result.get("email") or not result.get("phone")):
        # Build combined context: page-1 simple text + raw_text header
        _ollama_ctx_parts = []
        try:
            import pymupdf as _fz_mod
            _doc = _fz_mod.open(str(path))
            _p1_simple = _doc[0].get_text("text")
            _doc.close()
            if _p1_simple.strip():
                _ollama_ctx_parts.append(_p1_simple[:800])
        except Exception:
            pass
        _ollama_ctx_parts.append(raw_text[:800])
        _ollama_ctx = "\n---\n".join(filter(None, _ollama_ctx_parts))

        _valid_email_final = re.compile(
            r'^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$'
        )
        _valid_phone_final = re.compile(r'^\+?[\d][\d\s\-\.\(\)]{6,24}$')

        if not result.get("email"):
            try:
                _r = _ollama_chat(
                    "The resume below belongs to a person whose email address ends with "
                    ".com or .in. Find and return ONLY the email address — nothing else.\n\n"
                    "RESUME:\n" + _ollama_ctx,
                    as_json=False, num_predict=40, num_ctx=2048,
                )
                if _r and "@" in _r:
                    _em = re.search(
                        r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', _r
                    )
                    if _em and _valid_email_final.match(_em.group()):
                        result["email"] = _em.group()
                        logger.info(f"Ollama step-7 email: {result['email']}")
            except Exception as exc:
                logger.debug(f"Ollama step-7 email: {exc}")

        if not result.get("phone"):
            try:
                _r = _ollama_chat(
                    "The resume below belongs to a person with an Indian phone number "
                    "starting with +91 or a 10-digit mobile number. "
                    "Return ONLY the phone number — nothing else.\n\n"
                    "RESUME:\n" + _ollama_ctx,
                    as_json=False, num_predict=30, num_ctx=2048,
                )
                if _r:
                    _ph = re.search(r'[\+]?\d[\d\s\-\.]{7,20}\d', _r)
                    if _ph and _valid_phone_final.match(_ph.group()):
                        _ph_digits = re.sub(r'\D', '', _ph.group())
                        if 10 <= len(_ph_digits) <= 15:
                            result["phone"] = _ph.group().strip()
                            logger.info(f"Ollama step-7 phone: {result['phone']}")
            except Exception as exc:
                logger.debug(f"Ollama step-7 phone: {exc}")

    # ── 8. Ollama vision model contact scan ───────────────────────────────────
    # Renders page-1 top 30 % as an image and sends it to a vision-capable Ollama
    # model (moondream, llava, etc.).  This is the ONLY method that can read
    # email/phone from dark-background contact bars where the text layer is absent.
    #
    # To activate: run  `ollama pull moondream`  once in your terminal.
    # The code auto-detects any installed vision model — no config needed.
    if (_time.monotonic() - _ai_budget_start <= _OLLAMA_BUDGET_SECS
            and str(path).lower().endswith(".pdf")
            and (not result.get("email") or not result.get("phone"))):
        _vis = _ollama_vision_contact_scan(path)
        for _vf, _vv in _vis.items():
            if not result.get(_vf) and _vv:
                result[_vf] = _vv
                logger.info(f"Vision scan filled '{_vf}': {_vv}")

    return result, "llm_text"


def _parse_pdf_quick(path):
    """Fast rule-based PDF parse (no AI). Returns the parsed-fields dict.

    Used by bulk compare, where running the AI model on every file would be slow.
    """
    font_name = _extract_name_from_pdf_fonts(path)
    text = extract_resume_text(path, "pdf")
    return parse_resume_text(text, name_hint=font_name)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def home():
    return redirect(url_for("edit_resume"))


@app.route("/edit", methods=["GET", "POST"])
@app.route("/edit/<int:resume_id>", methods=["GET", "POST"])
def edit_resume(resume_id=None):
    with db_conn() as conn:
        if request.method == "POST":
            form_data = {k: request.form.get(k, "").strip() for k in [
                "full_name", "title", "email", "phone", "linkedin", "location", "summary",
                "skills", "experience", "education", "certifications", "projects",
            ]}

            if form_data.get("title"):
                _tc = re.split(r"[,|/\\–—]", form_data["title"])[0].strip()
                _tc = re.sub(r"[\(\[].*$", "", _tc).strip()
                _tc = re.sub(
                    r"^(experienced|skilled|dedicated|results.driven|dynamic|seasoned|"
                    r"highly experienced|passionate|motivated|proactive|hands.on)\s+",
                    "", _tc, flags=re.I,
                ).strip()
                _words = _tc.split()
                if len(_words) > 5:
                    _m = re.match(r"^((?:[A-Z][a-zA-Z]*(?:\s+|$)){1,5})", _tc)
                    _tc = _m.group(1).strip() if _m else " ".join(_words[:5])
                form_data["title"] = _tc[:80]

            resume_file = None
            if resume_id:
                existing = conn.execute(
                    "SELECT resume_file FROM resume WHERE id = %s", (resume_id,)
                ).fetchone()
                resume_file = existing["resume_file"] if existing else None

            uploaded = request.files.get("resume_file")
            parsed_data = {}

            if uploaded and uploaded.filename:
                if not allowed_file(uploaded.filename):
                    flash("Please upload a PDF or DOCX file only.", "error")
                    return redirect(
                        url_for("edit_resume", resume_id=resume_id) if resume_id
                        else url_for("edit_resume")
                    )

                name = secure_filename(uploaded.filename)
                ext = name.rsplit(".", 1)[1].lower()
                save_name = (
                    f"{slugify(form_data['full_name'] or Path(name).stem)}"
                    f"-{int(__import__('time').time())}.{ext}"
                )
                path = UPLOAD_FOLDER / save_name
                uploaded.seek(0)
                uploaded.save(path)
                resume_file = save_name

                # parse_mode: "llm" = Resume Intelligence (AI), anything else = Quick Parse
                parse_mode = request.form.get("parse_mode", "llm")
                try:
                    if ext == "pdf" and parse_mode == "llm":
                        logger.info(f"Using Resume Intelligence (AI) path for {path}")
                        try:
                            parsed_data, _ = parse_resume_with_llm_text(path)
                        except Exception as e:
                            logger.warning("Resume Intelligence failed (%s); using Quick Parse.", e, exc_info=True)
                            parsed_data = _parse_pdf_quick(path)
                    elif ext == "pdf":
                        parsed_data = _parse_pdf_quick(path)
                    else:
                        extracted = extract_resume_text(path, ext)
                        parsed_data = parse_resume_text(extracted)
                    flash("Resume uploaded and data extracted successfully.", "success")
                except Exception:
                    flash(
                        "File uploaded, but text could not be extracted. "
                        "You can still edit the profile manually.",
                        "warning",
                    )

            merged = merge_resume_data(form_data, parsed_data, overwrite=False)
            merged["resume_file"] = resume_file

            if resume_id:
                merged["slug"] = unique_slug(
                    conn, merged.get("full_name") or f"profile-{resume_id}", resume_id
                )
                merged["id"] = resume_id
                conn.execute(
                    """
                    UPDATE resume SET
                        full_name=%(full_name)s, title=%(title)s, email=%(email)s, phone=%(phone)s,
                        linkedin=%(linkedin)s, location=%(location)s, summary=%(summary)s, skills=%(skills)s,
                        experience=%(experience)s, education=%(education)s, certifications=%(certifications)s,
                        projects=%(projects)s, slug=%(slug)s, resume_file=%(resume_file)s,
                        updated_at=NOW()
                    WHERE id=%(id)s
                    """,
                    merged,
                )
                sync_skills(conn, resume_id, merged.get("skills", ""))
                return redirect(url_for("profile_detail", resume_id=resume_id))

            merged["slug"] = unique_slug(conn, merged.get("full_name") or "profile")
            cursor = conn.execute(
                """
                INSERT INTO resume (
                    full_name, title, email, phone, linkedin, location, summary, skills,
                    experience, education, certifications, projects, slug, resume_file,
                    created_at, updated_at
                ) VALUES (
                    %(full_name)s, %(title)s, %(email)s, %(phone)s, %(linkedin)s, %(location)s, %(summary)s, %(skills)s,
                    %(experience)s, %(education)s, %(certifications)s, %(projects)s, %(slug)s, %(resume_file)s,
                    NOW(), NOW()
                )
                RETURNING id
                """,
                merged,
            )
            new_id = cursor.fetchone()["id"]
            sync_skills(conn, new_id, merged.get("skills", ""))
            flash("Profile saved successfully! Here are your top matching roles.", "success")
            return redirect(url_for("profile_detail", resume_id=new_id))

        # GET
        if resume_id:
            resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()
            if not resume:
                return "Resume not found", 404
            resume = dict(resume)
        else:
            resume = {k: "" for k in [
                "full_name", "title", "email", "phone", "linkedin", "location",
                "summary", "skills", "experience", "education", "certifications",
                "projects", "slug", "resume_file",
            ]}
            resume["id"] = None

    return render_template("edit.html", resume=resume)


@app.route("/profile")
def view_current_profile():
    return redirect(url_for("profile_list"))


@app.route("/profiles")
def profile_list():
    with db_conn() as conn:
        resumes = conn.execute(
            """
            SELECT id, full_name, title, email, phone, location, resume_file, created_at, updated_at
            FROM resume
            ORDER BY updated_at DESC, id DESC
            """
        ).fetchall()
    return render_template("profile_list.html", resumes=resumes)


@app.route("/profile/<int:resume_id>")
def profile_detail(resume_id):
    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()
        if not resume:
            return "Profile not found", 404

        # Calculate top 3 matching JDs
        jds = conn.execute("SELECT * FROM job_description ORDER BY created_at DESC").fetchall()
        resume_dict = dict(resume)

        matches = []
        for jd in jds:
            jd_dict = dict(jd)
            score = calculate_match_score(resume_dict, jd_dict)
            matches.append({
                'jd_id': jd['id'],
                'jd_title': jd['title'],
                'match_percentage': score['match_percentage'],
                'matched_count': score['matched_count'],
                'total_jd_requirements': score['total_jd_requirements']
            })

        matches.sort(key=lambda x: x['match_percentage'], reverse=True)
        top_matches = matches[:3]
        is_weak = top_matches and top_matches[0]['match_percentage'] < 50

    return render_template("profile.html", resume=resume, top_matches=top_matches, is_weak=is_weak)


@app.route("/profile/<int:resume_id>/export-top-matches")
def export_top_matches(resume_id):
    """Export top 3 matching JDs with resume comparison as JSON."""
    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()
        if not resume:
            return jsonify({"error": "Profile not found"}), 404

        # Calculate top 3 matching JDs
        jds = conn.execute("SELECT * FROM job_description ORDER BY created_at DESC").fetchall()
        resume_dict = dict(resume)

        matches = []
        for jd in jds:
            jd_dict = dict(jd)
            score = calculate_match_score(resume_dict, jd_dict)
            matches.append({
                'jd': jd_dict,
                'jd_id': jd['id'],
                'jd_title': jd['title'],
                'match_percentage': score['match_percentage'],
                'matched_count': score['matched_count'],
                'total_jd_requirements': score['total_jd_requirements'],
                'matched_skills': score.get('matched_skills', []),
                'missing_skills': score.get('missing_skills', [])
            })

        matches.sort(key=lambda x: x['match_percentage'], reverse=True)
        top_matches = matches[:3]

    # Prepare export data
    export_data = {
        'resume': {
            'full_name': resume_dict.get('full_name'),
            'title': resume_dict.get('title'),
            'email': resume_dict.get('email'),
            'phone': resume_dict.get('phone'),
            'location': resume_dict.get('location'),
            'skills': resume_dict.get('skills'),
            'experience': resume_dict.get('experience'),
            'education': resume_dict.get('education')
        },
        'top_3_matching_jds': []
    }

    for match in top_matches:
        jd_data = {
            'rank': len(export_data['top_3_matching_jds']) + 1,
            'jd_id': match['jd_id'],
            'title': match['jd_title'],
            'category': match['jd'].get('category'),
            'role': match['jd'].get('role'),
            'match_percentage': match['match_percentage'],
            'skills_matched': f"{match['matched_count']}/{match['total_jd_requirements']}",
            'matched_skills': match.get('matched_skills', []),
            'missing_skills': match.get('missing_skills', []),
            'details': {
                'responsibilities': match['jd'].get('responsibilities'),
                'requirements': match['jd'].get('requirements'),
                'skills': match['jd'].get('skills'),
                'keywords': match['jd'].get('keywords')
            }
        }
        export_data['top_3_matching_jds'].append(jd_data)

    filename = f"{resume_dict.get('full_name', 'resume').replace(' ', '_')}_top_3_matches.json"
    return send_file(
        BytesIO(json.dumps(export_data, indent=2).encode()),
        mimetype='application/json',
        as_attachment=True,
        download_name=filename
    )


@app.route("/profile/<int:resume_id>/export-top-matches-pdf")
def export_top_matches_pdf(resume_id):
    """Export top 3 matching JDs with resume as PDF."""
    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()
        if not resume:
            return "Profile not found", 404

        # Calculate top 3 matching JDs
        jds = conn.execute("SELECT * FROM job_description ORDER BY created_at DESC").fetchall()
        resume_dict = dict(resume)

        matches = []
        for jd in jds:
            jd_dict = dict(jd)
            score = calculate_match_score(resume_dict, jd_dict)
            matches.append({
                'jd': jd_dict,
                'jd_id': jd['id'],
                'jd_title': jd['title'],
                'match_percentage': score['match_percentage'],
                'matched_count': score['matched_count'],
                'total_jd_requirements': score['total_jd_requirements']
            })

        matches.sort(key=lambda x: x['match_percentage'], reverse=True)
        top_matches = matches[:3]

    # Generate PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          topMargin=0.5*inch, bottomMargin=0.5*inch,
                          leftMargin=0.75*inch, rightMargin=0.75*inch)
    story = []
    styles = getSampleStyleSheet()

    # Resume Header
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=22,
        textColor='#1e293b',
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    story.append(Paragraph(f"Resume: {resume_dict.get('full_name', 'Resume')}", title_style))

    if resume_dict.get('title'):
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=11,
            textColor='#64748b',
            spaceAfter=16,
            alignment=TA_CENTER
        )
        story.append(Paragraph(resume_dict.get('title'), subtitle_style))

    # Contact info
    contact_parts = []
    if resume_dict.get('email'):
        contact_parts.append(resume_dict['email'])
    if resume_dict.get('phone'):
        contact_parts.append(resume_dict['phone'])
    if resume_dict.get('location'):
        contact_parts.append(resume_dict['location'])

    if contact_parts:
        contact_style = ParagraphStyle(
            'Contact',
            parent=styles['Normal'],
            fontSize=9,
            textColor='#475569',
            spaceAfter=20,
            alignment=TA_CENTER
        )
        story.append(Paragraph(' | '.join(contact_parts), contact_style))

    story.append(Spacer(1, 0.3*inch))

    # Top 3 Matches
    top3_style = ParagraphStyle(
        'Top3Title',
        parent=styles['Heading2'],
        fontSize=16,
        textColor='#1e293b',
        spaceAfter=12,
        spaceBefore=6,
        fontName='Helvetica-Bold'
    )
    story.append(Paragraph('TOP 3 MATCHING JOB DESCRIPTIONS', top3_style))

    section_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading3'],
        fontSize=11,
        textColor='#1e293b',
        spaceAfter=6,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=8.5,
        spaceAfter=6,
        alignment=TA_LEFT
    )

    for idx, match in enumerate(top_matches, 1):
        jd = match['jd']
        # Match header
        match_header = f"#{idx} - {jd.get('title', 'Job Description')} ({match['match_percentage']}% match)"
        story.append(Paragraph(match_header, section_style))

        # Match stats
        stats_text = f"<b>Skills Matched:</b> {match['matched_count']}/{match['total_jd_requirements']}"
        story.append(Paragraph(stats_text, body_style))

        if jd.get('category'):
            story.append(Paragraph(f"<b>Category:</b> {jd['category']}", body_style))

        story.append(Spacer(1, 0.1*inch))

        # Responsibilities
        if jd.get('responsibilities'):
            story.append(Paragraph('<b>Responsibilities:</b>', body_style))
            resp_text = jd['responsibilities'].replace('\n', ' ')[:300] + "..."
            story.append(Paragraph(resp_text, body_style))

        # Required Skills
        if jd.get('skills'):
            story.append(Paragraph('<b>Required Skills:</b>', body_style))
            skills_list = [s.strip() for s in jd['skills'].split('\n') if s.strip()][:5]
            skills_text = ', '.join(skills_list)
            story.append(Paragraph(skills_text, body_style))

        story.append(Spacer(1, 0.15*inch))

    # Build PDF
    doc.build(story)
    buffer.seek(0)

    filename = f"{resume_dict.get('full_name', 'resume').replace(' ', '_')}_top_3_matches.pdf"
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename
    )


@app.route("/profile/<int:resume_id>/export-rich-pdf")
def export_rich_profile_pdf(resume_id):
    """Export a professional assessment PDF: overview table + per-JD breakdown."""
    from reportlab.platypus import Table, TableStyle, HRFlowable, PageBreak
    from reportlab.lib import colors as rl_colors

    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()
        if not resume:
            return "Profile not found", 404
        jds = conn.execute("SELECT * FROM job_description ORDER BY created_at DESC").fetchall()

    resume_dict = dict(resume)
    matches = []
    for jd in jds:
        jd_dict = dict(jd)
        score = calculate_match_score(resume_dict, jd_dict)
        score['jd'] = jd_dict
        score['jd_title']    = str(jd_dict.get('title') or 'Unknown')
        score['jd_category'] = str(jd_dict.get('category') or '')
        matches.append(score)
    matches.sort(key=lambda x: x['match_percentage'], reverse=True)
    top_matches = matches[:3]

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)

    C_DARK    = rl_colors.HexColor('#1e293b')
    C_INDIGO  = rl_colors.HexColor('#4f46e5')
    C_GREEN   = rl_colors.HexColor('#10b981')
    C_ORANGE  = rl_colors.HexColor('#f59e0b')
    C_DORANG  = rl_colors.HexColor('#d97706')
    C_RED     = rl_colors.HexColor('#ef4444')
    C_MUTED   = rl_colors.HexColor('#64748b')
    C_BGBLUE  = rl_colors.HexColor('#eef2ff')
    C_BGGRN   = rl_colors.HexColor('#d1fae5')
    C_BGRED   = rl_colors.HexColor('#fee2e2')
    C_BGORG   = rl_colors.HexColor('#fef3c7')
    C_BGINDIG = rl_colors.HexColor('#ede9fe')
    C_WHITE   = rl_colors.white
    C_HDR     = rl_colors.HexColor('#1e293b')
    C_LINE    = rl_colors.HexColor('#3b82f6')

    def _scol(p):
        return C_GREEN if p >= 80 else C_INDIGO if p >= 60 else C_ORANGE if p >= 40 else C_RED

    def _slbl(p):
        if p >= 80: return 'Strong Match'
        if p >= 60: return 'Good Match'
        if p >= 40: return 'Partial Match'
        return 'Low Match'

    def _ps(nm, **kw):
        return ParagraphStyle(nm, **kw)

    W = A4[0] - 1.5 * inch
    story = []

    # ── PAGE 1: HEADER BAR ────────────────────────────────────────────────────
    hdr_bar = Table(
        [[Paragraph('CANDIDATE ASSESSMENT REPORT',
                    _ps('RP_HdrBar', fontName='Helvetica-Bold', fontSize=13,
                        textColor=C_WHITE, alignment=TA_CENTER))]],
        colWidths=[W]
    )
    hdr_bar.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, -1), C_HDR),
        ('TOPPADDING',    (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('LEFTPADDING',   (0, 0), (-1, -1), 10),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 10),
    ]))
    story.append(hdr_bar)
    story.append(Spacer(1, 0.18 * inch))

    cname  = str(resume_dict.get('full_name') or 'Candidate')
    ctitle = str(resume_dict.get('title') or '')
    cemail = str(resume_dict.get('email') or '')
    cphone = str(resume_dict.get('phone') or '')
    cloc   = str(resume_dict.get('location') or '')
    contact_parts = [x for x in [cemail, cphone, cloc] if x]

    story.append(Paragraph(cname,
                            _ps('RP_Name', fontName='Helvetica-Bold', fontSize=26,
                                textColor=C_DARK, leading=32, alignment=TA_CENTER)))
    if ctitle:
        story.append(Paragraph(ctitle,
                                _ps('RP_Title', fontName='Helvetica', fontSize=13,
                                    textColor=C_MUTED, leading=18, alignment=TA_CENTER,
                                    spaceAfter=2)))
    if contact_parts:
        story.append(Paragraph(' | '.join(contact_parts),
                                _ps('RP_Contact', fontName='Helvetica', fontSize=9,
                                    textColor=C_MUTED, leading=14, alignment=TA_CENTER,
                                    spaceAfter=6)))
    story.append(Spacer(1, 0.08 * inch))
    story.append(HRFlowable(width=W, thickness=2, color=C_LINE, spaceAfter=14))

    # ── OVERVIEW TABLE ────────────────────────────────────────────────────────
    story.append(Paragraph('TOP 3 MATCHING ROLES — OVERVIEW',
                            _ps('RP_OvSec', fontName='Helvetica-Bold', fontSize=12,
                                textColor=C_INDIGO, spaceAfter=8)))

    ov_hc = _ps('RP_OvHC', fontName='Helvetica-Bold', fontSize=9,
                textColor=C_WHITE, alignment=TA_CENTER)
    ov_hl = _ps('RP_OvHL', fontName='Helvetica-Bold', fontSize=9, textColor=C_WHITE)
    ov_data = [[
        Paragraph('Rank',        ov_hc),
        Paragraph('Job Role',    ov_hl),
        Paragraph('Category',    ov_hl),
        Paragraph('Match Score', ov_hc),
        Paragraph('Matched',     ov_hc),
        Paragraph('Missing',     ov_hc),
    ]]
    for ri, m in enumerate(top_matches, 1):
        p      = m['match_percentage']
        sc     = _scol(p)
        mc     = int(m['matched_count'])
        ms_cnt = int(m['missing_count'])
        ov_data.append([
            Paragraph(f'#{ri}',
                      _ps(f'RP_OvRk{ri}', fontName='Helvetica-Bold',
                          fontSize=10, textColor=C_DARK, alignment=TA_CENTER)),
            Paragraph(str(m['jd_title']),
                      _ps(f'RP_OvT{ri}', fontName='Helvetica', fontSize=9,
                          textColor=C_DARK, leading=12)),
            Paragraph(str(m['jd_category']),
                      _ps(f'RP_OvC{ri}', fontName='Helvetica', fontSize=9,
                          textColor=C_MUTED, leading=12)),
            Paragraph(f'{p}%  ({_slbl(p)})',
                      _ps(f'RP_OvSc{ri}', fontName='Helvetica-Bold',
                          fontSize=9, textColor=sc, alignment=TA_CENTER)),
            Paragraph(str(mc),
                      _ps(f'RP_OvMc{ri}', fontName='Helvetica-Bold',
                          fontSize=9, textColor=C_GREEN, alignment=TA_CENTER)),
            Paragraph(str(ms_cnt),
                      _ps(f'RP_OvMs{ri}', fontName='Helvetica-Bold',
                          fontSize=9, textColor=C_RED, alignment=TA_CENTER)),
        ])

    ov_cols  = [0.08*W, 0.28*W, 0.22*W, 0.24*W, 0.09*W, 0.09*W]
    ov_style = [
        ('BACKGROUND',    (0, 0), (-1, 0),  C_DARK),
        ('TOPPADDING',    (0, 0), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ('LEFTPADDING',   (0, 0), (-1, -1), 7),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 7),
        ('GRID',          (0, 0), (-1, -1), 0.5, rl_colors.HexColor('#e2e8f0')),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
    ]
    for ri in range(1, len(ov_data)):
        ov_style.append(('BACKGROUND', (0, ri), (-1, ri),
                         C_WHITE if ri % 2 == 1 else rl_colors.HexColor('#f8fafc')))
    ov_tbl = Table(ov_data, colWidths=ov_cols)
    ov_tbl.setStyle(TableStyle(ov_style))
    story.append(ov_tbl)
    story.append(Spacer(1, 0.2 * inch))

    # ── CANDIDATE KEY SKILLS ──────────────────────────────────────────────────
    skills_raw = str(resume_dict.get('skills') or '')
    skill_list = [s.strip() for s in skills_raw.split('\n') if s.strip()]

    if skill_list:
        story.append(Paragraph('Candidate Key Skills',
                                _ps('RP_SkSec', fontName='Helvetica-Bold', fontSize=12,
                                    textColor=C_DARK, spaceAfter=8)))
        tag_sty   = _ps('RP_SkTag', fontName='Helvetica', fontSize=8.5,
                        textColor=C_DARK, leading=13, alignment=TA_CENTER)
        empty_tag = Paragraph('', tag_sty)
        tag_rows  = []
        row       = []
        for i, sk in enumerate(skill_list):
            row.append(Paragraph(str(sk), tag_sty))
            if len(row) == 4 or i == len(skill_list) - 1:
                while len(row) < 4:
                    row.append(empty_tag)
                tag_rows.append(row)
                row = []
        if tag_rows:
            sk_tbl = Table(tag_rows, colWidths=[W / 4] * 4)
            sk_tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0, 0), (-1, -1), C_BGBLUE),
                ('TOPPADDING',    (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING',   (0, 0), (-1, -1), 6),
                ('RIGHTPADDING',  (0, 0), (-1, -1), 6),
                ('GRID',          (0, 0), (-1, -1), 0.5, rl_colors.HexColor('#c7d2fe')),
                ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(sk_tbl)

    # ── HELPER: tag grid (4-column) ───────────────────────────────────────────
    def _tag_grid(items, tag_bg, text_col, border_col, pfx):
        safe = list(items) if items else []
        if not safe:
            return None
        t_sty  = _ps(f'RP_TG{pfx}', fontName='Helvetica', fontSize=8.5,
                     textColor=text_col, leading=13, alignment=TA_CENTER)
        empty  = Paragraph('', t_sty)
        rows   = []
        row    = []
        for i, sk in enumerate(safe):
            row.append(Paragraph(str(sk), t_sty))
            if len(row) == 4 or i == len(safe) - 1:
                while len(row) < 4:
                    row.append(empty)
                rows.append(row)
                row = []
        tbl = Table(rows, colWidths=[W / 4] * 4)
        tbl.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (-1, -1), tag_bg),
            ('TOPPADDING',    (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('LEFTPADDING',   (0, 0), (-1, -1), 6),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 6),
            ('GRID',          (0, 0), (-1, -1), 0.5, border_col),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return tbl

    # ── HELPER: skills list (nested single-column Table) ─────────────────────
    def _skill_inner(items, label, label_col, text_col, marker, pfx):
        safe = list(items) if items else []
        rows = [[Paragraph(label,
                           _ps(f'RP_SL{pfx}', fontName='Helvetica-Bold', fontSize=9,
                               textColor=label_col, leading=13, spaceAfter=2))]]
        if safe:
            item_sty = _ps(f'RP_SI{pfx}', fontName='Helvetica', fontSize=8.5,
                           textColor=text_col, leading=13, leftIndent=4)
            for sk in safe[:20]:
                rows.append([Paragraph(f'{marker} {str(sk)}', item_sty)])
            if len(safe) > 20:
                rows.append([Paragraph(
                    f'+ {len(safe) - 20} more',
                    _ps(f'RP_SM{pfx}', fontName='Helvetica', fontSize=8,
                        textColor=C_MUTED, leading=11, leftIndent=4)
                )])
        else:
            rows.append([Paragraph('None',
                                   _ps(f'RP_SN{pfx}', fontName='Helvetica', fontSize=8.5,
                                       textColor=C_MUTED, leading=12, leftIndent=4))])
        inner = Table(rows, colWidths=[W * 0.5 - 28])
        inner.setStyle(TableStyle([
            ('TOPPADDING',    (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ('LEFTPADDING',   (0, 0), (-1, -1), 0),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ]))
        return inner

    # ── PER-JD DETAILED SECTIONS ──────────────────────────────────────────────
    RANKS = ['#1', '#2', '#3']

    for idx, m in enumerate(top_matches):
        p           = m['match_percentage']
        sc          = _scol(p)
        jdt         = str(m['jd_title'])
        sp          = int(m.get('skills_match_percentage', 0))
        ep          = int(m.get('experience_match_percentage', 0))
        exp_note    = str(m.get('experience_note') or '')
        res_yrs     = m.get('resume_years_estimated')
        strong      = list(m.get('strong_areas') or [])
        weak        = list(m.get('weak_areas') or [])
        matched     = list(m.get('matched_skills') or [])
        missing     = list(m.get('missing_skills') or [])
        match_count = int(m['matched_count'])
        total       = int(m['total_jd_requirements'])
        miss_count  = int(m['missing_count'])

        story.append(PageBreak())

        # JD banner: colored rank badge | title | % + label
        banner = Table([[
            Paragraph(RANKS[idx],
                      _ps(f'RP_Rk{idx}', fontName='Helvetica-Bold', fontSize=18,
                          textColor=C_WHITE, alignment=TA_CENTER, leading=22)),
            Paragraph(jdt,
                      _ps(f'RP_JdT{idx}', fontName='Helvetica-Bold', fontSize=13,
                          textColor=C_DARK, leading=18)),
            Paragraph(f'<b>{p}%</b><br/><font size="9">{_slbl(p)}</font>',
                      _ps(f'RP_JdP{idx}', fontName='Helvetica-Bold', fontSize=22,
                          textColor=sc, alignment=TA_RIGHT, leading=26)),
        ]], colWidths=[0.12 * W, 0.60 * W, 0.28 * W])
        banner.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (0, 0),   sc),
            ('BACKGROUND',    (1, 0), (-1, -1), C_WHITE),
            ('TOPPADDING',    (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('LEFTPADDING',   (0, 0), (0, 0),   6),
            ('LEFTPADDING',   (1, 0), (1, 0),   14),
            ('LEFTPADDING',   (2, 0), (2, 0),   8),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 10),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('BOX',           (0, 0), (-1, -1), 1.5, sc),
        ]))
        story.append(banner)
        story.append(Spacer(1, 0.15 * inch))

        # Match Score Breakdown table
        story.append(Paragraph('Match Score Breakdown',
                                _ps(f'RP_MBH{idx}', fontName='Helvetica-Bold', fontSize=11,
                                    textColor=C_DARK, spaceAfter=6)))
        mb_wh  = _ps(f'RP_MBW{idx}',  fontName='Helvetica-Bold', fontSize=9, textColor=C_WHITE)
        mb_whc = _ps(f'RP_MBWC{idx}', fontName='Helvetica-Bold', fontSize=9,
                     textColor=C_WHITE, alignment=TA_CENTER)
        mb_met = _ps(f'RP_MBM{idx}', fontName='Helvetica', fontSize=9, textColor=C_DARK)
        mb_det = _ps(f'RP_MBD{idx}', fontName='Helvetica', fontSize=9, textColor=C_MUTED)
        exp_det = exp_note[:90] if exp_note else 'N/A'

        mb_data = [
            [Paragraph('Metric', mb_wh),  Paragraph('Score', mb_whc),
             Paragraph('Detail', mb_wh)],
            [Paragraph('Overall Match', mb_met),
             Paragraph(f'{p}%', _ps(f'RP_MBOv{idx}', fontName='Helvetica-Bold', fontSize=9,
                       textColor=sc, alignment=TA_CENTER)),
             Paragraph(f'{match_count} of {total} requirements met', mb_det)],
            [Paragraph('Skills Match', mb_met),
             Paragraph(f'{sp}%', _ps(f'RP_MBSk{idx}', fontName='Helvetica-Bold', fontSize=9,
                       textColor=C_DARK, alignment=TA_CENTER)),
             Paragraph(f'{match_count} skills matched', mb_det)],
            [Paragraph('Experience Match', mb_met),
             Paragraph(f'{ep}%', _ps(f'RP_MBEx{idx}', fontName='Helvetica-Bold', fontSize=9,
                       textColor=C_DARK, alignment=TA_CENTER)),
             Paragraph(exp_det, mb_det)],
        ]
        mb_style = [
            ('BACKGROUND',    (0, 0), (-1, 0),  C_DARK),
            ('TOPPADDING',    (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING',   (0, 0), (-1, -1), 8),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 8),
            ('GRID',          (0, 0), (-1, -1), 0.5, rl_colors.HexColor('#e2e8f0')),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ]
        for ri in range(1, len(mb_data)):
            mb_style.append(('BACKGROUND', (0, ri), (-1, ri),
                             C_WHITE if ri % 2 == 1 else rl_colors.HexColor('#f8fafc')))
        mb_tbl = Table(mb_data, colWidths=[0.28 * W, 0.14 * W, 0.58 * W])
        mb_tbl.setStyle(TableStyle(mb_style))
        story.append(mb_tbl)
        story.append(Spacer(1, 0.15 * inch))

        # Matched / Missing skills side by side
        ms = Table([[
            _skill_inner(matched, f'Matched Skills ({match_count})',
                         C_GREEN, C_GREEN, '•', f'{idx}a'),
            _skill_inner(missing, f'Missing / Gap Skills ({miss_count})',
                         C_RED, C_RED, '×', f'{idx}b'),
        ]], colWidths=[W * 0.5, W * 0.5])
        ms.setStyle(TableStyle([
            ('BACKGROUND',    (0, 0), (0, 0), C_BGGRN),
            ('BACKGROUND',    (1, 0), (1, 0), C_BGRED),
            ('TOPPADDING',    (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING',   (0, 0), (-1, -1), 14),
            ('RIGHTPADDING',  (0, 0), (-1, -1), 14),
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ('LINEAFTER',     (0, 0), (0, 0),   1, rl_colors.HexColor('#d1d5db')),
        ]))
        story.append(ms)
        story.append(Spacer(1, 0.14 * inch))

        # Strong Areas
        if strong:
            story.append(Paragraph('Strong Areas',
                                    _ps(f'RP_StrH{idx}', fontName='Helvetica-Bold', fontSize=11,
                                        textColor=C_INDIGO, spaceAfter=6)))
            tbl = _tag_grid(strong, C_BGINDIG, C_INDIGO,
                            rl_colors.HexColor('#ddd6fe'), f'{idx}str')
            if tbl:
                story.append(tbl)
            story.append(Spacer(1, 0.12 * inch))

        # Skill Gaps
        if weak:
            story.append(Paragraph('Skill Gaps to Address',
                                    _ps(f'RP_GapH{idx}', fontName='Helvetica-Bold', fontSize=11,
                                        textColor=C_ORANGE, spaceAfter=6)))
            tbl = _tag_grid(weak, C_BGORG, C_DORANG,
                            rl_colors.HexColor('#fcd34d'), f'{idx}gap')
            if tbl:
                story.append(tbl)
            story.append(Spacer(1, 0.12 * inch))

        # Experience Assessment
        story.append(Paragraph('Experience Assessment',
                                _ps(f'RP_ExpH{idx}', fontName='Helvetica-Bold', fontSize=11,
                                    textColor=C_DARK, spaceAfter=6)))
        if exp_note:
            story.append(Paragraph(exp_note,
                                    _ps(f'RP_ExpN{idx}', fontName='Helvetica', fontSize=9,
                                        textColor=C_MUTED, leading=14, spaceAfter=4)))
        if res_yrs is not None:
            yr_row = Table([[
                Paragraph('Resume Shows',
                           _ps(f'RP_YrL{idx}', fontName='Helvetica-Bold',
                               fontSize=9, textColor=C_DARK)),
                Paragraph(f'~{res_yrs} year(s)',
                           _ps(f'RP_YrV{idx}', fontName='Helvetica-Bold',
                               fontSize=9, textColor=C_ORANGE)),
            ]], colWidths=[0.25 * W, 0.35 * W])
            yr_row.setStyle(TableStyle([
                ('TOPPADDING',    (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(yr_row)
        story.append(Spacer(1, 0.12 * inch))

        # Recommendations
        story.append(Paragraph('Recommendations',
                                _ps(f'RP_RecH{idx}', fontName='Helvetica-Bold', fontSize=11,
                                    textColor=C_DARK, spaceAfter=6)))
        if p >= 80:
            rec1 = 'Strong match — candidate meets most requirements.'
        elif p >= 60:
            rec1 = 'Good match — candidate meets key requirements with some gaps.'
        elif p >= 40:
            rec1 = 'Partial match — may require upskilling in the gap areas listed above.'
        else:
            rec1 = 'Low match — significant gaps to address before applying.'
        story.append(Paragraph(f'• {rec1}',
                                _ps(f'RP_Rec1{idx}', fontName='Helvetica', fontSize=9,
                                    textColor=C_DARK, leading=14, spaceAfter=4)))
        if weak:
            top_gaps = ', '.join(str(g) for g in weak[:4])
            if len(weak) > 4:
                top_gaps += '...'
            story.append(Paragraph(f'• Key areas to strengthen: {top_gaps}',
                                    _ps(f'RP_Rec2{idx}', fontName='Helvetica', fontSize=9,
                                        textColor=C_DARK, leading=14)))
        story.append(Spacer(1, 0.2 * inch))

    doc.build(story)
    buffer.seek(0)
    safe = (resume_dict.get('full_name') or 'candidate').replace(' ', '_')
    return send_file(buffer, mimetype='application/pdf',
                     as_attachment=True, download_name=f'{safe}_assessment_report.pdf')


@app.route("/profile/slug/<slug>")
def public_profile(slug):
    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE slug = %s", (slug,)).fetchone()
    if not resume:
        return "Profile not found", 404
    return render_template("profile.html", resume=resume)


def generate_resume_pdf(resume):
    """Generate a PDF from resume data using reportlab."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          topMargin=0.5*inch, bottomMargin=0.5*inch,
                          leftMargin=0.75*inch, rightMargin=0.75*inch)
    story = []
    styles = getSampleStyleSheet()

    # Header with name and title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#1e293b',
        spaceAfter=2,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    story.append(Paragraph(resume.get('full_name', 'Resume'), title_style))

    if resume.get('title'):
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor='#64748b',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        story.append(Paragraph(resume.get('title'), subtitle_style))

    # Contact Info
    contact_parts = []
    if resume.get('email'):
        contact_parts.append(resume['email'])
    if resume.get('phone'):
        contact_parts.append(resume['phone'])
    if resume.get('location'):
        contact_parts.append(resume['location'])

    if contact_parts:
        contact_style = ParagraphStyle(
            'Contact',
            parent=styles['Normal'],
            fontSize=9,
            textColor='#475569',
            spaceAfter=16,
            alignment=TA_CENTER
        )
        story.append(Paragraph(' | '.join(contact_parts), contact_style))

    # Section styling
    section_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='#1e293b',
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold',
        borderColor='#e2e8f0',
        borderWidth=0,
        borderPadding=0,
    )

    # Summary
    if resume.get('summary'):
        story.append(Paragraph('PROFESSIONAL SUMMARY', section_style))
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontSize=9,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        story.append(Paragraph(resume['summary'], body_style))

    # Skills
    if resume.get('skills'):
        story.append(Paragraph('TECHNICAL SKILLS', section_style))
        skills_list = [s.strip() for s in resume['skills'].split('\n') if s.strip()]
        skills_text = ' • '.join(skills_list[:20])  # Limit to 20 skills
        story.append(Paragraph(skills_text, ParagraphStyle(
            'Skills', parent=styles['Normal'], fontSize=9, spaceAfter=12
        )))

    # Experience
    if resume.get('experience'):
        story.append(Paragraph('WORK EXPERIENCE', section_style))
        exp_text = resume['experience'].replace('\n', '<br/>')
        story.append(Paragraph(exp_text, ParagraphStyle(
            'Experience', parent=styles['Normal'], fontSize=9, spaceAfter=12
        )))

    # Education
    if resume.get('education'):
        story.append(Paragraph('EDUCATION', section_style))
        edu_text = resume['education'].replace('\n', '<br/>')
        story.append(Paragraph(edu_text, ParagraphStyle(
            'Education', parent=styles['Normal'], fontSize=9, spaceAfter=12
        )))

    # Projects
    if resume.get('projects'):
        story.append(Paragraph('PROJECTS', section_style))
        proj_text = resume['projects'].replace('\n', '<br/>')
        story.append(Paragraph(proj_text, ParagraphStyle(
            'Projects', parent=styles['Normal'], fontSize=9, spaceAfter=12
        )))

    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route("/profile/<int:resume_id>/download-pdf")
def download_resume_pdf(resume_id):
    """Download resume as PDF."""
    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()

    if not resume:
        return "Profile not found", 404

    resume_dict = dict(resume)
    pdf_buffer = generate_resume_pdf(resume_dict)

    filename = f"{resume_dict.get('full_name', 'Resume').replace(' ', '_')}.pdf"
    return send_file(
        pdf_buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename
    )


@app.route("/profile/<int:resume_id>/extract")
def extract_resume_data(resume_id):
    """Extract resume data as JSON."""
    with db_conn() as conn:
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()

    if not resume:
        return jsonify({"error": "Profile not found"}), 404

    resume_dict = dict(resume)
    # Remove sensitive/internal fields
    resume_dict.pop('id', None)
    resume_dict.pop('created_at', None)
    resume_dict.pop('updated_at', None)

    filename = f"{resume_dict.get('full_name', 'resume').replace(' ', '_')}_extracted.json"

    return send_file(
        BytesIO(json.dumps(resume_dict, indent=2).encode()),
        mimetype='application/json',
        as_attachment=True,
        download_name=filename
    )


@app.route("/profile/<int:resume_id>/delete", methods=["POST"])
def delete_resume(resume_id):
    with db_conn() as conn:
        row = conn.execute("SELECT resume_file FROM resume WHERE id = %s", (resume_id,)).fetchone()
        if not row:
            return "Profile not found", 404
        if row["resume_file"]:
            (UPLOAD_FOLDER / row["resume_file"]).unlink(missing_ok=True)
        conn.execute("DELETE FROM resume WHERE id = %s", (resume_id,))
    flash("Profile deleted.", "success")
    return redirect(url_for("profile_list"))


@app.route("/uploads/<path:filename>")
def uploaded_file(filename):
    # Only serve allowed file types
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        abort(403)
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)


@app.route("/api/parse-resume", methods=["POST"])
def parse_resume_api():
    uploaded = request.files.get("resume_file")
    if not uploaded or not uploaded.filename:
        return jsonify({"success": False, "message": "Please choose a resume file."}), 400
    if not allowed_file(uploaded.filename):
        return jsonify({"success": False, "message": "Please upload a PDF or DOCX file only."}), 400

    uploaded.seek(0, 2)
    file_size = uploaded.tell()
    uploaded.seek(0)
    if file_size > 25 * 1024 * 1024:
        return jsonify({
            "success": False,
            "message": f"File size ({file_size / 1024 / 1024:.1f} MB) exceeds the 25 MB limit.",
        }), 400

    try:
        name = secure_filename(uploaded.filename)
        ext = name.rsplit(".", 1)[1].lower()
        temp_path = UPLOAD_FOLDER / f"tmp-parse.{ext}"
        uploaded.save(str(temp_path))
        # parse_mode: "llm" = Resume Intelligence (AI), anything else = Quick Parse
        parse_mode  = request.form.get("parse_mode", "llm")
        parser_used = "standard"
        logger.info(f"API parse-resume: parse_mode={parse_mode}, ext={ext}")
        try:
            if ext == "pdf" and parse_mode == "llm":
                try:
                    parsed, parser_used = parse_resume_with_llm_text(temp_path)
                except Exception as e:
                    logger.warning("Resume Intelligence failed (%s); using Quick Parse.", e, exc_info=True)
                    parsed = _parse_pdf_quick(temp_path)
                    parser_used = "text (fallback)"
            elif ext == "pdf":
                parsed = _parse_pdf_quick(temp_path)
                parser_used = "text"
            else:
                text = extract_resume_text(temp_path, ext)
                parsed = parse_resume_text(text)
                parser_used = "text"
        finally:
            temp_path.unlink(missing_ok=True)
        return jsonify({"success": True, "message": "Fields extracted from resume.", "data": parsed, "parser_used": parser_used})
    except Exception as e:
        logger.error(f"Error parsing resume: {e}", exc_info=True)
        return jsonify({"success": False, "message": f"Could not extract text: {e}"}), 422


@app.route("/api/resume")
def resume_api():
    with db_conn() as conn:
        resumes = conn.execute(
            "SELECT * FROM resume ORDER BY updated_at DESC, id DESC"
        ).fetchall()
    return jsonify([dict(row) for row in resumes])


# ── New: Role Definitions ─────────────────────────────────────────────────────

REGULATORY_ROLES = [
    "Design Control Consultant",
    "IFU Technical Writer",
    "IFU Team Lead",
    "Product Registration Specialist",
    "Labeling Specialist",
]

VALIDATION_ROLES = [
    "CSV Analyst", "CSV Lead", "Validation Engineer", "Validation Lead",
    "CQV Engineer", "CQV Lead", "Automation Engineer", "Automation Lead",
    "Tosca Engineer", "Tosca Lead", "Test Engineer", "Test Lead",
]

IT_ROLES = [
    "Software Developer", "Full Stack Developer", "Frontend Developer",
    "Backend Developer", "DevOps Engineer", "Cloud Engineer", "Data Analyst",
    "Data Engineer", "Business Analyst", "QA Engineer", "Automation Tester",
    "Project Manager", "Scrum Master", "UI/UX Designer", "Solution Architect",
    "Cybersecurity Engineer", "Database Administrator", "AI/ML Engineer",
    "SAP Consultant", "Salesforce Developer",
]

ALL_JD_ROLES = {
    "Regulatory Affairs": REGULATORY_ROLES,
    "Validation Roles": VALIDATION_ROLES,
    "IT Roles": IT_ROLES,
}

# ── New: Raw Upload Folder ────────────────────────────────────────────────────

RAW_UPLOAD_FOLDER = UPLOAD_FOLDER / "raw"
RAW_UPLOAD_FOLDER.mkdir(exist_ok=True)

_RAW_META_PATH = RAW_UPLOAD_FOLDER / "_meta.json"
_RAW_UPLOAD_ALLOWED = {"pdf", "docx", "doc"}  # allowed inside ZIP / direct


def _load_raw_meta():
    if _RAW_META_PATH.exists():
        try:
            return json.loads(_RAW_META_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _save_raw_meta(meta):
    try:
        _RAW_META_PATH.write_text(
            json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass


# ── New: Dashboard Route ──────────────────────────────────────────────────────

@app.route("/dashboard")
def dashboard():
    from datetime import datetime as _dt
    with db_conn() as conn:
        total_resumes = conn.execute("SELECT COUNT(*) AS count FROM resume").fetchone()["count"]

        today_resumes = conn.execute(
            "SELECT id, full_name, title, created_at FROM resume"
            " WHERE created_at::date = CURRENT_DATE ORDER BY created_at DESC"
        ).fetchall()

        role_counts = conn.execute(
            """
            SELECT COALESCE(NULLIF(title, ''), 'No Title') AS role, COUNT(*) AS count
            FROM resume
            GROUP BY COALESCE(NULLIF(title, ''), 'No Title')
            ORDER BY count DESC
            LIMIT 20
            """
        ).fetchall()

        recent_activity = conn.execute(
            "SELECT id, full_name, title, created_at, updated_at"
            " FROM resume ORDER BY updated_at DESC LIMIT 10"
        ).fetchall()

        total_roles = conn.execute(
            "SELECT COUNT(DISTINCT NULLIF(title, '')) AS count FROM resume"
        ).fetchone()["count"] or 0

    return render_template(
        "dashboard.html",
        total_resumes=total_resumes,
        today_resumes=list(today_resumes),
        role_counts=list(role_counts),
        recent_activity=list(recent_activity),
        total_roles=total_roles,
        all_jd_roles=ALL_JD_ROLES,
    )


# ── New: Upload Files Route ───────────────────────────────────────────────────

@app.route("/upload-files", methods=["GET", "POST"])
def upload_files():
    import time as _time
    from datetime import datetime as _dt
    if request.method == "POST":
        files = request.files.getlist("files")
        results = []
        meta = _load_raw_meta()

        for f in files:
            if not f or not f.filename:
                continue
            name = secure_filename(f.filename)
            ext = name.rsplit(".", 1)[1].lower() if "." in name else ""

            # ── ZIP: extract each resume inside it ───────────────────────────
            if ext == "zip":
                f.seek(0)
                tmp_zip = RAW_UPLOAD_FOLDER / f"_tmp_{int(_time.time())}.zip"
                f.save(str(tmp_zip))
                try:
                    with zipfile.ZipFile(tmp_zip, "r") as zf:
                        for member in zf.namelist():
                            member_name = Path(member).name
                            if not member_name or member_name.startswith("."):
                                continue
                            m_ext = member_name.rsplit(".", 1)[-1].lower() if "." in member_name else ""
                            if m_ext not in _RAW_UPLOAD_ALLOWED:
                                continue
                            m_stem = Path(member_name).stem
                            save_name = f"{slugify(m_stem)}-{int(_time.time())}.{m_ext}"
                            dest = RAW_UPLOAD_FOLDER / save_name
                            with zf.open(member) as src, open(str(dest), "wb") as dst:
                                dst.write(src.read())
                            now_str = _dt.now().strftime("%Y-%m-%d %H:%M")
                            meta[save_name] = {
                                "original": member_name, "source": "zip",
                                "zip_name": f.filename, "uploaded_at": now_str,
                            }
                            results.append({
                                "name": save_name, "original": member_name,
                                "size": dest.stat().st_size, "status": "success",
                                "source": "zip", "zip_name": f.filename,
                            })
                except Exception as e:
                    results.append({"name": f.filename, "status": "error",
                                    "message": f"ZIP error: {e}"})
                finally:
                    tmp_zip.unlink(missing_ok=True)
                continue

            # ── Existing PDF / DOCX logic (unchanged) ────────────────────────
            if not allowed_file(f.filename):
                results.append({"name": f.filename, "status": "error",
                                 "message": "Unsupported format (PDF/DOCX/ZIP only)"})
                continue
            stem = Path(name).stem
            save_name = f"{slugify(stem)}-{int(_time.time())}.{ext}"
            dest = RAW_UPLOAD_FOLDER / save_name
            f.seek(0)
            f.save(str(dest))
            now_str = _dt.now().strftime("%Y-%m-%d %H:%M")
            meta[save_name] = {
                "original": f.filename, "source": "direct", "uploaded_at": now_str,
            }
            results.append({
                "name": save_name, "original": f.filename,
                "size": dest.stat().st_size, "status": "success", "source": "direct",
            })

        _save_raw_meta(meta)
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify({"results": results})
        ok = sum(1 for r in results if r["status"] == "success")
        if ok:
            flash(f"Uploaded {ok} file(s) successfully.", "success")
        return redirect(url_for("upload_files"))

    # GET
    from datetime import datetime as _dt
    meta = _load_raw_meta()
    raw_files = []
    if RAW_UPLOAD_FOLDER.exists():
        for p in sorted(RAW_UPLOAD_FOLDER.iterdir(), key=lambda x: x.stat().st_mtime, reverse=True):
            if p.is_file() and p.suffix.lower().lstrip(".") in ALLOWED_EXTENSIONS:
                st = p.stat()
                fm = meta.get(p.name, {})
                raw_files.append({
                    "name": p.name,
                    "size": st.st_size,
                    "uploaded_at": fm.get("uploaded_at") or _dt.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M"),
                    "ext": p.suffix.lower().lstrip("."),
                    "source": fm.get("source", "direct"),
                    "original": fm.get("original", p.name),
                    "zip_name": fm.get("zip_name"),
                })

    # Pass JD list for the bulk-compare panel
    jds = []
    try:
        with db_conn() as conn:
            ensure_jd_table(conn)
            seed_jds(conn)
            jds = conn.execute(
                "SELECT id, title, category FROM job_description ORDER BY category, title"
            ).fetchall()
    except Exception:
        pass

    return render_template("upload_files.html", raw_files=raw_files, jds=list(jds))


@app.route("/uploads/raw/<path:filename>")
def raw_uploaded_file(filename):
    safe = secure_filename(filename)
    ext = safe.rsplit(".", 1)[-1].lower() if "." in safe else ""
    if ext not in ALLOWED_EXTENSIONS:
        abort(403)
    return send_from_directory(str(RAW_UPLOAD_FOLDER), safe)


@app.route("/api/delete-upload/<path:filename>", methods=["POST"])
def delete_raw_upload(filename):
    safe = secure_filename(filename)
    p = RAW_UPLOAD_FOLDER / safe
    if p.exists() and p.is_file():
        p.unlink()
        meta = _load_raw_meta()
        meta.pop(safe, None)
        _save_raw_meta(meta)
        return jsonify({"success": True})
    return jsonify({"success": False, "message": "File not found"}), 404


@app.route("/api/delete-all-uploads", methods=["POST"])
def delete_all_uploads():
    """Delete every raw uploaded file and clear metadata."""
    deleted = 0
    if RAW_UPLOAD_FOLDER.exists():
        for p in list(RAW_UPLOAD_FOLDER.iterdir()):
            if p.is_file() and p.name != '_meta.json':
                p.unlink()
                deleted += 1
    _save_raw_meta({})
    return jsonify({"success": True, "deleted": deleted})


@app.route("/api/export-compare-pdf", methods=["POST"])
def export_compare_pdf():
    """Generate a PDF from bulk-compare results for download."""
    from reportlab.platypus import Table, TableStyle, HRFlowable
    from reportlab.lib import colors as rl_colors

    data = request.get_json(silent=True) or {}
    results = data.get('results', [])
    jd_title = data.get('jd_title', 'Job Description')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            topMargin=0.5*inch, bottomMargin=0.5*inch,
                            leftMargin=0.65*inch, rightMargin=0.65*inch)

    COL_DARK   = rl_colors.HexColor('#1e293b')
    COL_ACCENT = rl_colors.HexColor('#6366f1')
    COL_GREEN  = rl_colors.HexColor('#10b981')
    COL_RED    = rl_colors.HexColor('#ef4444')
    COL_ORANGE = rl_colors.HexColor('#f59e0b')
    COL_MUTED  = rl_colors.HexColor('#64748b')
    COL_WHITE  = rl_colors.white
    COL_HDR    = rl_colors.HexColor('#312e81')

    def sc(pct):
        if pct >= 80: return COL_GREEN
        if pct >= 60: return COL_ACCENT
        if pct >= 40: return COL_ORANGE
        return COL_RED

    def lvl(pct):
        if pct >= 80: return 'Strong Match'
        if pct >= 60: return 'Good Match'
        if pct >= 40: return 'Partial Match'
        return 'Low Match'

    story = []
    W = A4[0] - 1.3*inch

    # Header
    hdr_tbl = Table([[
        Paragraph(f'Compare Report', ParagraphStyle('H', fontName='Helvetica-Bold',
                  fontSize=20, textColor=COL_WHITE, leading=24)),
        Paragraph(f'vs {jd_title}', ParagraphStyle('H2', fontName='Helvetica',
                  fontSize=10, textColor=rl_colors.HexColor('#c7d2fe'), leading=14)),
    ]], colWidths=[W*0.55, W*0.45])
    hdr_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), COL_HDR),
        ('TOPPADDING', (0,0), (-1,-1), 18),
        ('BOTTOMPADDING', (0,0), (-1,-1), 18),
        ('LEFTPADDING', (0,0), (-1,-1), 20),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(hdr_tbl)
    story.append(Spacer(1, 0.2*inch))

    if not results:
        story.append(Paragraph('No results to display.', ParagraphStyle('N', fontName='Helvetica', fontSize=11)))
    else:
        # Summary table
        hdr_row = [
            Paragraph('#', ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=COL_WHITE, alignment=TA_CENTER)),
            Paragraph('Candidate', ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=COL_WHITE)),
            Paragraph('Title', ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=COL_WHITE)),
            Paragraph('Score', ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=COL_WHITE, alignment=TA_CENTER)),
            Paragraph('Matched', ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=COL_WHITE, alignment=TA_CENTER)),
            Paragraph('Level', ParagraphStyle('TH', fontName='Helvetica-Bold', fontSize=8.5,
                      textColor=COL_WHITE, alignment=TA_CENTER)),
        ]
        tbl_data = [hdr_row]
        row_colors = []
        for i, r in enumerate(results):
            pct = r.get('match_percentage', 0)
            color = sc(pct)
            bg = rl_colors.HexColor('#f8fafc') if i % 2 == 0 else COL_WHITE
            row_colors.append(bg)
            matched_str = ', '.join(r.get('matched_skills', [])[:5])
            if len(r.get('matched_skills', [])) > 5:
                matched_str += f" +{len(r['matched_skills'])-5}"
            tbl_data.append([
                Paragraph(str(i+1), ParagraphStyle('TD', fontName='Helvetica-Bold', fontSize=9,
                          textColor=color, alignment=TA_CENTER)),
                Paragraph(r.get('candidate_name', r.get('file', '—')),
                          ParagraphStyle('TD2', fontName='Helvetica-Bold', fontSize=8.5,
                                         textColor=COL_DARK)),
                Paragraph(r.get('title', '—'),
                          ParagraphStyle('TD3', fontName='Helvetica', fontSize=8,
                                         textColor=COL_MUTED)),
                Paragraph(f'{pct}%', ParagraphStyle('TD4', fontName='Helvetica-Bold', fontSize=11,
                          textColor=color, alignment=TA_CENTER)),
                Paragraph(f"{r.get('matched_count',0)}/{r.get('total_jd_requirements',0)}",
                          ParagraphStyle('TD5', fontName='Helvetica', fontSize=8.5,
                                         textColor=COL_MUTED, alignment=TA_CENTER)),
                Paragraph(lvl(pct), ParagraphStyle('TD6', fontName='Helvetica', fontSize=8,
                          textColor=color, alignment=TA_CENTER)),
            ])

        col_w = [0.06*W, 0.28*W, 0.22*W, 0.12*W, 0.16*W, 0.16*W]
        main_tbl = Table(tbl_data, colWidths=col_w, repeatRows=1)
        ts = [
            ('BACKGROUND',    (0,0), (-1,0), COL_DARK),
            ('TOPPADDING',    (0,0), (-1,-1), 7),
            ('BOTTOMPADDING', (0,0), (-1,-1), 7),
            ('LEFTPADDING',   (0,0), (-1,-1), 8),
            ('RIGHTPADDING',  (0,0), (-1,-1), 8),
            ('GRID',          (0,0), (-1,-1), 0.4, rl_colors.HexColor('#e2e8f0')),
            ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
        ]
        for i, bg in enumerate(row_colors):
            ts.append(('BACKGROUND', (0, i+1), (-1, i+1), bg))
        main_tbl.setStyle(TableStyle(ts))
        story.append(main_tbl)
        story.append(Spacer(1, 0.3*inch))

        # Per-candidate matched/missing detail
        detail_sec = ParagraphStyle('DS', fontName='Helvetica-Bold', fontSize=11,
                                    textColor=COL_DARK, spaceBefore=8, spaceAfter=6)
        story.append(Paragraph('Skills Detail Per Candidate', detail_sec))
        story.append(HRFlowable(width=W, thickness=0.8,
                                color=rl_colors.HexColor('#e2e8f0'), spaceAfter=8))

        for r in results:
            if r.get('error'):
                continue
            pct = r.get('match_percentage', 0)
            color = sc(pct)
            name = r.get('candidate_name', r.get('file', '?'))
            matched = r.get('matched_skills', [])
            missing = r.get('missing_skills', [])

            cand_hdr = Table([[
                Paragraph(name, ParagraphStyle('CN', fontName='Helvetica-Bold',
                          fontSize=10, textColor=COL_WHITE, leading=14)),
                Paragraph(f'{pct}%  {lvl(pct)}',
                          ParagraphStyle('CPct', fontName='Helvetica-Bold', fontSize=10,
                                         textColor=COL_WHITE, alignment=TA_CENTER, leading=14)),
            ]], colWidths=[W*0.65, W*0.35])
            cand_hdr.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), color),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LEFTPADDING', (0,0), (-1,-1), 12),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ]))
            story.append(cand_hdr)

            m_text = ', '.join(matched) if matched else 'None'
            g_text = ', '.join(missing) if missing else 'None'
            detail_row = Table([[
                [Paragraph('Matched Skills', ParagraphStyle('ML', fontName='Helvetica-Bold',
                            fontSize=8, textColor=COL_GREEN, leading=12, spaceAfter=3)),
                 Paragraph(m_text, ParagraphStyle('MV', fontName='Helvetica', fontSize=7.5,
                            textColor=COL_DARK, leading=11))],
                [Paragraph('Gap / Missing Skills', ParagraphStyle('GL', fontName='Helvetica-Bold',
                            fontSize=8, textColor=COL_RED, leading=12, spaceAfter=3)),
                 Paragraph(g_text, ParagraphStyle('GV', fontName='Helvetica', fontSize=7.5,
                            textColor=COL_DARK, leading=11))],
            ]], colWidths=[W*0.5, W*0.5])
            detail_row.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (0,0), rl_colors.HexColor('#d1fae5')),
                ('BACKGROUND', (1,0), (1,0), rl_colors.HexColor('#fee2e2')),
                ('TOPPADDING', (0,0), (-1,-1), 8),
                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                ('LEFTPADDING', (0,0), (-1,-1), 10),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
                ('LINEAFTER', (0,0), (0,0), 0.5, rl_colors.HexColor('#d1d5db')),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ]))
            story.append(detail_row)
            story.append(Spacer(1, 0.12*inch))

    doc.build(story)
    buffer.seek(0)
    return send_file(buffer, mimetype='application/pdf',
                     as_attachment=True, download_name='compare_report.pdf')


@app.route("/api/bulk-compare", methods=["POST"])
def bulk_compare():
    data = request.get_json(silent=True) or {}
    filenames = data.get("files", [])
    jd_id = data.get("jd_id")

    if not filenames:
        return jsonify({"success": False, "message": "No files selected."}), 400
    if not jd_id:
        return jsonify({"success": False, "message": "No JD selected."}), 400

    try:
        with db_conn() as conn:
            ensure_jd_table(conn)
            jd = conn.execute(
                "SELECT * FROM job_description WHERE id = %s", (int(jd_id),)
            ).fetchone()
        if not jd:
            return jsonify({"success": False, "message": "JD not found."}), 404
        jd_dict = dict(jd)
    except Exception as e:
        return jsonify({"success": False, "message": f"DB error: {e}"}), 500

    results = []
    for filename in filenames:
        safe = secure_filename(filename)
        path = RAW_UPLOAD_FOLDER / safe
        if not path.exists():
            results.append({"file": filename, "candidate_name": safe,
                             "error": "File not found", "match_percentage": 0,
                             "match_level": "Error", "match_level_color": "gray"})
            continue
        ext = safe.rsplit(".", 1)[-1].lower() if "." in safe else ""
        if ext not in ALLOWED_EXTENSIONS:
            results.append({"file": filename, "candidate_name": safe,
                             "error": "Unsupported format", "match_percentage": 0,
                             "match_level": "Error", "match_level_color": "gray"})
            continue
        try:
            if ext == "pdf":
                parsed = _parse_pdf_quick(path)
            else:
                text = extract_resume_text(path, ext)
                parsed = parse_resume_text(text)
            score = calculate_match_score(parsed, jd_dict)
            results.append({
                "file": filename,
                "candidate_name": parsed.get("full_name") or Path(filename).stem,
                "title": parsed.get("title", ""),
                "email": parsed.get("email", ""),
                "match_percentage": score["match_percentage"],
                "match_level": score["match_level"],
                "match_level_color": score["match_level_color"],
                "matched_count": score["matched_count"],
                "missing_count": score["missing_count"],
                "total_jd_requirements": score["total_jd_requirements"],
                "matched_skills": score["matched_skills"][:10],
                "missing_skills": score["missing_skills"][:10],
                "strong_areas": score["strong_areas"][:6],
                "weak_areas": score["weak_areas"][:6],
                "experience_note": score["experience_note"],
            })
        except Exception as e:
            results.append({
                "file": filename,
                "candidate_name": Path(filename).stem,
                "error": str(e),
                "match_percentage": 0,
                "match_level": "Error",
                "match_level_color": "gray",
                "matched_count": 0,
                "missing_count": 0,
                "total_jd_requirements": 0,
                "matched_skills": [],
                "missing_skills": [],
            })

    results.sort(key=lambda x: x.get("match_percentage", 0), reverse=True)
    return jsonify({
        "success": True,
        "jd_title": jd_dict["title"],
        "jd_category": jd_dict.get("category", ""),
        "total": len(results),
        "results": results,
    })


# ── New: Groups Route ─────────────────────────────────────────────────────────

@app.route("/groups")
@app.route("/groups/<path:role>")
def groups(role=None):
    profiles = []
    if role:
        with db_conn() as conn:
            profiles = conn.execute(
                "SELECT id, full_name, title, email, phone, location, updated_at"
                " FROM resume WHERE title ILIKE %s ORDER BY updated_at DESC",
                (f"%{role}%",),
            ).fetchall()
    return render_template(
        "groups.html",
        all_roles=ALL_JD_ROLES,
        selected_role=role,
        profiles=list(profiles),
    )


# ── New: JD Management & Resume Matching ─────────────────────────────────────

JD_UPLOAD_FOLDER = UPLOAD_FOLDER / "jd"
JD_UPLOAD_FOLDER.mkdir(exist_ok=True)


def ensure_jd_table(conn):
    conn.execute("""
        CREATE TABLE IF NOT EXISTS job_description (
            id               SERIAL PRIMARY KEY,
            title            VARCHAR(200) NOT NULL,
            role             VARCHAR(200) DEFAULT '',
            category         VARCHAR(100) DEFAULT 'General',
            responsibilities TEXT DEFAULT '',
            requirements     TEXT DEFAULT '',
            skills           TEXT DEFAULT '',
            keywords         TEXT DEFAULT '',
            jd_file          VARCHAR(500),
            created_at       TIMESTAMP DEFAULT NOW(),
            updated_at       TIMESTAMP DEFAULT NOW()
        )
    """)


PREDEFINED_JDS = [
    # ── From IFU JD PDF ──────────────────────────────────────────────────────
    {
        "title": "IFU Technical Writer",
        "role": "IFU Technical Writer",
        "category": "Regulatory Affairs",
        "responsibilities": (
            "Create, revise, and publish high-quality Instructions for Use (IFU)\n"
            "Ensure documentation complies with QSR, ISO, and internal QMS standards\n"
            "Coordinate translation activities with external agencies and internal stakeholders\n"
            "Maintain adherence to company style guides, templates, and QMS\n"
            "Collaborate with Product Development, Regulatory Affairs, Quality Assurance\n"
            "Support CAPA and complaint investigations by providing documentation expertise\n"
            "Contribute to continuous improvement of documentation standards and templates"
        ),
        "requirements": (
            "Bachelor's or Master's degree in Technical Communication, English, or Life Sciences\n"
            "2-3 years of technical writing experience within life sciences, medical device, or IVD\n"
            "Excellent English speaking, writing, and editing skills\n"
            "Strong understanding of regulated content development including labeling requirements\n"
            "Hands-on experience with CMS/CCMS, ideally DITA/XML structured authoring\n"
            "Experience working with translations and simplified English\n"
            "Experience with EU IVDR documentation requirements"
        ),
        "skills": (
            "Technical Writing\nIFU\nQSR\nISO\nQMS\nCMS\nCCMS\nDITA\nXML\n"
            "EU IVDR\nEU MDR\nTranslation\nRegulatory Compliance\nLabeling\nDocumentation\n"
            "Style Guide\nCAPA\nRisk Communication\nMedical Device\nIVD\n"
            "Simplified English\nStructured Authoring\nLife Sciences"
        ),
        "keywords": "IFU, technical writer, DITA, XML, IVDR, medical device, regulatory, QMS, labeling",
    },
    # ── From CSV JD PDF – Lead ────────────────────────────────────────────────
    {
        "title": "CSV Lead",
        "role": "CSV Lead",
        "category": "Validation Roles",
        "responsibilities": (
            "Lead and perform end-to-end validation activities for complex projects\n"
            "Author end-to-end validation deliverables\n"
            "Support automation, cloud qualification, and new IT projects\n"
            "Provide compliance consulting for GxP IT systems\n"
            "Perform CAPA and RCA using 5 x Why approach\n"
            "Manage test execution and reporting\n"
            "High level of stakeholder management across multiple concurrent tasks\n"
            "Participate in audits and inspections"
        ),
        "requirements": (
            "5+ years of experience in CSV or Computer System Validation\n"
            "Excellence in CSV and Equipment Qualification\n"
            "Hands-on experience with Agile, Waterfall, and hybrid models\n"
            "Expert in HP ALM, Azure, EDMS\n"
            "Strong IT compliance and GxP knowledge\n"
            "21 CFR Part 11 compliance experience"
        ),
        "skills": (
            "CSV\nComputer System Validation\nEquipment Qualification\nHP ALM\nAzure\nEDMS\n"
            "Agile\nWaterfall\nCAPA\nRCA\nGxP\nGMP\n21 CFR Part 11\nEU Annex 11\n"
            "GAMP 5\nRegulatory Compliance\nIT Compliance\nRisk Assessment\nStakeholder Management\n"
            "Cloud Qualification\nAutomation\nTOSCA\nValgenesis\nIQ\nOQ\nPQ\nURS"
        ),
        "keywords": "CSV Lead, validation, GxP, compliance, 21CFR, HP ALM, Azure, EDMS, GAMP5",
    },
    # ── From CSV JD PDF – Analyst ─────────────────────────────────────────────
    {
        "title": "CSV Analyst",
        "role": "CSV Analyst",
        "category": "Validation Roles",
        "responsibilities": (
            "Perform validation activities for systems/applications from various business areas\n"
            "Develop plans, execute, and deliver services with acceptable service evaluations\n"
            "Perform CAPA and RCA using 5 x Why approach\n"
            "Prepare validation documents including URS, IRA, FRA, IQ, OQ, PQ protocols\n"
            "Responsible for Software Validation, test script preparation, execution and review\n"
            "Ensure Risk Management Plan aligned to ICH Q9\n"
            "Perform 21CFR Part 11 Gap assessment for Computerized systems"
        ),
        "requirements": (
            "M.Sc. or B.Sc. in Analytical Chemistry, Life Sciences, or related field\n"
            "3-5 years of experience in CSV or Computer System Validation\n"
            "Good knowledge in GxP, 21CFR PART 11, EU Annex 11, GAMP 5\n"
            "Experience in V-model, Agile, Waterfall methodologies\n"
            "Knowledge of EDMS systems such as Valgenesis\n"
            "Experience with LIMS software\n"
            "Understanding of equipment qualification (IQ, OQ, PQ)"
        ),
        "skills": (
            "CSV\nComputer System Validation\nGAMP 5\n21 CFR Part 11\nEU Annex 11\n"
            "GxP\nGMP\nGDP\nIQ\nOQ\nPQ\nURS\nDQ\nFRA\nIRA\nCAPA\nRisk Assessment\n"
            "EDMS\nValgenesis\nLIMS\nSAP\nHP ALM\nAgile\nWaterfall\nV-model\nICH Q9\n"
            "Equipment Qualification\nSoftware Validation"
        ),
        "keywords": "CSV, GAMP5, validation, pharmaceutical, compliance, regulatory, 21CFR, GxP, LIMS",
    },
    # ── From CSV JD PDF – Compliance ──────────────────────────────────────────
    {
        "title": "IT Compliance Consultant (CSV)",
        "role": "Validation Engineer",
        "category": "Validation Roles",
        "responsibilities": (
            "Provide compliance consulting for IT projects and deliver CSV compliance deliverables\n"
            "Perform stakeholder management involving multiple stakeholders\n"
            "Manage multiple concurrent tasks with flexibility\n"
            "Perform IT risk assessments\n"
            "Execute risk-based validation strategy and deliverables\n"
            "Participate in audits and inspections related to the projects"
        ),
        "requirements": (
            "Experience in IT Compliance, CSV, and Application Support\n"
            "CSV, Regulatory Compliance, 21 CFR Part 11 expertise\n"
            "Strong IT compliance and GxP knowledge\n"
            "Written and Verbal communication skills"
        ),
        "skills": (
            "CSV\nIT Compliance\nRegulatory Compliance\n21 CFR Part 11\nGxP\n"
            "Application Support\nStakeholder Management\nRisk Assessment\nCAPA\nRCA\n"
            "Audit\nValidation\nComputer System Validation\nGMP"
        ),
        "keywords": "CSV, IT compliance, 21CFR, GxP, regulatory, validation, audit",
    },
    # ── From CSV JD PDF – Testing / TOSCA ────────────────────────────────────
    {
        "title": "CSV Testing Engineer (TOSCA/LeapWork)",
        "role": "Tosca Engineer",
        "category": "Validation Roles",
        "responsibilities": (
            "Creation and execution of manual and automated scripts using LeapWork/TOSCA\n"
            "End-to-end testing including requirement analysis, test data, test case creation, execution, defect management\n"
            "Organizing and monitoring defect management process and regression tests\n"
            "Define and implement test automation strategy including roadmap and tools\n"
            "Act as SME around Testing across business units\n"
            "Ensure quality of deliverables in alignment with stakeholder expectations"
        ),
        "requirements": (
            "2+ years of experience in CSV and Testing\n"
            "TOSCA or LeapWork experience required\n"
            "CSV, STLC Principles\n"
            "Regulatory Compliance, 21 CFR Part 11\n"
            "Azure DevOps, TIMS experience\n"
            "Experience in Agile Set-up\n"
            "ITIL understanding"
        ),
        "skills": (
            "TOSCA\nLeapWork\nLeap Work\nTest Automation\nSTLC\nCSV\n"
            "21 CFR Part 11\nAzure DevOps\nTIMS\nITIL\nAgile\n"
            "Regression Testing\nDefect Management\nTest Case Creation\nTest Scripts\n"
            "Regulatory Compliance\nComputer System Validation\nManual Testing"
        ),
        "keywords": "TOSCA, LeapWork, test automation, CSV, testing, 21CFR, Azure DevOps, STLC",
    },
    # ── From Product Registration Specialist PDF ──────────────────────────────
    {
        "title": "Product Registration Specialist",
        "role": "Product Registration Specialist",
        "category": "Regulatory Affairs",
        "responsibilities": (
            "Lead global regulatory strategies for product registrations, renewals, and change controls\n"
            "Prepare, review, and submit regulatory dossiers under MDR, IVDR, and regional frameworks\n"
            "Manage end-to-end submission lifecycle including gap analysis and response coordination\n"
            "Partner with R&D, Quality, and Manufacturing teams\n"
            "Coordinate with Notified Bodies and health authorities\n"
            "Maintain regulatory documentation and databases\n"
            "Support labeling reviews and post-market compliance activities"
        ),
        "requirements": (
            "Minimum 12 years of experience in global Regulatory Affairs for IVDs and Medical Devices\n"
            "Proven track record in hands-on submission execution\n"
            "Strong understanding of MDR, IVDR, FDA 510(k)/PMA, and regional product registration pathways\n"
            "Degree in Life Sciences, Regulatory Affairs, or Biomedical Engineering"
        ),
        "skills": (
            "MDR\nIVDR\nFDA 510k\nPMA\nRegulatory Affairs\nIVD\nMedical Device\n"
            "Notified Bodies\nRegulatory Dossiers\nGap Analysis\nChange Control\n"
            "Post-market Compliance\nLabeling\nDocument Control\nRegulatory Strategy\n"
            "Submission Management\nLife Sciences\nBiomedical Engineering\nAPAC\nLATAM\nEMEA"
        ),
        "keywords": "regulatory affairs, MDR, IVDR, FDA, medical device, IVD, submission, registration",
    },
    # ── Generic Validation Roles ──────────────────────────────────────────────
    {
        "title": "Validation Engineer",
        "role": "Validation Engineer",
        "category": "Validation Roles",
        "responsibilities": (
            "Execute validation protocols and prepare validation deliverables\n"
            "Prepare and review IQ, OQ, PQ protocols\n"
            "Author validation plans, reports, and traceability matrices\n"
            "Support CAPA and deviation investigations\n"
            "Participate in risk assessments and risk management activities\n"
            "Maintain compliance with GxP guidelines"
        ),
        "requirements": (
            "B.Sc. or M.Sc. in Life Sciences, Engineering, or related field\n"
            "2-4 years of experience in validation\n"
            "Knowledge of 21 CFR Part 11 and GAMP 5\n"
            "Experience with equipment and software validation"
        ),
        "skills": (
            "Validation\nIQ\nOQ\nPQ\nURS\nFRA\nGAMP 5\n21 CFR Part 11\n"
            "GxP\nGMP\nCAPA\nRisk Management\nEquipment Qualification\n"
            "Software Validation\nEDMS\nCompliance"
        ),
        "keywords": "validation engineer, GxP, IQ OQ PQ, GAMP5, compliance, pharmaceutical",
    },
    {
        "title": "Automation Engineer (Validation)",
        "role": "Automation Engineer",
        "category": "Validation Roles",
        "responsibilities": (
            "Design, develop, and maintain test automation frameworks\n"
            "Create and execute automated test scripts\n"
            "Integrate automation into CI/CD pipelines\n"
            "Perform root cause analysis and defect management\n"
            "Support regulatory compliance activities"
        ),
        "requirements": (
            "3+ years of test automation experience\n"
            "Proficiency in TOSCA, Selenium, or similar tools\n"
            "Knowledge of CSV and regulatory compliance\n"
            "Good scripting skills"
        ),
        "skills": (
            "Test Automation\nTOSCA\nSelenium\nCI/CD\nSTLC\nCSV\n"
            "Agile\nDefect Management\nRegression Testing\nScripting\n"
            "Python\nJava\nJavaScript\nJira\nAzure DevOps"
        ),
        "keywords": "automation engineer, TOSCA, Selenium, test automation, CI/CD, scripting",
    },
    # ── IT Roles ──────────────────────────────────────────────────────────────
    {
        "title": "Software Developer",
        "role": "Software Developer",
        "category": "IT Roles",
        "responsibilities": (
            "Design, develop, test, and maintain software applications\n"
            "Write clean, efficient, and well-documented code\n"
            "Collaborate with cross-functional teams to define and implement features\n"
            "Participate in code reviews and provide constructive feedback\n"
            "Troubleshoot, debug, and upgrade existing systems"
        ),
        "requirements": (
            "Bachelor's degree in Computer Science or related field\n"
            "3+ years of software development experience\n"
            "Proficiency in one or more programming languages\n"
            "Experience with databases and SQL\n"
            "Knowledge of software development lifecycle (SDLC)"
        ),
        "skills": (
            "Java\nPython\nJavaScript\nC#\nSQL\nGit\nREST API\n"
            "Microservices\nAgile\nScrum\nSDLC\nOOP\nUnit Testing\nCI/CD\nDocker"
        ),
        "keywords": "software developer, programming, Java, Python, JavaScript, REST API, Agile, SDLC",
    },
    {
        "title": "Full Stack Developer",
        "role": "Full Stack Developer",
        "category": "IT Roles",
        "responsibilities": (
            "Develop both front-end and back-end components of web applications\n"
            "Design and implement RESTful APIs and microservices\n"
            "Create responsive and user-friendly UI components\n"
            "Manage databases and ensure data integrity"
        ),
        "requirements": (
            "3+ years of full stack development experience\n"
            "Proficiency in frontend frameworks such as React, Angular, or Vue\n"
            "Backend experience with Node.js, Python, Java, or similar\n"
            "Experience with relational and NoSQL databases"
        ),
        "skills": (
            "React\nAngular\nVue.js\nNode.js\nPython\nJava\nJavaScript\nTypeScript\n"
            "HTML\nCSS\nREST API\nGraphQL\nSQL\nMongoDB\nPostgreSQL\nDocker\nKubernetes\nGit\nCI/CD"
        ),
        "keywords": "full stack, React, Node.js, REST API, JavaScript, frontend, backend, TypeScript",
    },
    {
        "title": "DevOps Engineer",
        "role": "DevOps Engineer",
        "category": "IT Roles",
        "responsibilities": (
            "Design and implement CI/CD pipelines\n"
            "Manage cloud infrastructure and containerized environments\n"
            "Monitor system performance and ensure availability\n"
            "Automate infrastructure provisioning using IaC tools"
        ),
        "requirements": (
            "3+ years of DevOps or SRE experience\n"
            "Experience with cloud platforms such as AWS, Azure, or GCP\n"
            "Proficiency in Docker and Kubernetes\n"
            "Infrastructure as Code experience with Terraform or Ansible"
        ),
        "skills": (
            "Docker\nKubernetes\nAWS\nAzure\nGCP\nCI/CD\nJenkins\nGitLab CI\nGitHub Actions\n"
            "Terraform\nAnsible\nPython\nBash\nLinux\nMonitoring\nPrometheus\nGrafana"
        ),
        "keywords": "DevOps, CI/CD, Docker, Kubernetes, AWS, Azure, Terraform, Jenkins, automation",
    },
    {
        "title": "QA Engineer",
        "role": "QA Engineer",
        "category": "IT Roles",
        "responsibilities": (
            "Design, develop, and execute test cases and test plans\n"
            "Perform manual and automated testing of software applications\n"
            "Identify, document, and track defects\n"
            "Implement and maintain test automation frameworks"
        ),
        "requirements": (
            "2+ years of QA testing experience\n"
            "Experience with test automation frameworks such as Selenium, Cypress, or TOSCA\n"
            "Knowledge of STLC and software testing methodologies\n"
            "Proficiency in defect tracking tools such as Jira or Azure DevOps"
        ),
        "skills": (
            "Selenium\nCypress\nTOSCA\nJira\nAzure DevOps\nSTLC\n"
            "Test Automation\nManual Testing\nAPI Testing\nPostman\n"
            "Regression Testing\nPerformance Testing\nDefect Management\nAgile\nSQL"
        ),
        "keywords": "QA engineer, testing, Selenium, automation, Jira, STLC, defect management, Agile",
    },
    {
        "title": "Data Analyst",
        "role": "Data Analyst",
        "category": "IT Roles",
        "responsibilities": (
            "Collect, process, and analyze large datasets to extract meaningful insights\n"
            "Create dashboards and visualizations using BI tools\n"
            "Prepare reports and presentations for business stakeholders\n"
            "Ensure data quality and accuracy"
        ),
        "requirements": (
            "Bachelor's degree in Statistics, Mathematics, Computer Science, or related field\n"
            "2+ years of data analysis experience\n"
            "Proficiency in SQL and data querying\n"
            "Experience with BI tools such as Power BI or Tableau\n"
            "Python or R for data analysis"
        ),
        "skills": (
            "SQL\nPython\nR\nPower BI\nTableau\nExcel\nData Analysis\n"
            "Statistics\nVisualization\nETL\nPandas\nNumPy\nBusiness Intelligence\nKPI\nDashboard"
        ),
        "keywords": "data analyst, SQL, Power BI, Tableau, Python, statistics, visualization, BI",
    },
    {
        "title": "AI/ML Engineer",
        "role": "AI/ML Engineer",
        "category": "IT Roles",
        "responsibilities": (
            "Design, develop, and deploy machine learning models\n"
            "Work with large datasets to train and evaluate ML models\n"
            "Implement MLOps pipelines for model deployment\n"
            "Research and apply state-of-the-art ML techniques"
        ),
        "requirements": (
            "Bachelor's or Master's degree in Computer Science, Mathematics, or related field\n"
            "3+ years of ML/AI development experience\n"
            "Proficiency in Python and ML frameworks\n"
            "Strong knowledge of statistics and mathematics"
        ),
        "skills": (
            "Python\nTensorFlow\nPyTorch\nMachine Learning\nDeep Learning\nNLP\nScikit-learn\n"
            "Pandas\nNumPy\nMLOps\nDocker\nKubernetes\nAWS\nStatistics\nSpark\nSQL"
        ),
        "keywords": "AI, ML, machine learning, deep learning, Python, TensorFlow, PyTorch, NLP, MLOps",
    },
    {
        "title": "SAP Consultant",
        "role": "SAP Consultant",
        "category": "IT Roles",
        "responsibilities": (
            "Implement, configure, and support SAP modules\n"
            "Analyze business requirements and translate them into SAP solutions\n"
            "Perform system integration testing and user acceptance testing\n"
            "Provide end-user training and documentation"
        ),
        "requirements": (
            "3+ years of SAP consulting experience\n"
            "Expertise in one or more SAP modules such as S/4HANA, MM, SD, FICO\n"
            "Good understanding of business processes\n"
            "Strong analytical and problem-solving skills"
        ),
        "skills": (
            "SAP\nS/4HANA\nSAP MM\nSAP SD\nSAP FICO\nSAP HR\nSAP ERP\n"
            "ABAP\nFIORI\nSAP BTP\nBusiness Analysis\nIntegration Testing\nUAT\n"
            "Data Migration\nSAP Activate\nAgile"
        ),
        "keywords": "SAP, S/4HANA, SAP consultant, ERP, ABAP, MM, SD, FICO, implementation",
    },
    {
        "title": "Design Control Consultant (IVD)",
        "role": "Design Control Consultant",
        "category": "Regulatory Affairs",
        "responsibilities": (
            "Review and strengthen Design Control processes in accordance with FDA 21 CFR Part 820, ISO 13485, and EU IVDR (2017/746)\n"
            "Ensure compliance across all Design and Development lifecycle stages (planning, inputs, outputs, verification, validation, transfer, DHF)\n"
            "Conduct detailed gap assessments against regulatory requirements and internal procedures\n"
            "Identify compliance risks and deficiencies in Design History Files (DHF) and related documentation\n"
            "Develop and execute remediation strategies and action plans\n"
            "Harmonize design control processes across multiple product lines (especially IVDs) and global/regional regulatory frameworks\n"
            "Standardize SOPs, templates, and workflows to ensure consistency and scalability\n"
            "Deliver gap assessment report with actionable recommendations\n"
            "Produce audit-ready documentation including remediated and compliant Design Control processes"
        ),
        "requirements": (
            "Bachelor's or Master's degree in Engineering, Life Sciences, or related field\n"
            "5 to 10 years of experience in Medical Devices with strong exposure to IVD products\n"
            "Proven experience in Design Control implementation and remediation\n"
            "Proven experience conducting gap assessments and process harmonization across global standards\n"
            "Strong knowledge of FDA 21 CFR Part 820 / QSR, ISO 13485, and EU IVDR requirements\n"
            "Experience with Design History Files (DHF) and lifecycle management"
        ),
        "skills": (
            "Design Control\nIVD\nIn Vitro Diagnostics\nFDA 21 CFR Part 820\nISO 13485\n"
            "EU IVDR\nQSR\nDHF\nDesign History File\nGap Assessment\nProcess Harmonization\n"
            "SOPs\nMedical Devices\nRegulatory Compliance\nRisk Assessment\nIQ OQ PQ\n"
            "Validation\nDocumentation\nAudit Readiness\nSDLC\nGxP"
        ),
        "keywords": (
            "Design Control, IVD, In Vitro Diagnostics, FDA 21 CFR Part 820, ISO 13485, "
            "EU IVDR, DHF, gap assessment, design history file, SOPs, medical devices, QSR, "
            "regulatory compliance, process harmonization, validation, audit ready"
        ),
    },
    # ── From IFU Team Lead JD PDF ─────────────────────────────────────────────
    {
        "title": "IFU Team Lead",
        "role": "IFU Team Lead",
        "category": "Regulatory Affairs",
        "responsibilities": (
            "Lead and mentor a team of IFU writers, manuals writers, and label creators\n"
            "Represent the team in cross-functional projects and management meetings\n"
            "Drive team activities and deliverables within defined scope, ensuring timelines and quality\n"
            "Create, revise, and publish high-quality Instructions for Use (IFU) and User Manuals\n"
            "Ensure compliance with QSR, ISO, IVDD/IVDR, FDA 21 CFR Part 11, ISO 13485\n"
            "Manage and coordinate translations to ensure accuracy and consistency\n"
            "Collaborate with global teams: Product Development, Regulatory Affairs, QA, Marketing\n"
            "Continuously improve documentation standards, templates, and processes\n"
            "Support quality investigations including deviations, CAPAs, and complaints"
        ),
        "requirements": (
            "Bachelor's or Master's degree in Technical Communication, English, Journalism, or Life Sciences\n"
            "5-8 years of technical writing experience within life sciences, medical device, or IVD\n"
            "Proven expertise in developing IFUs, DFUs, product labels, and user manuals\n"
            "Strong understanding of regulated content development, usability, and risk communication\n"
            "Experience working with translations and simplified English\n"
            "Familiarity with visual and multimedia tools (e.g., Adobe Illustrator)\n"
            "Prior training in QSR; experience with EU IVDR / MDR is a plus\n"
            "Excellent English speaking, writing, and editing skills"
        ),
        "skills": (
            "IFU\nUser Manuals\nDFU\nTechnical Writing\nTeam Lead\nMentoring\n"
            "QSR\nISO 13485\nFDA 21 CFR Part 11\nEU IVDR\nEU MDR\nIVDD\n"
            "Translation Management\nSimplified English\nCMS\nCCMS\nDITA\nXML\n"
            "Adobe Illustrator\nLabeling\nDocumentation\nCAPA\nDeviation\n"
            "Medical Device\nIVD\nRegulatory Compliance\nLife Sciences\nStyle Guide"
        ),
        "keywords": "IFU Team Lead, IFU writer, technical writing, IVDR, MDR, medical device, regulatory, QMS, labeling, team lead",
    },
    # ── From Labeling JD PDF ──────────────────────────────────────────────────
    {
        "title": "Labeling Specialist",
        "role": "Labeling Specialist",
        "category": "Regulatory Affairs",
        "responsibilities": (
            "Create and update labels and box prints in alignment with procedures, regulatory expectations, and market needs\n"
            "Ensure documentation complies with QSR, ISO, and internal quality and regulatory standards\n"
            "Maintain adherence to company style guides, templates, and quality management systems (QMS)\n"
            "Collaborate with Product Development, Regulatory Affairs, Quality Assurance, Marketing, and Global Operations\n"
            "Drive and manage activities within the assigned area of responsibility\n"
            "Contribute to continuous improvement of documentation standards and templates\n"
            "Support Deviation, CAPA, and complaint investigations by providing documentation expertise"
        ),
        "requirements": (
            "University degree in the life science field and/or relevant work experience\n"
            "Excellent English speaking, writing, and editing skills\n"
            "2-3 years of label creation experience in life sciences, medical devices, or IVD industry\n"
            "Familiarity with visual and multimedia tools (e.g., Adobe InDesign, Illustrator)\n"
            "Proven ability to manage multiple documentation projects in a global, regulated environment\n"
            "Experience with EU IVDR documentation requirements"
        ),
        "skills": (
            "Labeling\nLabel Design\nBox Print\nAdobe InDesign\nAdobe Illustrator\n"
            "QSR\nISO\nQMS\nEU IVDR\nEU MDR\nRegulatory Compliance\nDocumentation\n"
            "Style Guide\nCAPA\nDeviation\nMedical Device\nIVD\nLife Sciences\n"
            "Technical Writing\nMultilingual\nTranslation\nCross-functional Collaboration"
        ),
        "keywords": "labeling, label design, IVD, medical device, IVDR, regulatory, QMS, Adobe InDesign, documentation",
    },
]


def seed_jds(conn):
    existing_titles = {
        row["title"]
        for row in conn.execute("SELECT title FROM job_description").fetchall()
    }
    for jd in PREDEFINED_JDS:
        if jd["title"] not in existing_titles:
            conn.execute(
                """
                INSERT INTO job_description
                    (title, role, category, responsibilities, requirements, skills, keywords)
                VALUES
                    (%(title)s, %(role)s, %(category)s, %(responsibilities)s,
                     %(requirements)s, %(skills)s, %(keywords)s)
                """,
                jd,
            )


# ── Matching Algorithm ────────────────────────────────────────────────────────

def _normalize(text):
    return re.sub(r'\s+', ' ', (text or '').lower()).strip()


def _parse_jd_items(skills_text, requirements_text, keywords_text):
    combined = "\n".join(filter(None, [
        str(skills_text or ''), str(requirements_text or ''), str(keywords_text or '')
    ]))
    seen, items = set(), []
    for line in combined.splitlines():
        line = re.sub(r'^[•\-–*►◆▸▪\d\.\)\s]+', '', line).strip()
        if not line:
            continue
        for part in re.split(r'[,;]', line):
            part = part.strip()
            key = _normalize(part)
            if 2 <= len(part) <= 80 and key not in seen:
                seen.add(key)
                items.append(part)
    return items


def _skill_matches(skill, corpus):
    skill_n = _normalize(skill)
    corpus_n = _normalize(corpus)
    if skill_n in corpus_n:
        return True
    words = [w for w in skill_n.split() if len(w) > 3]
    if len(words) >= 2 and all(w in corpus_n for w in words[:2]):
        return True
    return False


def _extract_years_required(text):
    for pat in [
        r'minimum\s+(\d+)\s*years?',
        r'at\s+least\s+(\d+)\s*years?',
        r'(\d+)\+\s*years?',
        r'(\d+)\s*[-–]\s*\d+\s*years?',
        r'(\d+)\s*years?\s+of\s+experience',
    ]:
        m = re.search(pat, text, re.I)
        if m:
            try:
                return int(m.group(1))
            except Exception:
                pass
    return None


def _estimate_exp_years(experience_text):
    from datetime import datetime as _dt2
    if not experience_text:
        return None
    current_year = _dt2.now().year
    ranges = re.findall(
        r'(\d{4})\s*[-–—to]+\s*(present|current|till\s*date|\d{4})',
        experience_text, re.I
    )
    total = 0
    for start, end in ranges:
        try:
            s = int(start)
            e = current_year if re.match(r'present|current|till', end, re.I) else int(end)
            if 1990 <= s <= current_year and e >= s:
                total += e - s
        except Exception:
            pass
    return total if total > 0 else None


def calculate_match_score(resume_d, jd_d):
    corpus = " ".join(str(resume_d.get(k) or '') for k in
                      ['skills', 'experience', 'summary', 'projects', 'certifications', 'education', 'title'])
    items = _parse_jd_items(jd_d.get('skills'), jd_d.get('requirements'), jd_d.get('keywords'))
    if not items:
        return {
            'match_percentage': 0, 'skills_match_percentage': 0, 'experience_match_percentage': 0,
            'matched_skills': [], 'missing_skills': [], 'total_jd_requirements': 0,
            'matched_count': 0, 'missing_count': 0, 'match_level': 'No Data',
            'match_level_color': 'gray', 'strong_areas': [], 'weak_areas': [],
            'experience_note': 'No skills data found in this JD.',
        }
    matched, missing = [], []
    for item in items:
        (matched if _skill_matches(item, corpus) else missing).append(item)
    skills_pct = len(matched) / len(items) * 100
    jd_exp_text = " ".join(str(jd_d.get(k) or '') for k in ['requirements', 'responsibilities'])
    jd_yrs = _extract_years_required(jd_exp_text)
    res_yrs = _estimate_exp_years(str(resume_d.get('experience') or ''))
    if jd_yrs and res_yrs is not None:
        exp_pct = min(100.0, res_yrs / jd_yrs * 100)
        exp_note = f"Estimated {res_yrs} yr(s) in resume; JD requires {jd_yrs}+ yr(s)"
    elif jd_yrs:
        exp_pct = 60.0
        exp_note = f"JD requires {jd_yrs}+ yr(s); experience timeline not detected in resume"
    else:
        exp_pct = 80.0
        exp_note = "No specific experience years requirement stated in JD"
    final = min(100.0, skills_pct * 0.80 + exp_pct * 0.20)
    if final >= 80:
        level, color = 'Strong Match', 'green'
    elif final >= 60:
        level, color = 'Good Match', 'blue'
    elif final >= 40:
        level, color = 'Partial Match', 'orange'
    else:
        level, color = 'Low Match', 'red'
    return {
        'match_percentage': round(final),
        'skills_match_percentage': round(skills_pct),
        'experience_match_percentage': round(exp_pct),
        'matched_skills': matched,
        'missing_skills': missing,
        'total_jd_requirements': len(items),
        'matched_count': len(matched),
        'missing_count': len(missing),
        'match_level': level,
        'match_level_color': color,
        'strong_areas': matched[:12],
        'weak_areas': missing[:12],
        'experience_note': exp_note,
        'jd_years_required': jd_yrs,
        'resume_years_estimated': res_yrs,
    }


# ── JD Routes ─────────────────────────────────────────────────────────────────

@app.route("/jd-management")
def jd_management():
    with db_conn() as conn:
        ensure_jd_table(conn)
        seed_jds(conn)
        jds = conn.execute(
            "SELECT id, title, role, category, created_at"
            " FROM job_description ORDER BY category, title"
        ).fetchall()
    return render_template("jd_management.html", jds=list(jds))


@app.route("/jd/add", methods=["GET", "POST"])
def jd_add():
    if request.method == "POST":
        import time as _t
        data = {k: request.form.get(k, "").strip() for k in
                ["title", "role", "category", "responsibilities", "requirements", "skills", "keywords"]}
        data["jd_file"] = None
        uploaded = request.files.get("jd_file")
        if uploaded and uploaded.filename and allowed_file(uploaded.filename):
            name = secure_filename(uploaded.filename)
            ext = name.rsplit(".", 1)[1].lower()
            save_name = f"{slugify(Path(name).stem)}-{int(_t.time())}.{ext}"
            dest = JD_UPLOAD_FOLDER / save_name
            uploaded.seek(0)
            uploaded.save(str(dest))
            data["jd_file"] = save_name
            if not data["responsibilities"]:
                try:
                    txt = extract_resume_text(dest, ext)
                    data["responsibilities"] = (txt or "")[:4000]
                except Exception:
                    pass
        with db_conn() as conn:
            ensure_jd_table(conn)
            conn.execute(
                """
                INSERT INTO job_description
                    (title, role, category, responsibilities, requirements, skills, keywords, jd_file)
                VALUES
                    (%(title)s, %(role)s, %(category)s, %(responsibilities)s,
                     %(requirements)s, %(skills)s, %(keywords)s, %(jd_file)s)
                """,
                data,
            )
        flash(f"Job Description '{data['title']}' added.", "success")
        return redirect(url_for("jd_management"))
    return render_template("jd_form.html", jd=None, all_roles=ALL_JD_ROLES)


@app.route("/jd/<int:jd_id>")
def jd_detail(jd_id):
    with db_conn() as conn:
        ensure_jd_table(conn)
        jd = conn.execute(
            "SELECT * FROM job_description WHERE id = %s", (jd_id,)
        ).fetchone()
        if not jd:
            return "Job Description not found", 404
        resumes = conn.execute(
            "SELECT id, full_name, title FROM resume ORDER BY updated_at DESC LIMIT 30"
        ).fetchall()
    return render_template("jd_detail.html", jd=dict(jd), resumes=list(resumes))


@app.route("/jd/<int:jd_id>/download-pdf")
def download_jd_pdf(jd_id):
    """Download JD as PDF."""
    with db_conn() as conn:
        ensure_jd_table(conn)
        jd = conn.execute(
            "SELECT * FROM job_description WHERE id = %s", (jd_id,)
        ).fetchone()

    if not jd:
        return "Job Description not found", 404

    jd_dict = dict(jd)

    # Generate PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          topMargin=0.5*inch, bottomMargin=0.5*inch,
                          leftMargin=0.75*inch, rightMargin=0.75*inch)
    story = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#1e293b',
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    story.append(Paragraph(jd_dict.get('title', 'Job Description'), title_style))

    # Category & Role
    if jd_dict.get('category') or jd_dict.get('role'):
        meta_parts = []
        if jd_dict.get('category'):
            meta_parts.append(f"Category: {jd_dict['category']}")
        if jd_dict.get('role'):
            meta_parts.append(f"Role: {jd_dict['role']}")
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#64748b',
            spaceAfter=16,
            alignment=TA_CENTER
        )
        story.append(Paragraph(' | '.join(meta_parts), meta_style))

    # Section styling
    section_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor='#1e293b',
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )

    # Responsibilities
    if jd_dict.get('responsibilities'):
        story.append(Paragraph('RESPONSIBILITIES', section_style))
        resp_text = jd_dict['responsibilities'].replace('\n', '<br/>')
        story.append(Paragraph(resp_text, body_style))

    # Requirements
    if jd_dict.get('requirements'):
        story.append(Paragraph('REQUIREMENTS / QUALIFICATIONS', section_style))
        req_text = jd_dict['requirements'].replace('\n', '<br/>')
        story.append(Paragraph(req_text, body_style))

    # Skills
    if jd_dict.get('skills'):
        story.append(Paragraph('REQUIRED SKILLS', section_style))
        skills_list = [s.strip() for s in jd_dict['skills'].split('\n') if s.strip()]
        skills_text = ' • '.join(skills_list)
        story.append(Paragraph(skills_text, body_style))

    # Keywords
    if jd_dict.get('keywords'):
        story.append(Paragraph('KEYWORDS', section_style))
        story.append(Paragraph(jd_dict['keywords'], body_style))

    # Build PDF
    doc.build(story)
    buffer.seek(0)

    filename = f"{jd_dict.get('title', 'JD').replace(' ', '_')}.pdf"
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=filename
    )


@app.route("/jd/<int:jd_id>/extract")
def extract_jd_data(jd_id):
    """Extract JD data as JSON."""
    with db_conn() as conn:
        ensure_jd_table(conn)
        jd = conn.execute(
            "SELECT * FROM job_description WHERE id = %s", (jd_id,)
        ).fetchone()

    if not jd:
        return jsonify({"error": "Job Description not found"}), 404

    jd_dict = dict(jd)
    # Remove internal fields
    jd_dict.pop('id', None)
    jd_dict.pop('created_at', None)
    jd_dict.pop('updated_at', None)
    jd_dict.pop('jd_file', None)

    filename = f"{jd_dict.get('title', 'jd').replace(' ', '_')}_details.json"

    return send_file(
        BytesIO(json.dumps(jd_dict, indent=2).encode()),
        mimetype='application/json',
        as_attachment=True,
        download_name=filename
    )


@app.route("/jd/<int:jd_id>/edit", methods=["GET", "POST"])
def jd_edit(jd_id):
    with db_conn() as conn:
        ensure_jd_table(conn)
        if request.method == "POST":
            data = {k: request.form.get(k, "").strip() for k in
                    ["title", "role", "category", "responsibilities", "requirements", "skills", "keywords"]}
            data["id"] = jd_id
            conn.execute(
                """
                UPDATE job_description SET
                    title=%(title)s, role=%(role)s, category=%(category)s,
                    responsibilities=%(responsibilities)s, requirements=%(requirements)s,
                    skills=%(skills)s, keywords=%(keywords)s, updated_at=NOW()
                WHERE id=%(id)s
                """,
                data,
            )
            flash("Job Description updated.", "success")
            return redirect(url_for("jd_detail", jd_id=jd_id))
        jd = conn.execute(
            "SELECT * FROM job_description WHERE id = %s", (jd_id,)
        ).fetchone()
        if not jd:
            return "Job Description not found", 404
    return render_template("jd_form.html", jd=dict(jd), all_roles=ALL_JD_ROLES)


@app.route("/jd/<int:jd_id>/delete", methods=["POST"])
def jd_delete(jd_id):
    with db_conn() as conn:
        ensure_jd_table(conn)
        row = conn.execute(
            "SELECT jd_file FROM job_description WHERE id = %s", (jd_id,)
        ).fetchone()
        if row and row["jd_file"]:
            (JD_UPLOAD_FOLDER / (row["jd_file"] or "")).unlink(missing_ok=True)
        conn.execute("DELETE FROM job_description WHERE id = %s", (jd_id,))
    flash("Job Description deleted.", "success")
    return redirect(url_for("jd_management"))


@app.route("/uploads/jd/<path:filename>")
def jd_uploaded_file(filename):
    safe = secure_filename(filename)
    ext = safe.rsplit(".", 1)[-1].lower() if "." in safe else ""
    if ext not in ALLOWED_EXTENSIONS:
        abort(403)
    return send_from_directory(str(JD_UPLOAD_FOLDER), safe)


@app.route("/compare/<int:resume_id>")
def compare_select_jd(resume_id):
    with db_conn() as conn:
        ensure_jd_table(conn)
        seed_jds(conn)
        resume = conn.execute(
            "SELECT id, full_name, title FROM resume WHERE id = %s", (resume_id,)
        ).fetchone()
        if not resume:
            return "Resume not found", 404
        jds = conn.execute(
            "SELECT id, title, role, category FROM job_description ORDER BY category, title"
        ).fetchall()
    return render_template("compare_select.html", resume=dict(resume), jds=list(jds))


@app.route("/compare/<int:resume_id>/<int:jd_id>")
def compare_result(resume_id, jd_id):
    with db_conn() as conn:
        ensure_jd_table(conn)
        resume = conn.execute("SELECT * FROM resume WHERE id = %s", (resume_id,)).fetchone()
        jd = conn.execute("SELECT * FROM job_description WHERE id = %s", (jd_id,)).fetchone()
        if not resume or not jd:
            return "Resume or JD not found", 404
        all_jds = conn.execute(
            "SELECT id, title, category FROM job_description ORDER BY category, title"
        ).fetchall()
    result = calculate_match_score(dict(resume), dict(jd))
    return render_template(
        "compare_result.html",
        resume=dict(resume),
        jd=dict(jd),
        result=result,
        all_jds=list(all_jds),
    )


if __name__ == "__main__":
    app.run(debug=True, use_reloader=False, port=5001)
